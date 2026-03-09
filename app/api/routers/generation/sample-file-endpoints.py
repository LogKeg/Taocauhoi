"""
Sample file management endpoints.

GET /sample-folders - List available sample folders.
GET /sample-files - List sample files in a subject folder.
GET /sample-content - Get content of a sample file.
POST /parse-sample-urls - Parse sample questions from URLs.
"""
import io
from pathlib import Path
from typing import List

import httpx
from docx import Document
from fastapi import APIRouter

from app.core import ParseSamplesRequest
from app.parsers.docx import extract_docx_content
from app.services.generation import split_questions
from app.api.routers.generation.helpers import SAMPLE_DIR

router = APIRouter(tags=["generation"])


def _read_sample_file(path: Path) -> str:
    """Read content from a sample file."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = Document(str(path))
        return extract_docx_content(doc)
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_sample_url(url: str) -> str:
    """Read content from a URL."""
    suffix = Path(url.split("?")[0]).suffix.lower()
    with httpx.Client(timeout=30) as client:
        resp = client.get(url)
        resp.raise_for_status()
        data = resp.content
    if suffix == ".docx":
        doc = Document(io.BytesIO(data))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(lines)
    return data.decode("utf-8", errors="ignore")


@router.get("/sample-folders")
def list_sample_folders() -> dict:
    """List available sample folders."""
    if SAMPLE_DIR is None or not SAMPLE_DIR.exists():
        return {"folders": []}
    folders = [p.name for p in SAMPLE_DIR.iterdir() if p.is_dir()]
    folders.sort()
    return {"folders": folders}


@router.get("/sample-files")
def list_sample_files(subject: str) -> dict:
    """List sample files in a subject folder."""
    if SAMPLE_DIR is None or not subject or not SAMPLE_DIR.exists():
        return {"files": []}
    subject_dir = (SAMPLE_DIR / subject).resolve()
    if SAMPLE_DIR not in subject_dir.parents or not subject_dir.exists():
        return {"files": []}
    files = [
        p.name
        for p in subject_dir.iterdir()
        if p.is_file() and p.suffix.lower() in {".txt", ".docx", ".md"}
    ]
    files.sort()
    return {"files": files}


@router.get("/sample-content")
def sample_content(subject: str, filename: str) -> dict:
    """Get content of a sample file."""
    if not subject or not filename:
        return {"content": ""}
    if SAMPLE_DIR is None:
        return {"content": ""}
    subject_dir = (SAMPLE_DIR / subject).resolve()
    if not SAMPLE_DIR.exists() or SAMPLE_DIR not in subject_dir.parents:
        return {"content": ""}
    target = (subject_dir / filename).resolve()
    if subject_dir not in target.parents or not target.exists() or not target.is_file():
        return {"content": ""}
    content = _read_sample_file(target)
    return {"content": content}


@router.post("/parse-sample-urls")
def parse_sample_urls(payload: ParseSamplesRequest) -> dict:
    """Parse sample questions from URLs."""
    contents: List[str] = []
    for url in payload.urls:
        if not url:
            continue
        try:
            contents.append(_read_sample_url(url))
        except Exception:
            continue
    merged = "\n".join(contents)
    return {"content": merged, "samples": split_questions(merged)}
