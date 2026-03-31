"""
Professional exam PDF/DOCX exporter with school header, exam info,
answer sheet, and answer key.

POST /api/export-exam-pro - Export a complete, print-ready exam document.
"""
import io
import json
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fastapi import APIRouter, Request, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas

from app.api.routers.generation.helpers import _get_pdf_font, _wrap_text

router = APIRouter(tags=["generation"])


def _draw_header_pdf(c, width, height, info: dict, font_name: str):
    """Draw school header and exam info on PDF."""
    y = height - 40

    # Left column: School info
    c.setFont(font_name, 10)
    school = info.get("school_name", "")
    department = info.get("department", "")
    if school:
        c.drawString(50, y, school.upper())
        y -= 14
    if department:
        c.drawString(50, y, department)
        y -= 14

    # Right column: Exam code
    exam_code = info.get("exam_code", "")
    if exam_code:
        c.setFont(font_name, 11)
        code_text = f"Mã đề: {exam_code}"
        code_width = pdfmetrics.stringWidth(code_text, font_name, 11)
        c.drawString(width - 50 - code_width, height - 40, code_text)

    # Horizontal line
    y -= 6
    c.setStrokeColor(colors.HexColor("#333333"))
    c.setLineWidth(1.2)
    c.line(50, y, width - 50, y)
    y -= 20

    # Title block - centered
    title = info.get("title", "ĐỀ KIỂM TRA")
    c.setFont(font_name, 16)
    title_width = pdfmetrics.stringWidth(title.upper(), font_name, 16)
    c.drawString((width - title_width) / 2, y, title.upper())
    y -= 20

    # Subtitle line
    subtitle_parts = []
    if info.get("subject"):
        subtitle_parts.append(f"Môn: {info['subject']}")
    if info.get("grade"):
        subtitle_parts.append(f"Lớp: {info['grade']}")
    if info.get("duration"):
        subtitle_parts.append(f"Thời gian: {info['duration']} phút")

    if subtitle_parts:
        c.setFont(font_name, 11)
        subtitle = "  |  ".join(subtitle_parts)
        sub_width = pdfmetrics.stringWidth(subtitle, font_name, 11)
        c.drawString((width - sub_width) / 2, y, subtitle)
        y -= 16

    # Date
    date_str = info.get("date", "")
    if date_str:
        c.setFont(font_name, 10)
        date_width = pdfmetrics.stringWidth(date_str, font_name, 10)
        c.drawString((width - date_width) / 2, y, date_str)
        y -= 16

    # Thin line separator
    y -= 4
    c.setLineWidth(0.5)
    c.setStrokeColor(colors.HexColor("#cccccc"))
    c.line(50, y, width - 50, y)
    y -= 8

    # Student info fields
    c.setFont(font_name, 11)
    c.drawString(50, y, "Họ và tên: .................................................................")
    c.drawString(width / 2 + 20, y, "Lớp: ................   SBD: ................")
    y -= 24

    # Another separator
    c.setLineWidth(0.3)
    c.line(50, y, width - 50, y)
    y -= 16

    return y


def _draw_answer_sheet_pdf(c, y, width, height, total_questions: int, font_name: str, cols: int = 5):
    """Draw bubble answer sheet grid on PDF."""
    c.setFont(font_name, 11)
    c.drawString(50, y, "PHIẾU TRẢ LỜI")
    y -= 6

    c.setLineWidth(0.5)
    c.setStrokeColor(colors.HexColor("#999999"))
    c.line(50, y, width - 50, y)
    y -= 16

    c.setFont(font_name, 9)

    # Calculate grid layout
    rows_per_col = (total_questions + cols - 1) // cols
    col_width = (width - 100) / cols
    start_y = y

    for col in range(cols):
        x = 50 + col * col_width
        row_y = start_y

        for row in range(rows_per_col):
            q_num = col * rows_per_col + row + 1
            if q_num > total_questions:
                break

            c.setFont(font_name, 9)
            c.drawString(x, row_y, f"{q_num:2d}.")

            # Draw option circles A B C D
            for opt_idx, opt_label in enumerate("ABCD"):
                cx = x + 24 + opt_idx * 22
                cy = row_y + 3
                c.circle(cx, cy, 5, stroke=1, fill=0)
                c.setFont(font_name, 7)
                label_w = pdfmetrics.stringWidth(opt_label, font_name, 7)
                c.drawString(cx - label_w / 2, row_y + 0.5, opt_label)

            row_y -= 16

    y = start_y - rows_per_col * 16 - 10
    return y


def _draw_questions_pdf(c, y, width, height, questions: list, font_name: str, sections: list = None):
    """Draw questions on PDF with optional section headers."""
    from app.services.math_renderer import render_text_with_math, has_latex

    font_size = 11
    q_num = 0

    items = []
    if sections:
        for sec in sections:
            items.append({"type": "section", "label": sec.get("topic_label", ""), "difficulty": sec.get("difficulty_label", "")})
            for q in sec.get("questions", []):
                q_num += 1
                items.append({"type": "question", "num": q_num, "text": q})
    else:
        for q in questions:
            q_num += 1
            text = q if isinstance(q, str) else (q.get("content", "") or q.get("question", ""))
            items.append({"type": "question", "num": q_num, "text": text})

    for item in items:
        # Page break check
        if y < 80:
            c.showPage()
            c.setFont(font_name, font_size)
            y = height - 50

        if item["type"] == "section":
            # Section header
            c.setFont(font_name, 11)
            label = f"{item['label']}"
            if item.get("difficulty"):
                label += f" - {item['difficulty']}"
            c.setFillColor(colors.HexColor("#4f46e5"))
            c.drawString(50, y, label)
            c.setFillColor(colors.HexColor("#000000"))
            y -= 6
            c.setLineWidth(0.5)
            c.setStrokeColor(colors.HexColor("#a5b4fc"))
            c.line(50, y, 50 + pdfmetrics.stringWidth(label, font_name, 11) + 10, y)
            y -= 14
            c.setFont(font_name, font_size)
            continue

        # Question
        text = item["text"]
        lines = text.split('\n')

        # Question number + first line
        prefix = f"Câu {item['num']}. "
        c.setFont(font_name, font_size)
        c.drawString(50, y, prefix)
        x_offset = 50 + pdfmetrics.stringWidth(prefix, font_name, font_size)

        first_line = render_text_with_math(lines[0])
        remaining_width = width - x_offset - 50
        if pdfmetrics.stringWidth(first_line, font_name, font_size) <= remaining_width:
            c.drawString(x_offset, y, first_line)
        else:
            wrapped = _wrap_text(first_line, remaining_width, font_name, font_size)
            c.drawString(x_offset, y, wrapped[0])
            for part in wrapped[1:]:
                y -= 15
                if y < 60:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = height - 50
                c.drawString(50, y, part)
        y -= 17

        # Options
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            if y < 60:
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - 50
            opt_text = render_text_with_math(line)
            c.drawString(70, y, opt_text)
            y -= 15

        y -= 6  # Extra space between questions

    return y


def _draw_answer_key_pdf(c, y, width, height, answers: str, font_name: str):
    """Draw answer key section on a new page."""
    c.showPage()
    c.setFont(font_name, 14)
    y = height - 50

    c.drawString(50, y, "ĐÁP ÁN")
    y -= 8
    c.setLineWidth(1)
    c.setStrokeColor(colors.HexColor("#333333"))
    c.line(50, y, width - 50, y)
    y -= 20

    c.setFont(font_name, 11)
    if not answers:
        c.drawString(50, y, "(Không có đáp án)")
        return

    # Parse answers into grid
    answer_list = []
    for line in answers.split('\n'):
        line = line.strip()
        m = re.match(r'^(\d+)\.\s*([A-Da-d])', line, re.IGNORECASE)
        if m:
            answer_list.append((int(m.group(1)), m.group(2).upper()))

    if not answer_list:
        c.drawString(50, y, answers[:500])
        return

    # Draw in 5-column grid
    cols = 5
    rows_per_col = (len(answer_list) + cols - 1) // cols
    col_width = (width - 100) / cols

    for idx, (num, ans) in enumerate(answer_list):
        col = idx // rows_per_col
        row = idx % rows_per_col

        x = 50 + col * col_width
        cy = y - row * 18

        if cy < 60:
            break

        c.setFont(font_name, 11)
        c.drawString(x, cy, f"Câu {num}: {ans}")

    return y - rows_per_col * 18 - 10


@router.post("/api/export-exam-pro")
async def export_exam_professional(request: Request):
    """
    Export a professional, print-ready exam PDF/DOCX.

    Request body JSON:
    {
        "questions": [...],           # List of question strings or {content, options} objects
        "sections": [...],            # Optional: sections from matrix generation
        "answers": "1. A\n2. B...",   # Answer key text
        "format": "pdf" | "docx",
        "exam_info": {
            "school_name": "TRƯỜNG THPT ABC",
            "department": "TỔ TOÁN - TIN",
            "title": "ĐỀ KIỂM TRA GIỮA KỲ 1",
            "subject": "Toán",
            "grade": "12",
            "duration": "45",
            "exam_code": "132",
            "date": "Năm học 2025-2026",
            "include_answer_sheet": true,
            "include_answer_key": true
        }
    }
    """
    data = await request.json()
    questions = data.get("questions", [])
    sections = data.get("sections", [])
    answers = data.get("answers", "")
    fmt = data.get("format", "pdf")
    info = data.get("exam_info", {})

    if not questions and not sections:
        raise HTTPException(status_code=400, detail="Không có câu hỏi để xuất")

    # Count total questions
    total_q = len(questions) if questions else sum(len(s.get("questions", [])) for s in sections)

    if fmt == "pdf":
        font_name, _ = _get_pdf_font()
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Page 1: Header + questions
        y = _draw_header_pdf(c, width, height, info, font_name)

        # Optional: Answer sheet
        if info.get("include_answer_sheet", True) and total_q > 0:
            y = _draw_answer_sheet_pdf(c, y, width, height, total_q, font_name)
            y -= 10

        # Questions
        if y < 120:
            c.showPage()
            y = height - 50

        y = _draw_questions_pdf(c, y, width, height, questions, font_name, sections=sections if sections else None)

        # Optional: Answer key on last page
        if info.get("include_answer_key", True) and answers:
            _draw_answer_key_pdf(c, y, width, height, answers, font_name)

        c.save()
        buffer.seek(0)

        filename = f"de_thi_{info.get('exam_code', 'export')}.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    if fmt == "docx":
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Header section
        school = info.get("school_name", "")
        department = info.get("department", "")

        if school or department:
            header_table = doc.add_table(rows=1, cols=2)
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Left cell: school info
            left_cell = header_table.rows[0].cells[0]
            if school:
                p = left_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(school.upper())
                run.bold = True
                run.font.size = Pt(11)
            if department:
                p = left_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(department)
                run.font.size = Pt(10)

            # Right cell: exam code
            right_cell = header_table.rows[0].cells[1]
            exam_code = info.get("exam_code", "")
            if exam_code:
                p = right_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Mã đề: {exam_code}")
                run.bold = True
                run.font.size = Pt(12)

        # Title
        title = info.get("title", "ĐỀ KIỂM TRA")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(16)

        # Subtitle
        subtitle_parts = []
        if info.get("subject"):
            subtitle_parts.append(f"Môn: {info['subject']}")
        if info.get("grade"):
            subtitle_parts.append(f"Lớp: {info['grade']}")
        if info.get("duration"):
            subtitle_parts.append(f"Thời gian: {info['duration']} phút")
        if subtitle_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("  |  ".join(subtitle_parts))
            run.font.size = Pt(11)

        # Date
        date_str = info.get("date", "")
        if date_str:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(date_str)
            run.font.size = Pt(10)
            run.italic = True

        # Student info
        doc.add_paragraph("Họ và tên: ................................................................   Lớp: ................   SBD: ................")
        doc.add_paragraph("").paragraph_format.space_after = Pt(6)

        # Questions
        q_num = 0
        if sections:
            for sec in sections:
                p = doc.add_paragraph()
                run = p.add_run(f"{sec.get('topic_label', '')} - {sec.get('difficulty_label', '')}")
                run.bold = True
                run.font.color.rgb = RGBColor(79, 70, 229)

                for q_text in sec.get("questions", []):
                    q_num += 1
                    _add_question_docx(doc, q_num, q_text)
        else:
            for q in questions:
                q_num += 1
                text = q if isinstance(q, str) else (q.get("content", "") or q.get("question", ""))
                _add_question_docx(doc, q_num, text)

        # Answer key
        if info.get("include_answer_key", True) and answers:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run("ĐÁP ÁN")
            run.bold = True
            run.font.size = Pt(14)

            for line in answers.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"de_thi_{info.get('exam_code', 'export')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    raise HTTPException(status_code=400, detail="Format không hỗ trợ")


def _add_question_docx(doc, num: int, text: str):
    """Add a question to DOCX document."""
    lines = text.split('\n')
    p = doc.add_paragraph()
    run = p.add_run(f"Câu {num}. ")
    run.bold = True
    p.add_run(lines[0])

    for line in lines[1:]:
        line = line.strip()
        if line:
            opt_p = doc.add_paragraph()
            opt_p.paragraph_format.left_indent = Inches(0.4)
            opt_p.add_run(line)
