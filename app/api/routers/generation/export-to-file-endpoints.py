"""
Export endpoints for questions to various file formats.

POST /export - Export generated questions to txt/csv/docx/pdf.
POST /api/export-exam - Export exam questions with full content and options.
"""
import io
import re
from typing import List

from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, Form, HTTPException, Request
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas

from app.api.routers.generation.helpers import _get_pdf_font, _wrap_text

router = APIRouter(tags=["generation"])


def _draw_text_with_math_pdf(c, text: str, x: float, y: float, font_name: str, font_size: int, max_width: float) -> float:
    """Draw text with math formulas on PDF canvas. Returns the final x position."""
    from app.services.math_renderer import extract_latex_parts, latex_to_image, latex_to_unicode
    from reportlab.lib.utils import ImageReader

    parts = extract_latex_parts(text)
    current_x = x

    for part_text, is_latex in parts:
        if is_latex:
            # Try to render as image
            img_bytes = latex_to_image(part_text, dpi=150, fontsize=font_size + 2)
            if img_bytes:
                try:
                    img = ImageReader(io.BytesIO(img_bytes))
                    img_width, img_height = img.getSize()

                    # Scale down for PDF
                    scale = 0.48
                    draw_width = img_width * scale
                    draw_height = img_height * scale

                    # Check if fits in remaining width
                    if current_x + draw_width > x + max_width:
                        # Would overflow - use unicode fallback
                        unicode_text = latex_to_unicode(part_text)
                        c.drawString(current_x, y, unicode_text)
                        current_x += pdfmetrics.stringWidth(unicode_text, font_name, font_size)
                    else:
                        # Center vertically relative to text baseline
                        y_offset = (draw_height - font_size) / 3

                        c.drawImage(img, current_x, y - y_offset,
                                   width=draw_width, height=draw_height,
                                   mask='auto')
                        current_x += draw_width + 2
                except Exception:
                    # Fallback to Unicode
                    unicode_text = latex_to_unicode(part_text)
                    c.drawString(current_x, y, unicode_text)
                    current_x += pdfmetrics.stringWidth(unicode_text, font_name, font_size)
            else:
                # Fallback to Unicode
                unicode_text = latex_to_unicode(part_text)
                c.drawString(current_x, y, unicode_text)
                current_x += pdfmetrics.stringWidth(unicode_text, font_name, font_size)
        else:
            c.drawString(current_x, y, part_text)
            current_x += pdfmetrics.stringWidth(part_text, font_name, font_size)

    return current_x


def _add_text_with_math(paragraph, text: str):
    """Add text with math formulas to a Word paragraph using native OMML equations."""
    from app.services.math_renderer import extract_latex_parts, latex_to_unicode, add_math_to_paragraph

    parts = extract_latex_parts(text)

    for part_text, is_latex in parts:
        if is_latex:
            # Try to add as native OMML equation
            success = add_math_to_paragraph(paragraph, part_text)
            if not success:
                # Fallback to Unicode representation
                unicode_text = latex_to_unicode(part_text)
                run = paragraph.add_run(unicode_text)
                run.italic = True
        else:
            paragraph.add_run(part_text)


@router.post("/export")
def export(
    samples: str = Form(...),
    topic: str = Form(...),
    custom_keywords: str = Form(""),
    paraphrase: bool = Form(False),
    change_numbers: bool = Form(False),
    change_context: bool = Form(False),
    variants_per_question: int = Form(1),
    fmt: str = Form("txt"),
    use_ai: bool = Form(False),
    ai_engine: str = Form("openai"),
):
    """Export generated questions to various formats."""
    from app.services.math_renderer import render_text_with_math, has_latex

    # Export directly without regenerating - just use the samples as-is
    questions = [s.strip() for s in samples.split("\n\n") if s.strip()]
    if not questions:
        # Fallback: split by single newline if no double newlines
        questions = [s.strip() for s in samples.split("\n") if s.strip()]

    if fmt == "csv":
        # Convert LaTeX to Unicode for CSV
        csv_questions = [render_text_with_math(q) for q in questions]
        content = "id,question\n" + "\n".join(
            f"{i+1},\"{q.replace('\"', '\"\"')}\"" for i, q in enumerate(csv_questions)
        )
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=questions.csv"},
        )

    if fmt == "docx":
        doc = Document()
        doc.add_heading("Danh sách câu hỏi", level=1)
        for i, q in enumerate(questions, 1):
            # Split question into lines (first line = question, rest = options)
            lines = q.split('\n')
            p = doc.add_paragraph()
            p.add_run(f"Câu {i}. ").bold = True

            # Render first line with math
            _add_text_with_math(p, lines[0])

            # Add options on separate lines
            for line in lines[1:]:
                line = line.strip()
                if line:
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.3)
                    _add_text_with_math(opt_p, line)
            doc.add_paragraph("")  # Empty line between questions
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=questions.docx"},
        )

    if fmt == "pdf":
        font_name, _ = _get_pdf_font()
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 60
        font_size = 12
        c.setFont(font_name, font_size)
        c.drawString(50, y, "Danh sách câu hỏi")
        y -= 30
        for i, q in enumerate(questions, 1):
            # Split question into lines (first line = question, rest = options)
            lines = q.split('\n')

            # Check page break
            if y < 80:
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - 60

            # Write question number
            c.drawString(50, y, f"Câu {i}. ")
            x_offset = 50 + pdfmetrics.stringWidth(f"Câu {i}. ", font_name, font_size)

            # Write question content with math
            if has_latex(lines[0]):
                _draw_text_with_math_pdf(c, lines[0], x_offset, y, font_name, font_size, width - x_offset - 50)
            else:
                # No LaTeX - use simple text wrapping
                q_text = render_text_with_math(lines[0])
                remaining_width = width - x_offset - 50
                if pdfmetrics.stringWidth(q_text, font_name, font_size) <= remaining_width:
                    c.drawString(x_offset, y, q_text)
                else:
                    # Need to wrap
                    wrapped = _wrap_text(q_text, remaining_width, font_name, font_size)
                    c.drawString(x_offset, y, wrapped[0])
                    for part in wrapped[1:]:
                        y -= 18
                        if y < 60:
                            c.showPage()
                            c.setFont(font_name, font_size)
                            y = height - 60
                        c.drawString(50, y, part)
            y -= 20

            # Write options
            for line in lines[1:]:
                line = line.strip()
                if line:
                    if y < 60:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = height - 60

                    if has_latex(line):
                        c.drawString(70, y, "")  # Position marker
                        _draw_text_with_math_pdf(c, line, 70, y, font_name, font_size, width - 120)
                    else:
                        opt_text = render_text_with_math(line)
                        c.drawString(70, y, opt_text)
                    y -= 16
            y -= 10  # Space between questions
        c.save()
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=questions.pdf"},
        )

    # TXT format - convert LaTeX to Unicode
    txt_questions = [render_text_with_math(q) for q in questions]
    content = "\n".join(txt_questions)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=questions.txt"},
    )


@router.post("/api/export-exam")
async def export_exam(request: Request):
    """Export exam questions with full content and options."""
    from app.services.math_renderer import render_text_with_math, has_latex

    data = await request.json()
    questions = data.get("questions", [])
    fmt = data.get("format", "docx")
    title = data.get("title", "Đề thi")

    if not questions:
        raise HTTPException(status_code=400, detail="Không có câu hỏi để xuất")

    if fmt == "docx":
        doc = Document()
        doc.add_heading(title, level=1)

        for i, q in enumerate(questions, 1):
            content = q.get("content", "") or q.get("question", "")
            options = q.get("options", [])

            p = doc.add_paragraph()
            p.add_run(f"Câu {i}. ").bold = True
            _add_text_with_math(p, content)

            if options:
                labels = ['A', 'B', 'C', 'D', 'E']
                for j, opt in enumerate(options[:5]):
                    opt_clean = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', str(opt).strip())
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.3)
                    _add_text_with_math(opt_p, f"{labels[j]}) {opt_clean}")

            doc.add_paragraph("")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=de_thi.docx"},
        )

    if fmt == "pdf":
        font_name, _ = _get_pdf_font()
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 60
        font_size = 12

        c.setFont(font_name, 16)
        c.drawString(50, y, title)
        y -= 40
        c.setFont(font_name, font_size)

        for i, q in enumerate(questions, 1):
            content = q.get("content", "") or q.get("question", "")
            options = q.get("options", [])

            # Check page break
            if y < 80:
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - 60

            # Write question number
            q_prefix = f"Câu {i}. "
            c.drawString(50, y, q_prefix)
            x_offset = 50 + pdfmetrics.stringWidth(q_prefix, font_name, font_size)

            # Write question content with math rendering
            if has_latex(content):
                _draw_text_with_math_pdf(c, content, x_offset, y, font_name, font_size, width - x_offset - 50)
            else:
                content_text = render_text_with_math(content)
                remaining_width = width - x_offset - 50
                if pdfmetrics.stringWidth(content_text, font_name, font_size) <= remaining_width:
                    c.drawString(x_offset, y, content_text)
                else:
                    wrapped = _wrap_text(content_text, remaining_width, font_name, font_size)
                    c.drawString(x_offset, y, wrapped[0])
                    for part in wrapped[1:]:
                        y -= 18
                        if y < 60:
                            c.showPage()
                            c.setFont(font_name, font_size)
                            y = height - 60
                        c.drawString(50, y, part)
            y -= 20

            # Write options with math rendering
            if options:
                labels = ['A', 'B', 'C', 'D', 'E']
                for j, opt in enumerate(options[:5]):
                    opt_clean = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', str(opt).strip())
                    opt_line = f"{labels[j]}) {opt_clean}"

                    if y < 60:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = height - 60

                    if has_latex(opt_line):
                        _draw_text_with_math_pdf(c, opt_line, 70, y, font_name, font_size, width - 120)
                    else:
                        opt_text = render_text_with_math(opt_line)
                        c.drawString(70, y, opt_text)
                    y -= 16

            y -= 10

        c.save()
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=de_thi.pdf"},
        )

    raise HTTPException(status_code=400, detail="Format không hỗ trợ")
