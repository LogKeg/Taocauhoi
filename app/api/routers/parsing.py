"""
Word document parsing API endpoints.
"""
import io
import json
import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from app.core import ANSWER_TEMPLATES

# Import parsers directly
from app.parsers.docx import (
    parse_cell_based_questions,
    extract_docx_lines_with_options,
    extract_docx_content,
    parse_bilingual_questions,
)

router = APIRouter(tags=["parsing"])


def _clean_option_label(opt: str) -> str:
    """Remove leading option labels like 'A)', 'A.', 'Option A:' from option text."""
    return re.sub(r'^(?:Option\s+)?[A-E][).:]\s*', '', opt.strip())


def _extract_questions_from_json(text: str) -> list:
    """Extract question objects from AI response that may contain multiple JSON arrays."""
    # 1. Try parsing as single JSON array (greedy match)
    json_match = re.search(r'\[[\s\S]*\]', text)
    if json_match:
        try:
            result = json.loads(json_match.group())
            if isinstance(result, list):
                # Clean option labels
                for q in result:
                    if 'options' in q:
                        q['options'] = [_clean_option_label(o) for o in q['options']]
                return result
        except json.JSONDecodeError:
            pass

    # 2. Find individual JSON objects with 'question' key
    questions = []
    for m in re.finditer(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', text):
        try:
            obj = json.loads(m.group())
            if 'question' in obj:
                if 'options' in obj:
                    obj['options'] = [_clean_option_label(o) for o in obj['options']]
                questions.append(obj)
        except (json.JSONDecodeError, ValueError):
            pass
    return questions


def _is_bilingual_question(text: str) -> bool:
    """Check if a question text is bilingual (contains both English and Vietnamese)."""
    if not text:
        return False
    # Check for common bilingual separators
    if ' / ' in text or ' - ' in text:
        parts = re.split(r'\s*/\s*|\s*-\s*', text, maxsplit=1)
        if len(parts) >= 2:
            # Check if one part has Vietnamese characters and other has English
            has_viet = any('\u00c0' <= c <= '\u1ef9' for c in text)
            has_english = bool(re.search(r'[a-zA-Z]{3,}', text))
            return has_viet and has_english
    return False


def _has_vietnamese(text: str) -> bool:
    """Check if text contains Vietnamese characters."""
    return any('\u00c0' <= c <= '\u1ef9' for c in text)


def _has_english(text: str) -> bool:
    """Check if text contains English words."""
    return bool(re.search(r'[a-zA-Z]{3,}', text))


def _detect_exam_info(questions: list) -> tuple:
    """Detect subject and language from parsed questions.

    Returns:
        tuple: (detected_subject, detected_language)
    """
    if not questions:
        return ("general", "unknown")

    # Combine text from first 10 questions for analysis
    all_text = ""
    for q in questions[:10]:
        all_text += q.get('question', '') + " " + " ".join(q.get('options', []))
    all_text_lower = all_text.lower()

    # Detect subject
    detected_subject = "general"
    subject_keywords = {
        "science": ['science', 'khoa học', 'biology', 'sinh học', 'chemistry', 'hóa học',
                    'physics', 'vật lý', 'organism', 'cell', 'atom', 'molecule', 'energy'],
        "math": ['math', 'toán', 'calculate', 'tính', 'equation', 'phương trình',
                 'number', 'số', 'algebra', 'geometry', 'hình học'],
        "history": ['history', 'lịch sử', 'war', 'chiến tranh', 'dynasty', 'triều đại',
                    'king', 'vua', 'emperor', 'revolution', 'cách mạng'],
        "geography": ['geography', 'địa lý', 'country', 'quốc gia', 'continent', 'châu lục',
                      'river', 'sông', 'mountain', 'núi', 'capital', 'thủ đô'],
        "english": ['grammar', 'ngữ pháp', 'vocabulary', 'từ vựng', 'tense', 'thì',
                    'adjective', 'tính từ', 'verb', 'động từ', 'noun', 'danh từ'],
    }

    max_score = 0
    for subject, keywords in subject_keywords.items():
        score = sum(1 for kw in keywords if kw in all_text_lower)
        if score > max_score:
            max_score = score
            detected_subject = subject

    # Detect language
    has_viet = _has_vietnamese(all_text)
    has_eng = _has_english(all_text)
    has_bilingual_separator = ' / ' in all_text

    if has_bilingual_separator and has_viet and has_eng:
        detected_language = "bilingual"
    elif has_viet and has_eng:
        detected_language = "mixed"
    elif has_viet:
        detected_language = "vietnamese"
    elif has_eng:
        detected_language = "english"
    else:
        detected_language = "unknown"

    return (detected_subject, detected_language)


def _translate_to_bilingual(questions: list, call_ai_func, ai_engine: str) -> list:
    """Translate English questions to bilingual format (English / Vietnamese)."""
    result = []

    # Process in small batches of 3 for better translation accuracy
    TRANS_BATCH = 3
    for i in range(0, len(questions), TRANS_BATCH):
        batch = questions[i:i + TRANS_BATCH]

        # Build numbered translation prompt for better alignment
        items_to_translate = []
        for q in batch:
            items_to_translate.append(q.get('question', ''))
            for opt in q.get('options', []):
                items_to_translate.append(opt)

        prompt = f"Dịch {len(items_to_translate)} dòng sau sang tiếng Việt. CHỈ trả về bản dịch, giữ nguyên đánh số, KHÔNG thêm lời giới thiệu.\n\n"
        for idx, item in enumerate(items_to_translate, 1):
            prompt += f"{idx}. {item}\n"

        try:
            response, error = call_ai_func(prompt, ai_engine)
            if not error and response:
                # Remove thinking tags if present
                response = re.sub(r'<think>[\s\S]*?</think>', '', response).strip()
                # Build dict keyed by line number for reliable alignment
                trans_dict = {}
                for line in response.strip().split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    # Only accept lines starting with a number
                    m = re.match(r'^(\d+)[.):\s]+(.+)', line)
                    if m:
                        num = int(m.group(1))
                        trans_dict[num] = m.group(2).strip()

                # Convert to ordered list
                translations = []
                for idx in range(1, len(items_to_translate) + 1):
                    translations.append(trans_dict.get(idx, ''))

                idx = 0
                for q in batch:
                    eng_question = q.get('question', '')
                    eng_options = q.get('options', [])
                    answer = q.get('answer', 'A')

                    # Get Vietnamese translation for question
                    viet_question = translations[idx] if idx < len(translations) else eng_question
                    idx += 1
                    # Clean any leftover label prefix
                    viet_question = re.sub(r'^(?:Option\s+)?[A-E][).:]\s*', '', viet_question)

                    # Get Vietnamese translations for options
                    bilingual_options = []
                    for opt in eng_options:
                        viet_opt = translations[idx] if idx < len(translations) else opt
                        idx += 1
                        # Clean up any A), B) prefix from translation
                        viet_opt = re.sub(r'^(?:Option\s+)?[A-E][).:]\s*', '', viet_opt)
                        bilingual_options.append(f"{opt} / {viet_opt}")

                    result.append({
                        'question': f"{eng_question}\n{viet_question}",
                        'options': bilingual_options,
                        'answer': answer,
                    })
                continue
        except Exception:
            pass

        # Fallback: keep original if translation fails
        result.extend(batch)

    return result


def _ensure_bilingual_format(questions: list) -> list:
    """Post-process to ensure questions are in bilingual format."""
    result = []
    for q in questions:
        question_text = q.get('question', '')
        options = q.get('options', [])
        answer = q.get('answer', 'A')

        # Check if question already has both languages
        has_slash = ' / ' in question_text
        has_viet = _has_vietnamese(question_text)
        has_eng = _has_english(question_text)

        # If question is only in one language, try to detect and note it
        if not has_slash or not (has_viet and has_eng):
            # If only English, add Vietnamese placeholder
            if has_eng and not has_viet:
                question_text = question_text + " / [Cần dịch sang tiếng Việt]"
            # If only Vietnamese, add English placeholder
            elif has_viet and not has_eng:
                question_text = "[Need English translation] / " + question_text

        # Process options similarly
        processed_options = []
        for opt in options:
            opt_has_slash = ' / ' in opt
            opt_has_viet = _has_vietnamese(opt)
            opt_has_eng = _has_english(opt)

            if not opt_has_slash or not (opt_has_viet and opt_has_eng):
                if opt_has_eng and not opt_has_viet:
                    opt = opt + " / [Cần dịch]"
                elif opt_has_viet and not opt_has_eng:
                    opt = "[Translation] / " + opt
            processed_options.append(opt)

        result.append({
            'question': question_text,
            'options': processed_options,
            'answer': answer
        })

    return result


def _save_questions_to_bank(
    questions: List[dict],
    subject: str = "general",
    source: str = "imported",
    question_type: str = "mcq",
) -> int:
    """Save parsed questions to the question bank."""
    from app.database import SessionLocal, QuestionCRUD

    saved = 0
    db = SessionLocal()
    try:
        for q in questions:
            content = q.get('question', '')
            if not content.strip():
                continue

            # Convert options list to JSON string
            options = q.get('options', [])
            options_json = json.dumps(options) if options else None
            answer = q.get('answer', q.get('correct_answer', ''))

            QuestionCRUD.create(
                db,
                content=content.strip(),
                subject=subject,
                source=source,
                difficulty="medium",
                question_type=question_type,
                options=options_json,
                answer=answer,
            )
            saved += 1
    finally:
        db.close()
    return saved


def _get_parsing_functions():
    """Lazy import of parsing functions from main module."""
    from app import main
    from app.services.ai import call_ai, load_saved_settings
    return {
        'parse_cell_based_questions': parse_cell_based_questions,
        'extract_docx_lines': extract_docx_lines_with_options,
        'extract_docx_content': extract_docx_content,
        'parse_math_exam_questions': main._parse_math_exam_questions,
        'parse_english_exam_questions': main._parse_english_exam_questions,
        'parse_envie_questions': main._parse_envie_questions,
        'parse_bilingual_questions': parse_bilingual_questions,
        'save_questions_to_bank': _save_questions_to_bank,
        'call_ai': call_ai,
        'load_ai_settings': load_saved_settings,
    }


@router.post("/api/parse-exam")
async def parse_exam_file(file: UploadFile):
    """Parse questions from uploaded Word document."""
    funcs = _get_parsing_functions()

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    content = await file.read()
    doc = Document(io.BytesIO(content))

    # Try cell-based parser first (for ASMO Science format)
    questions = funcs['parse_cell_based_questions'](doc)
    if questions:
        # Detect subject and language
        detected_subject, detected_language = _detect_exam_info(questions)
        return {
            "ok": True,
            "filename": file.filename,
            "total_lines": len(questions),
            "questions": questions,
            "count": len(questions),
            "detected_subject": detected_subject,
            "detected_language": detected_language,
        }

    # Extract lines from document
    lines, table_options = funcs['extract_docx_lines'](doc)

    # Detect format and use appropriate parser
    is_math_format = any(re.match(r'^Question\s+\d+', line, re.IGNORECASE) for line in lines[:20])
    is_english_level_format = (
        any('Section A' in line or 'Section B' in line for line in lines[:15]) and
        file.filename and 'LEVEL' in file.filename.upper()
    )
    is_envie_format = file.filename and 'EN-VIE' in file.filename.upper()

    # Detect bilingual science (IKSC) format
    is_bilingual_science_format = False
    if not is_math_format and not is_english_level_format and not is_envie_format:
        bilingual_option_count = sum(
            1 for line in lines[:80]
            if re.match(r'^[A-E][.)]\s*.+\s*/\s*.+', line)
        )
        bilingual_pair_count = 0
        for li in range(len(lines) - 1):
            l1, l2 = lines[li], lines[li + 1]
            if (re.match(r'^\d+\.\s*', l1)
                    and not any(ord(c) > 127 for c in l1[:200])
                    and any(ord(c) > 127 for c in l2[:200])
                    and not re.match(r'^[A-E][.)]\s*', l2)):
                bilingual_pair_count += 1
                if bilingual_pair_count >= 3:
                    break
        if bilingual_option_count >= 3 or bilingual_pair_count >= 3:
            is_bilingual_science_format = True
            is_envie_format = True  # Use envie parser

    # Detect English reading/competition exam format
    is_english_reading_format = False
    if not is_math_format and not is_english_level_format and not is_envie_format:
        text_sample = ' '.join(lines[:50]).lower()
        has_instruction = any(
            kw in text_sample
            for kw in ['for each question', 'read the text', 'choose the correct answer', 'choose the best answer']
        )
        has_merged_options = any(
            re.search(r'[A-D]\)\s*\w', line) for line in lines[:80]
        )
        if has_instruction or has_merged_options:
            is_english_reading_format = True

    if is_math_format:
        questions = funcs['parse_math_exam_questions'](lines)
    elif is_english_level_format:
        questions = funcs['parse_english_exam_questions'](doc)
    elif is_envie_format or is_english_reading_format:
        questions = funcs['parse_envie_questions'](doc)
    else:
        questions = funcs['parse_bilingual_questions'](lines, table_options)

    # Detect subject and language
    detected_subject, detected_language = _detect_exam_info(questions)

    # For science subject, apply bilingual dedup (IKSC format: EN+VN pairs)
    if detected_subject == 'science' or is_bilingual_science_format:
        from app.main import _dedup_bilingual_science
        questions = _dedup_bilingual_science(questions)

    return {
        "ok": True,
        "filename": file.filename,
        "total_lines": len(lines),
        "questions": questions,
        "count": len(questions),
        "detected_subject": detected_subject,
        "detected_language": detected_language,
    }


def _parse_answer_key(answer_key_content: bytes, answer_key_filename: str, question_filename: str) -> Dict[int, int]:
    """Parse answer key file and return answers for the matching level.

    Returns dict mapping question_number -> answer_number (A=1, B=2, C=3, D=4, E=5).
    """
    answer_map = {'A': 1, 'B': 2, 'C': 3, 'D': 4, 'E': 5}
    answers = {}

    try:
        if answer_key_filename.lower().endswith('.docx'):
            doc = Document(io.BytesIO(answer_key_content))
            for table in doc.tables:
                if len(table.rows) < 2 or len(table.columns) < 2:
                    continue

                # Parse header row to find level columns
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]

                # Find matching level based on question filename
                # e.g., "Benjamin.docx" -> match "Benjamin" in header
                q_name = question_filename.rsplit('.', 1)[0].lower().strip()
                level_col = -1
                for col_idx, header in enumerate(header_cells):
                    if header.strip().lower() == q_name:
                        level_col = col_idx
                        break

                if level_col < 0:
                    # Try partial match (e.g., "PreEcolier" in "Pre-Ecolier")
                    for col_idx, header in enumerate(header_cells):
                        h = header.strip().lower().replace('-', '').replace(' ', '')
                        qn = q_name.replace('-', '').replace(' ', '')
                        if h == qn:
                            level_col = col_idx
                            break

                if level_col < 0:
                    continue

                # Answer is in the column after the question number column
                # Format: [Q#, Answer, Q#, Answer, ...]
                # Ensure level_col is the question number column (even index in pair)
                if level_col % 2 == 1:
                    level_col -= 1  # Go back to question number column

                answer_col = level_col + 1

                # Parse answer rows
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if answer_col >= len(cells):
                        continue
                    q_num_text = cells[level_col].strip()
                    answer_text = cells[answer_col].strip().upper()
                    if q_num_text and q_num_text.isdigit() and answer_text in answer_map:
                        answers[int(q_num_text)] = answer_map[answer_text]

        elif answer_key_filename.lower().endswith(('.xlsx', '.xls')):
            # Excel support: columns are levels, rows are Q# + answer
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(answer_key_content), read_only=True)
            ws = wb.active

            rows_data = list(ws.iter_rows(values_only=True))
            if rows_data:
                # Row 0 is header: [Q#_or_empty, Level1, Level2, ...]
                headers = [str(h).strip().lower() if h else '' for h in rows_data[0]]

                # Match level based on question filename
                q_name = question_filename.rsplit('.', 1)[0].lower().strip()
                # Remove common suffixes like " Edited 1-2026"
                q_name_clean = re.sub(r'\s*edited.*$', '', q_name, flags=re.IGNORECASE).strip()

                level_col = -1
                for col_idx, header in enumerate(headers):
                    h = header.replace('-', '').replace(' ', '')
                    qn = q_name_clean.replace('-', '').replace(' ', '')
                    if h == qn or h == q_name.replace('-', '').replace(' ', ''):
                        level_col = col_idx
                        break

                if level_col >= 0:
                    for row in rows_data[1:]:
                        if len(row) <= level_col:
                            continue
                        q_num_val = row[0]
                        answer_val = row[level_col]
                        if q_num_val is not None and answer_val is not None:
                            q_num_str = str(q_num_val).strip()
                            answer_text = str(answer_val).strip().upper()
                            if q_num_str.isdigit() and answer_text in answer_map:
                                answers[int(q_num_str)] = answer_map[answer_text]
            wb.close()

        elif answer_key_filename.lower().endswith('.pdf'):
            # PDF support: extract text and find answer patterns
            try:
                import fitz  # PyMuPDF
                pdf_doc = fitz.open(stream=answer_key_content, filetype="pdf")
                full_text = ""
                for page in pdf_doc:
                    full_text += page.get_text()
                pdf_doc.close()

                # Find answer patterns like "1. A", "2. B", "1 A", "1\tA"
                for match in re.finditer(r'(\d+)\s*[.\t)]\s*([A-E])\b', full_text):
                    q_num = int(match.group(1))
                    answer_letter = match.group(2).upper()
                    if answer_letter in answer_map:
                        answers[q_num] = answer_map[answer_letter]
            except ImportError:
                pass  # PyMuPDF not installed, skip PDF parsing

    except Exception:
        pass  # Return empty answers on any error

    return answers


@router.post("/convert-word-to-excel")
def convert_word_to_excel(
    file: UploadFile = Form(...),
    use_latex: str = Form("0"),
    subject: str = Form("general"),
    answer_key: Optional[UploadFile] = File(None),
) -> StreamingResponse:
    """Convert a Word file containing questions to Excel format."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from docx.oxml.ns import qn as docx_qn

    funcs = _get_parsing_functions()

    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    latex_enabled = use_latex == "1"

    try:
        content = file.file.read()
        doc = Document(io.BytesIO(content))

        # Try cell-based parser first
        questions = funcs['parse_cell_based_questions'](doc)

        is_math_format = False
        is_english_level_format = False
        is_envie_format = False

        if not questions:
            lines, table_options = funcs['extract_docx_lines'](doc, include_textboxes=True, use_latex=latex_enabled)

            is_math_format = any(re.match(r'^Question\s+\d+', line, re.IGNORECASE) for line in lines[:20])
            is_english_level_format = (
                any('Section A' in line or 'Section B' in line for line in lines[:15]) and
                file.filename and 'LEVEL' in file.filename.upper()
            )
            is_envie_format = file.filename and 'EN-VIE' in file.filename.upper()

            # Detect bilingual science (IKSC) format
            if not is_envie_format:
                # Check bilingual options (e.g., "A. Refraction / Khúc xạ")
                bilingual_option_count = sum(
                    1 for line in lines[:80]
                    if re.match(r'^[A-E][.)]\s*.+\s*/\s*.+', line)
                )
                # Also check for EN+VN question pairs (numbered EN → VN translation)
                bilingual_pair_count = 0
                for li in range(len(lines) - 1):
                    l1, l2 = lines[li], lines[li + 1]
                    if (re.match(r'^\d+\.\s*', l1)
                            and not any(ord(c) > 127 for c in l1[:200])
                            and any(ord(c) > 127 for c in l2[:200])
                            and not re.match(r'^[A-E][.)]\s*', l2)):
                        bilingual_pair_count += 1
                        if bilingual_pair_count >= 3:
                            break
                if bilingual_option_count >= 3 or bilingual_pair_count >= 3:
                    is_envie_format = True

            # Check for Word numbered list format
            is_word_numbered_format = False
            if not is_envie_format:
                body = doc._element.body
                num_level_0 = 0
                num_level_1 = 0
                for child in body:
                    tag = child.tag.split('}')[-1]
                    if tag == 'p':
                        pPr = child.find(docx_qn('w:pPr'))
                        if pPr is not None:
                            numPr = pPr.find(docx_qn('w:numPr'))
                            if numPr is not None:
                                ilvl_elem = numPr.find(docx_qn('w:ilvl'))
                                if ilvl_elem is not None:
                                    ilvl = int(ilvl_elem.get(docx_qn('w:val')) or '0')
                                    if ilvl == 0:
                                        num_level_0 += 1
                                    elif ilvl == 1:
                                        num_level_1 += 1
                if num_level_0 > 10 and num_level_0 == num_level_1:
                    is_word_numbered_format = True

            # Detect English reading/competition exam format
            is_english_reading_format = False
            if not is_math_format and not is_english_level_format and not is_envie_format and not is_word_numbered_format:
                text_sample = ' '.join(lines[:50]).lower()
                has_instruction = any(
                    kw in text_sample
                    for kw in ['for each question', 'read the text', 'choose the correct answer', 'choose the best answer']
                )
                has_merged_options = any(
                    re.search(r'[A-D]\)\s*\w', line) for line in lines[:80]
                )
                if has_instruction or has_merged_options:
                    is_english_reading_format = True

            if is_math_format:
                questions = funcs['parse_math_exam_questions'](lines)
            elif is_english_level_format:
                questions = funcs['parse_english_exam_questions'](doc)
            elif is_envie_format or is_word_numbered_format or is_english_reading_format:
                questions = funcs['parse_envie_questions'](doc)
            else:
                questions = funcs['parse_bilingual_questions'](lines, table_options)

        # For science subject, apply bilingual dedup (IKSC format: EN+VN pairs)
        if subject == 'science' and questions:
            from app.main import _dedup_bilingual_science
            questions = _dedup_bilingual_science(questions)

        # Save to question bank
        detected_subject = subject if subject in ('english', 'math', 'science', 'general') else 'general'
        funcs['save_questions_to_bank'](
            questions,
            subject=detected_subject,
            source=f"word-import:{file.filename}",
            question_type='mcq' if questions and questions[0].get('options') else 'blank',
        )

        # Sort questions by number if most questions have a number assigned
        numbered_count = sum(1 for q in questions if q.get('number') is not None)
        if numbered_count > len(questions) * 0.5:
            questions.sort(key=lambda q: q.get('number') or float('inf'))

        # Parse answer key if provided
        answer_key_data = {}
        if answer_key is not None and answer_key.filename:
            ak_content = answer_key.file.read()
            if ak_content:
                answer_key_data = _parse_answer_key(ak_content, answer_key.filename, file.filename)

        # Create Excel workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Questions"

        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        highlight_fill = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Standard format: Question Type, Question, Option 1-5, Correct Answer,
        # Default Marks, Default Time To Solve, Difficulty Level, Hint, Solution
        headers = [
            "Question Type", "Question",
            "Option 1", "Option 2", "Option 3", "Option 4", "Option 5",
            "Correct Answer", "Default Marks", "Default Time To Solve",
            "Difficulty Level", "Hint", "Solution",
        ]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border

        # Determine level from filename for IKLC/IKSC scoring
        raw_name = file.filename.rsplit('.', 1)[0].strip().lower() if file.filename else ''
        total_questions = len(questions)

        # Extract level name from filename (e.g., "Benjamin Edited 1-2026" → "benjamin")
        level_name = raw_name
        for lvl in ['preecolier', 'pre-ecolier', 'ecolier', 'benjamin', 'cadet', 'junior', 'student']:
            if lvl in raw_name.replace('-', '').replace(' ', ''):
                level_name = lvl
                break

        # IKLC English scoring rules per level
        iklc_marks = {
            'preecolier': {'correct': 2.00, 'wrong': 0},
            'pre-ecolier': {'correct': 2.00, 'wrong': 0},
            'ecolier': {'correct': 2.00, 'wrong': 0},
            'benjamin': {'correct': 1.00, 'wrong': -0.50},
            'cadet': {'correct': 1.25, 'wrong': -0.50},
            'junior': {'correct': 1.50, 'wrong': -0.50},
            'student': {'correct': 1.75, 'wrong': -0.50},
        }

        # IKSC Science scoring: difficulty-based (easy=3, medium=4, hard=5), wrong=-1
        # Questions split into 3 equal groups by position

        def get_default_marks(q_idx_0based):
            """Return (default_mark, difficulty_level) for a question."""
            if subject == 'english' and level_name in iklc_marks:
                marks_info = iklc_marks[level_name]
                return marks_info['correct'], "EASY"
            elif subject == 'science':
                # IKSC: Split into 3 equal groups (easy=3, medium=4, hard=5)
                group_size = total_questions // 3 if total_questions >= 3 else total_questions
                if q_idx_0based < group_size:
                    return 3, "EASY"
                elif q_idx_0based < group_size * 2:
                    return 4, "MEDIUM"
                else:
                    return 5, "HARD"
            return None, "EASY"

        for row_idx, q in enumerate(questions, 2):
            question_text = q.get('question', '')
            options = q.get('options', [])
            answer = q.get('answer', '')
            q_idx_0based = row_idx - 2  # 0-based question index

            # Determine question type: MSA = trắc nghiệm, SAQ = tự luận
            q_type = "MSA" if options else "SAQ"

            # A: Question Type
            cell = ws.cell(row=row_idx, column=1, value=q_type)
            cell.alignment = center_align
            cell.border = thin_border

            # B: Question
            cell = ws.cell(row=row_idx, column=2, value=question_text)
            cell.alignment = left_align
            cell.border = thin_border

            # Determine correct answer from answer key or parsed answer
            # answer_key_data maps question_number -> answer_number (1-5)
            q_num = row_idx - 1  # 1-based question number
            correct_answer_num = answer_key_data.get(q_num)

            # Fallback to parsed answer if no answer key
            if correct_answer_num is None and answer and isinstance(answer, str):
                answer_upper = answer.upper().strip()
                if answer_upper in ('A', 'B', 'C', 'D', 'E'):
                    correct_answer_num = ord(answer_upper) - ord('A') + 1

            # C-G: Option 1-5
            for opt_idx in range(5):
                cell = ws.cell(row=row_idx, column=3 + opt_idx,
                               value=options[opt_idx] if opt_idx < len(options) else None)
                cell.alignment = left_align
                cell.border = thin_border

                # Highlight correct answer option
                if opt_idx < len(options) and correct_answer_num and (opt_idx + 1) == correct_answer_num:
                    cell.fill = highlight_fill

            # H: Correct Answer (A=1, B=2, C=3, D=4, E=5)
            cell = ws.cell(row=row_idx, column=8, value=correct_answer_num)
            cell.alignment = center_align
            cell.border = thin_border

            # I: Default Marks & K: Difficulty Level (based on subject/level)
            default_mark, difficulty = get_default_marks(q_idx_0based)
            cell = ws.cell(row=row_idx, column=9, value=default_mark)
            cell.alignment = center_align
            cell.border = thin_border

            # J: Default Time To Solve
            cell = ws.cell(row=row_idx, column=10, value=30)
            cell.alignment = center_align
            cell.border = thin_border

            # K: Difficulty Level
            cell = ws.cell(row=row_idx, column=11, value=difficulty)
            cell.alignment = center_align
            cell.border = thin_border

            # L: Hint
            cell = ws.cell(row=row_idx, column=12, value=None)
            cell.border = thin_border

            # M: Solution
            cell = ws.cell(row=row_idx, column=13, value=None)
            cell.border = thin_border

        # Adjust column widths
        ws.column_dimensions["A"].width = 15   # Question Type
        ws.column_dimensions["B"].width = 60   # Question
        ws.column_dimensions["C"].width = 25   # Option 1
        ws.column_dimensions["D"].width = 25   # Option 2
        ws.column_dimensions["E"].width = 25   # Option 3
        ws.column_dimensions["F"].width = 25   # Option 4
        ws.column_dimensions["G"].width = 25   # Option 5
        ws.column_dimensions["H"].width = 15   # Correct Answer
        ws.column_dimensions["I"].width = 14   # Default Marks
        ws.column_dimensions["J"].width = 20   # Default Time To Solve
        ws.column_dimensions["K"].width = 16   # Difficulty Level
        ws.column_dimensions["L"].width = 20   # Hint
        ws.column_dimensions["M"].width = 20   # Solution

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        base_name = Path(file.filename).stem
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base_name}_questions.xlsx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý file: {str(e)}")


@router.post("/api/analyze-exam")
async def analyze_exam_with_ai(
    file: UploadFile,
    ai_engine: str = Form("openai"),
):
    """Analyze an exam file using AI to extract and categorize questions."""
    funcs = _get_parsing_functions()

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    content = await file.read()
    doc = Document(io.BytesIO(content))

    text_content = funcs['extract_docx_content'](doc)
    settings = funcs['load_ai_settings']()

    prompt = f"""Phân tích đề thi sau và trích xuất các câu hỏi. Trả về JSON array với format:
[
    {{
        "number": 1,
        "content": "nội dung câu hỏi",
        "options": ["A", "B", "C", "D"],
        "correct_answer": "A",
        "topic": "chủ đề",
        "difficulty": "easy|medium|hard"
    }}
]

Nội dung đề thi:
{text_content[:8000]}

Chỉ trả về JSON array, không giải thích."""

    try:
        response, error = funcs['call_ai'](prompt, ai_engine)
        if error:
            return {"ok": False, "error": error}
        json_match = re.search(r'\[[\s\S]*\]', response)
        if json_match:
            questions = json.loads(json_match.group())
            return {
                "ok": True,
                "filename": file.filename,
                "questions": questions,
                "count": len(questions),
            }
        return {"ok": False, "error": "Không thể phân tích kết quả AI"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@router.post("/api/generate-similar-exam")
async def generate_similar_exam(
    file: UploadFile,
    difficulty: str = Form("same"),
    subject: str = Form("auto"),
    bilingual: str = Form("auto"),
    ai_engine: str = Form("openai"),
):
    """Generate similar questions based on an exam file."""
    funcs = _get_parsing_functions()

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    content = await file.read()
    doc = Document(io.BytesIO(content))

    # Try cell-based parser first
    sample_questions = funcs['parse_cell_based_questions'](doc)

    if not sample_questions:
        lines, table_options = funcs['extract_docx_lines'](doc)

        is_math_format = any(re.match(r'^Question\s+\d+', line, re.IGNORECASE) for line in lines[:20])
        is_english_level_format = (
            any('Section A' in line or 'Section B' in line for line in lines[:15]) and
            file.filename and 'LEVEL' in file.filename.upper()
        )
        is_envie_format = file.filename and 'EN-VIE' in file.filename.upper()

        if is_math_format:
            sample_questions = funcs['parse_math_exam_questions'](lines)
        elif is_english_level_format:
            sample_questions = funcs['parse_english_exam_questions'](doc)
        elif is_envie_format:
            sample_questions = funcs['parse_envie_questions'](doc)
        else:
            sample_questions = funcs['parse_bilingual_questions'](lines, table_options)

    if not sample_questions:
        return {"ok": False, "error": "Không tìm thấy câu hỏi trong file"}

    # Build difficulty instruction
    difficulty_text = {
        "same": "tương đương về độ khó",
        "easier": "DỄ HƠN (đơn giản hơn)",
        "harder": "KHÓ HƠN (phức tạp hơn)"
    }.get(difficulty, "tương đương về độ khó")

    settings = funcs['load_ai_settings']()

    subject_names = {
        "science": "Science/Khoa học",
        "math": "Math/Toán học",
        "history": "History/Lịch sử",
        "geography": "Geography/Địa lý",
        "english": "English/Tiếng Anh",
        "general": "General",
        "auto": "General"
    }

    # Auto-detect subject if needed
    if subject == "auto":
        all_sample_text = ""
        for q in sample_questions[:10]:
            all_sample_text += q.get('question', '') + " " + " ".join(q.get('options', []))
        all_sample_lower = all_sample_text.lower()

        detected_subject = "general"
        if any(kw in all_sample_lower for kw in ['science', 'khoa học', 'biology', 'chemistry', 'physics']):
            detected_subject = "science"
        elif any(kw in all_sample_lower for kw in ['math', 'toán', 'calculate', 'equation', 'number']):
            detected_subject = "math"
        elif any(kw in all_sample_lower for kw in ['history', 'lịch sử', 'war', 'dynasty']):
            detected_subject = "history"
        elif any(kw in all_sample_lower for kw in ['geography', 'địa lý', 'country', 'continent']):
            detected_subject = "geography"
        elif any(kw in all_sample_lower for kw in ['english', 'tiếng anh', 'vocabulary', 'grammar']):
            detected_subject = "english"
    else:
        detected_subject = subject

    # Detect if questions are bilingual
    is_bilingual = bilingual in ("yes", "bilingual") or (bilingual == "auto" and any(
        _is_bilingual_question(q.get('question', ''))
        for q in sample_questions[:5]
    ))

    # Process in batches
    BATCH_SIZE = 5
    all_generated = []

    for batch_start in range(0, len(sample_questions), BATCH_SIZE):
        batch = sample_questions[batch_start:batch_start + BATCH_SIZE]
        batch_num = batch_start // BATCH_SIZE + 1

        # Build prompt for batch
        if is_bilingual:
            # Step 1: Create English questions first
            batch_prompt = f"""Create {len(batch)} {subject_names.get(detected_subject, 'science')} multiple choice questions in ENGLISH ONLY.

Requirements:
- Create NEW questions similar in style and topic to the samples below
- Difficulty: {difficulty_text}
- Each question has 5 options (A through E)
- Return a SINGLE JSON array with ALL {len(batch)} questions

SAMPLES:
"""
        else:
            batch_prompt = f"""Bạn là chuyên gia tạo đề thi môn {subject_names.get(detected_subject, 'General')}.

Yêu cầu:
- Tạo câu hỏi MỚI HOÀN TOÀN cho từng câu hỏi mẫu bên dưới
- Câu hỏi mới phải {difficulty_text}
- KHÔNG được copy nguyên văn câu hỏi gốc
- Giữ nguyên số lượng đáp án và format
"""
        for i, q in enumerate(batch, 1):
            q_text = q.get('question', '')
            opts = q.get('options', [])
            # For bilingual, show only English part of sample
            if is_bilingual and ' / ' in q_text:
                q_text = q_text.split(' / ')[0]
            batch_prompt += f"\n--- Sample {batch_start + i} ---\n{q_text}\n"
            if opts:
                for j, opt in enumerate(opts):
                    opt_text = opt.split(' / ')[0] if is_bilingual and ' / ' in opt else opt
                    batch_prompt += f"{chr(65+j)}) {opt_text}\n"

        if is_bilingual:
            batch_prompt += f"""

IMPORTANT: Return ONLY a single JSON array with exactly {len(batch)} question objects. No extra text.
Format: [{{"question": "...", "options": ["A text", "B text", "C text", "D text", "E text"], "answer": "C"}}]"""
        else:
            batch_prompt += f"""

Trả về JSON array với format:
[
    {{"question": "câu hỏi mới", "options": ["A", "B", "C", "D"], "answer": "A"}}
]

Tạo ĐÚNG {len(batch)} câu hỏi mới. Chỉ trả về JSON array."""

        try:
            # Try up to 2 times to get enough questions
            batch_generated = []
            for attempt in range(2):
                response, error = funcs['call_ai'](batch_prompt, ai_engine)
                if not error and response:
                    response = re.sub(r'<think>[\s\S]*?</think>', '', response).strip()
                    batch_generated = _extract_questions_from_json(response)
                if len(batch_generated) >= len(batch):
                    break
            if batch_generated:
                all_generated.extend(batch_generated[:len(batch)])
                continue
            # Add placeholder for failed batch
            for _ in batch:
                all_generated.append({
                    "question": "Lỗi tạo câu hỏi",
                    "options": [],
                    "answer": ""
                })
        except Exception:
            for _ in batch:
                all_generated.append({
                    "question": "Lỗi tạo câu hỏi",
                    "options": [],
                    "answer": ""
                })

    # Step 2: If bilingual, translate to Vietnamese
    if is_bilingual and all_generated:
        all_generated = _translate_to_bilingual(all_generated, funcs['call_ai'], ai_engine)

    return {
        "ok": True,
        "filename": file.filename,
        "original_count": len(sample_questions),
        "generated_count": len(all_generated),
        "questions": all_generated,
        "detected_subject": detected_subject,
    }


@router.get("/api/answer-templates")
def get_answer_templates():
    """Get available answer templates for grading."""
    templates = []
    for key, value in ANSWER_TEMPLATES.items():
        templates.append({
            "key": key,
            "name": value.get("name", key),
            "questions": value.get("questions", 30),
            "options": value.get("options", 4),
            "layout": value.get("layout", "row")
        })
    return {"templates": templates}
