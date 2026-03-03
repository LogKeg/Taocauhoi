"""
Question generation API endpoints.
"""
import io
import os
import random
import re
from pathlib import Path
from typing import List

import httpx
import docx
from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, Form, HTTPException, Request
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from app.core import GenerateRequest, ParseSamplesRequest
from app.services.ai import call_ai, BASE_DIR
from app.services.text import normalize_name, normalize_question
from app.services.generation import (
    generate_variants,
    split_questions,
    load_questions_from_subject,
    build_topic_prompt,
    normalize_ai_blocks,
    retrieve_similar_questions,
)
from app.parsers.docx import extract_docx_content

router = APIRouter(tags=["generation"])


# ============================================================================
# Helper functions
# ============================================================================

def _resolve_sample_dir() -> Path:
    """Resolve the sample directory path."""
    env_dir = os.getenv("SAMPLE_DIR")
    if env_dir:
        path = Path(env_dir).expanduser().resolve()
        if path.exists() and path.is_dir():
            return path

    target = normalize_name("đề mẫu")
    for entry in BASE_DIR.iterdir():
        if entry.is_dir() and normalize_name(entry.name) == target:
            return entry
    return None


SAMPLE_DIR = _resolve_sample_dir()


def _read_sample_file(path: Path) -> str:
    """Read content from a sample file."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = Document(str(path))
        return extract_docx_content(doc)
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_sample_url(url: str) -> str:
    """Read content from a URL."""
    suffix = Path(url.split("?")[0]).suffix.lower()
    with httpx.Client(timeout=30) as client:
        resp = client.get(url)
        resp.raise_for_status()
        data = resp.content
    if suffix == ".docx":
        doc = Document(io.BytesIO(data))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(lines)
    return data.decode("utf-8", errors="ignore")


def _is_engine_available(engine: str) -> bool:
    """Check if an AI engine is available."""
    from app.services.ai import OPENAI_API_KEY, GEMINI_API_KEY
    if engine == "gemini":
        return bool(GEMINI_API_KEY)
    if engine == "ollama":
        return True
    return bool(OPENAI_API_KEY)


def _get_pdf_font():
    """Get a font that supports Vietnamese characters."""
    candidates = [
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/Library/Fonts/Times New Roman.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            font_name = "VietnameseFont"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, path))
            return font_name, path
    return "Helvetica", None


def _wrap_text(text: str, max_width: float, font_name: str, font_size: int) -> List[str]:
    """Wrap text to fit within max_width."""
    words = text.split()
    if not words:
        return [text]
    lines: List[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if pdfmetrics.stringWidth(trial, font_name, font_size) <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _save_text_questions_to_bank(
    questions: List[str],
    subject: str = "general",
    source: str = "generated",
    difficulty: str = "medium",
) -> int:
    """Save generated text questions to the question bank."""
    from app.database import SessionLocal, QuestionCRUD

    saved = 0
    db = SessionLocal()
    try:
        for q in questions:
            if not q.strip():
                continue
            QuestionCRUD.create(
                db,
                content=q.strip(),
                subject=subject,
                source=source,
                difficulty=difficulty,
            )
            saved += 1
    finally:
        db.close()
    return saved


# ============================================================================
# Generation endpoints
# ============================================================================

@router.post("/generate")
def generate(payload: GenerateRequest) -> dict:
    """Generate question variants from samples."""
    engine = payload.ai_engine
    if payload.use_ai and not _is_engine_available(engine):
        engine_names = {"openai": "OPENAI_API_KEY", "gemini": "GEMINI_API_KEY", "ollama": "Ollama"}
        questions = generate_variants(payload)
        saved = _save_text_questions_to_bank(
            questions,
            subject=payload.topic or "general",
            source="generated",
        )
        return {
            "questions": questions,
            "message": f"Chưa cấu hình {engine_names.get(engine, engine)} nên AI không được dùng.",
            "saved_to_bank": saved,
        }

    questions = generate_variants(payload)
    saved = _save_text_questions_to_bank(
        questions,
        subject=payload.topic or "general",
        source="generated-ai" if payload.use_ai else "generated",
    )

    if payload.use_ai and _is_engine_available(engine):
        src = {normalize_question(s) for s in payload.samples if s.strip()}
        out = {normalize_question(q) for q in questions if q.strip()}
        if out and out.issubset(src):
            return {
                "questions": questions,
                "message": "AI đang trả về câu gần giống câu gốc. Hệ thống đã thêm biến thể đơn giản.",
                "saved_to_bank": saved,
            }
    return {"questions": questions, "saved_to_bank": saved}


@router.post("/generate-topic")
def generate_topic(
    subject: str = Form(...),
    grade: int = Form(1),
    qtype: str = Form("mcq"),
    count: int = Form(10),
    ai_engine: str = Form("ollama"),
    topic: str = Form(""),
    difficulty: str = Form("medium"),
    use_rag: bool = Form(True),
    rag_count: int = Form(5),
) -> dict:
    """Generate questions by topic using AI."""
    if not _is_engine_available(ai_engine):
        engine_names = {"openai": "OPENAI_API_KEY", "gemini": "GEMINI_API_KEY", "ollama": "Ollama"}
        return {"questions": [], "message": f"Chưa cấu hình {engine_names.get(ai_engine, ai_engine)} nên AI không được dùng."}

    count = max(1, min(50, count))
    grade = max(1, min(12, grade))
    rag_count = max(1, min(10, rag_count))

    # Retrieve RAG examples from question bank
    rag_examples = []
    if use_rag:
        rag_examples = retrieve_similar_questions(
            subject=subject,
            topic=topic,
            difficulty=difficulty,
            question_type=qtype,
            limit=rag_count,
        )

    prompt = build_topic_prompt(subject, grade, qtype, count, topic, difficulty, rag_examples=rag_examples)
    text, err = call_ai(prompt, ai_engine)

    if not text:
        msg = f"Không nhận được phản hồi từ AI. {err}" if err else "Không nhận được phản hồi từ AI."
        return {"questions": [], "answers": "", "message": msg}

    # Split answers section from questions
    answers = ""
    answer_pattern = re.compile(
        r"\n\s*-{0,3}\s*(?:ĐÁP ÁN|Đáp án|đáp án|ANSWERS|Answers|Answer Key|answer key)\s*-{0,3}\s*:?\s*\n",
        re.IGNORECASE,
    )
    match = answer_pattern.search(text)
    if match:
        raw_answers = text[match.end():].strip()
        text = text[:match.start()]
        # Clean up answers: extract only answer lines like "1. A", "2. B", etc.
        answer_lines = []
        for line in raw_answers.splitlines():
            line = line.strip()
            # Match patterns like "1. A", "1) B", "1: C", "Câu 1: A", etc.
            ans_match = re.match(r'^(?:Câu\s*)?(\d+)[\.\):\s]+([A-Da-d])\b', line, re.IGNORECASE)
            if ans_match:
                answer_lines.append(f"{ans_match.group(1)}. {ans_match.group(2).upper()}")
        answers = "\n".join(answer_lines) if answer_lines else raw_answers

    questions = normalize_ai_blocks(text)
    questions = [q for q in questions if q.strip()]
    final_questions = questions[:count]

    saved = _save_text_questions_to_bank(
        final_questions,
        subject=subject,
        source="generated-topic",
        difficulty=difficulty,
    )

    return {"questions": final_questions, "answers": answers, "saved_to_bank": saved}


@router.post("/auto-generate")
def auto_generate(
    subject: str = Form(...),
    count: int = Form(10),
    ai_ratio: int = Form(30),
    topic: str = Form(""),
    custom_keywords: str = Form(""),
    paraphrase: bool = Form(True),
    change_numbers: bool = Form(True),
    change_context: bool = Form(True),
    use_ai: bool = Form(False),
    samples_text: str = Form(""),
    ai_engine: str = Form("openai"),
) -> dict:
    """Auto-generate questions with configurable AI ratio."""
    if samples_text.strip():
        samples = split_questions(samples_text)
    else:
        samples = load_questions_from_subject(subject)

    if not samples:
        return {"questions": [], "message": "Không tìm thấy câu hỏi mẫu."}

    count = max(1, min(200, count))
    ai_ratio = max(0, min(100, ai_ratio))
    ai_count = int(round(count * (ai_ratio / 100)))
    rule_count = count - ai_count

    random.shuffle(samples)
    selected = samples[: min(len(samples), max(1, rule_count))]
    if len(selected) < rule_count:
        selected = selected * (rule_count // max(1, len(selected)) + 1)
        selected = selected[:rule_count]

    req = GenerateRequest(
        samples=selected,
        topic=topic or subject,
        custom_keywords=[s for s in custom_keywords.split(",") if s.strip()],
        paraphrase=paraphrase,
        change_numbers=change_numbers,
        change_context=change_context,
        variants_per_question=1,
        use_ai=False,
    )
    questions = generate_variants(req)

    if use_ai and ai_count > 0 and not _is_engine_available(ai_engine):
        engine_names = {"openai": "OPENAI_API_KEY", "gemini": "GEMINI_API_KEY", "ollama": "Ollama"}
        return {
            "questions": questions[:count],
            "message": f"Chưa cấu hình {engine_names.get(ai_engine, ai_engine)} nên AI không được dùng.",
        }

    if use_ai and ai_count > 0:
        ai_req = GenerateRequest(
            samples=[random.choice(samples)] if samples else [],
            topic=topic or subject,
            custom_keywords=[s for s in custom_keywords.split(",") if s.strip()],
            paraphrase=paraphrase,
            change_numbers=change_numbers,
            change_context=change_context,
            variants_per_question=max(1, ai_count),
            use_ai=True,
            ai_engine=ai_engine,
        )
        ai_questions = generate_variants(ai_req)
        if ai_questions:
            questions.extend(ai_questions[:ai_count])

    random.shuffle(questions)
    final_questions = questions[:count]

    saved = _save_text_questions_to_bank(
        final_questions,
        subject=subject,
        source="auto-generated",
    )

    result = {"questions": final_questions, "saved_to_bank": saved}
    if use_ai and _is_engine_available(ai_engine):
        src = {normalize_question(s) for s in samples if s.strip()}
        out = {normalize_question(q) for q in result["questions"] if q.strip()}
        if out and out.issubset(src):
            result["message"] = "AI đang trả về câu gần giống câu gốc. Hệ thống đã thêm biến thể đơn giản."
    return result


@router.post("/export")
def export(
    samples: str = Form(...),
    topic: str = Form(...),
    custom_keywords: str = Form(""),
    paraphrase: bool = Form(False),
    change_numbers: bool = Form(False),
    change_context: bool = Form(False),
    variants_per_question: int = Form(1),
    fmt: str = Form("txt"),
    use_ai: bool = Form(False),
    ai_engine: str = Form("openai"),
):
    """Export generated questions to various formats."""
    from app.services.math_renderer import render_text_with_math, extract_latex_parts, has_latex

    # Export directly without regenerating - just use the samples as-is
    questions = [s.strip() for s in samples.split("\n\n") if s.strip()]
    if not questions:
        # Fallback: split by single newline if no double newlines
        questions = [s.strip() for s in samples.split("\n") if s.strip()]

    if fmt == "csv":
        # Convert LaTeX to Unicode for CSV
        csv_questions = [render_text_with_math(q) for q in questions]
        content = "id,question\n" + "\n".join(
            f"{i+1},\"{q.replace('\"', '\"\"')}\"" for i, q in enumerate(csv_questions)
        )
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=questions.csv"},
        )

    if fmt == "docx":
        doc = Document()
        doc.add_heading("Danh sách câu hỏi", level=1)
        for i, q in enumerate(questions, 1):
            # Split question into lines (first line = question, rest = options)
            lines = q.split('\n')
            p = doc.add_paragraph()
            p.add_run(f"Câu {i}. ").bold = True

            # Render first line with math
            _add_text_with_math(p, lines[0])

            # Add options on separate lines
            for line in lines[1:]:
                line = line.strip()
                if line:
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.3)
                    _add_text_with_math(opt_p, line)
            doc.add_paragraph("")  # Empty line between questions
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=questions.docx"},
        )

    if fmt == "pdf":
        from app.services.math_renderer import has_latex

        font_name, _ = _get_pdf_font()
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 60
        font_size = 12
        c.setFont(font_name, font_size)
        c.drawString(50, y, "Danh sách câu hỏi")
        y -= 30
        for i, q in enumerate(questions, 1):
            # Split question into lines (first line = question, rest = options)
            lines = q.split('\n')

            # Check page break
            if y < 80:
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - 60

            # Write question number
            c.drawString(50, y, f"Câu {i}. ")
            x_offset = 50 + pdfmetrics.stringWidth(f"Câu {i}. ", font_name, font_size)

            # Write question content with math
            if has_latex(lines[0]):
                _draw_text_with_math_pdf(c, lines[0], x_offset, y, font_name, font_size, width - x_offset - 50)
            else:
                # No LaTeX - use simple text wrapping
                q_text = render_text_with_math(lines[0])
                remaining_width = width - x_offset - 50
                if pdfmetrics.stringWidth(q_text, font_name, font_size) <= remaining_width:
                    c.drawString(x_offset, y, q_text)
                else:
                    # Need to wrap
                    wrapped = _wrap_text(q_text, remaining_width, font_name, font_size)
                    c.drawString(x_offset, y, wrapped[0])
                    for part in wrapped[1:]:
                        y -= 18
                        if y < 60:
                            c.showPage()
                            c.setFont(font_name, font_size)
                            y = height - 60
                        c.drawString(50, y, part)
            y -= 20

            # Write options
            for line in lines[1:]:
                line = line.strip()
                if line:
                    if y < 60:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = height - 60

                    if has_latex(line):
                        c.drawString(70, y, "")  # Position marker
                        _draw_text_with_math_pdf(c, line, 70, y, font_name, font_size, width - 120)
                    else:
                        opt_text = render_text_with_math(line)
                        c.drawString(70, y, opt_text)
                    y -= 16
            y -= 10  # Space between questions
        c.save()
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=questions.pdf"},
        )

    # TXT format - convert LaTeX to Unicode
    txt_questions = [render_text_with_math(q) for q in questions]
    content = "\n".join(txt_questions)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=questions.txt"},
    )


def _draw_text_with_math_pdf(c, text: str, x: float, y: float, font_name: str, font_size: int, max_width: float) -> float:
    """
    Draw text with math formulas on PDF canvas.
    Returns the final x position after drawing.
    """
    from app.services.math_renderer import extract_latex_parts, latex_to_image, latex_to_unicode
    from reportlab.lib.utils import ImageReader

    parts = extract_latex_parts(text)
    current_x = x

    for part_text, is_latex in parts:
        if is_latex:
            # Try to render as image
            img_bytes = latex_to_image(part_text, dpi=150, fontsize=font_size + 2)
            if img_bytes:
                try:
                    img = ImageReader(io.BytesIO(img_bytes))
                    img_width, img_height = img.getSize()

                    # Scale down for PDF
                    scale = 0.48
                    draw_width = img_width * scale
                    draw_height = img_height * scale

                    # Check if fits in remaining width
                    if current_x + draw_width > x + max_width:
                        # Would overflow - use unicode fallback
                        unicode_text = latex_to_unicode(part_text)
                        c.drawString(current_x, y, unicode_text)
                        current_x += pdfmetrics.stringWidth(unicode_text, font_name, font_size)
                    else:
                        # Center vertically relative to text baseline
                        y_offset = (draw_height - font_size) / 3

                        c.drawImage(img, current_x, y - y_offset,
                                   width=draw_width, height=draw_height,
                                   mask='auto')
                        current_x += draw_width + 2
                except Exception:
                    # Fallback to Unicode
                    unicode_text = latex_to_unicode(part_text)
                    c.drawString(current_x, y, unicode_text)
                    current_x += pdfmetrics.stringWidth(unicode_text, font_name, font_size)
            else:
                # Fallback to Unicode
                unicode_text = latex_to_unicode(part_text)
                c.drawString(current_x, y, unicode_text)
                current_x += pdfmetrics.stringWidth(unicode_text, font_name, font_size)
        else:
            c.drawString(current_x, y, part_text)
            current_x += pdfmetrics.stringWidth(part_text, font_name, font_size)

    return current_x


def _add_text_with_math(paragraph, text: str):
    """Add text with math formulas to a Word paragraph using native OMML equations."""
    from app.services.math_renderer import extract_latex_parts, latex_to_unicode, add_math_to_paragraph

    parts = extract_latex_parts(text)

    for part_text, is_latex in parts:
        if is_latex:
            # Try to add as native OMML equation
            success = add_math_to_paragraph(paragraph, part_text)
            if not success:
                # Fallback to Unicode representation
                unicode_text = latex_to_unicode(part_text)
                run = paragraph.add_run(unicode_text)
                run.italic = True
        else:
            paragraph.add_run(part_text)


@router.post("/api/export-exam")
async def export_exam(request: Request):
    """Export exam questions with full content and options."""
    from app.services.math_renderer import render_text_with_math

    data = await request.json()
    questions = data.get("questions", [])
    fmt = data.get("format", "docx")
    title = data.get("title", "Đề thi")

    if not questions:
        raise HTTPException(status_code=400, detail="Không có câu hỏi để xuất")

    if fmt == "docx":
        doc = Document()
        doc.add_heading(title, level=1)

        for i, q in enumerate(questions, 1):
            content = q.get("content", "")
            options = q.get("options", [])

            p = doc.add_paragraph()
            p.add_run(f"Câu {i}. ").bold = True
            _add_text_with_math(p, content)

            if options:
                labels = ['A', 'B', 'C', 'D', 'E']
                for j, opt in enumerate(options[:5]):
                    opt_clean = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', str(opt).strip())
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.3)
                    _add_text_with_math(opt_p, f"{labels[j]}) {opt_clean}")

            doc.add_paragraph("")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=de_thi.docx"},
        )

    if fmt == "pdf":
        from app.services.math_renderer import has_latex

        font_name, _ = _get_pdf_font()
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 60
        font_size = 12

        c.setFont(font_name, 16)
        c.drawString(50, y, title)
        y -= 40
        c.setFont(font_name, font_size)

        for i, q in enumerate(questions, 1):
            content = q.get("content", "")
            options = q.get("options", [])

            # Check page break
            if y < 80:
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - 60

            # Write question number
            q_prefix = f"Câu {i}. "
            c.drawString(50, y, q_prefix)
            x_offset = 50 + pdfmetrics.stringWidth(q_prefix, font_name, font_size)

            # Write question content with math rendering
            if has_latex(content):
                _draw_text_with_math_pdf(c, content, x_offset, y, font_name, font_size, width - x_offset - 50)
            else:
                content_text = render_text_with_math(content)
                remaining_width = width - x_offset - 50
                if pdfmetrics.stringWidth(content_text, font_name, font_size) <= remaining_width:
                    c.drawString(x_offset, y, content_text)
                else:
                    wrapped = _wrap_text(content_text, remaining_width, font_name, font_size)
                    c.drawString(x_offset, y, wrapped[0])
                    for part in wrapped[1:]:
                        y -= 18
                        if y < 60:
                            c.showPage()
                            c.setFont(font_name, font_size)
                            y = height - 60
                        c.drawString(50, y, part)
            y -= 20

            # Write options with math rendering
            if options:
                labels = ['A', 'B', 'C', 'D', 'E']
                for j, opt in enumerate(options[:5]):
                    opt_clean = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', str(opt).strip())
                    opt_line = f"{labels[j]}) {opt_clean}"

                    if y < 60:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = height - 60

                    if has_latex(opt_line):
                        _draw_text_with_math_pdf(c, opt_line, 70, y, font_name, font_size, width - 120)
                    else:
                        opt_text = render_text_with_math(opt_line)
                        c.drawString(70, y, opt_text)
                    y -= 16

            y -= 10

        c.save()
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=de_thi.pdf"},
        )

    raise HTTPException(status_code=400, detail="Format không hỗ trợ")


# ============================================================================
# Sample file endpoints
# ============================================================================

@router.get("/sample-folders")
def list_sample_folders() -> dict:
    """List available sample folders."""
    if SAMPLE_DIR is None or not SAMPLE_DIR.exists():
        return {"folders": []}
    folders = [p.name for p in SAMPLE_DIR.iterdir() if p.is_dir()]
    folders.sort()
    return {"folders": folders}


@router.get("/sample-files")
def list_sample_files(subject: str) -> dict:
    """List sample files in a subject folder."""
    if SAMPLE_DIR is None or not subject or not SAMPLE_DIR.exists():
        return {"files": []}
    subject_dir = (SAMPLE_DIR / subject).resolve()
    if SAMPLE_DIR not in subject_dir.parents or not subject_dir.exists():
        return {"files": []}
    files = [
        p.name
        for p in subject_dir.iterdir()
        if p.is_file() and p.suffix.lower() in {".txt", ".docx", ".md"}
    ]
    files.sort()
    return {"files": files}


@router.get("/sample-content")
def sample_content(subject: str, filename: str) -> dict:
    """Get content of a sample file."""
    if not subject or not filename:
        return {"content": ""}
    if SAMPLE_DIR is None:
        return {"content": ""}
    subject_dir = (SAMPLE_DIR / subject).resolve()
    if not SAMPLE_DIR.exists() or SAMPLE_DIR not in subject_dir.parents:
        return {"content": ""}
    target = (subject_dir / filename).resolve()
    if subject_dir not in target.parents or not target.exists() or not target.is_file():
        return {"content": ""}
    content = _read_sample_file(target)
    return {"content": content}


@router.post("/parse-sample-urls")
def parse_sample_urls(payload: ParseSamplesRequest) -> dict:
    """Parse sample questions from URLs."""
    contents: List[str] = []
    for url in payload.urls:
        if not url:
            continue
        try:
            contents.append(_read_sample_url(url))
        except Exception:
            continue
    merged = "\n".join(contents)
    return {"content": merged, "samples": split_questions(merged)}
