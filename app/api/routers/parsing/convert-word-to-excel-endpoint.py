"""
Word to Excel conversion endpoint.

POST /convert-word-to-excel - Convert Word file containing questions to Excel format.
"""
import io
import re
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.oxml.ns import qn as docx_qn
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from app.api.routers.parsing.helpers import _get_parsing_functions

router = APIRouter(tags=["parsing"])


def _parse_answer_key(answer_key_content: bytes, answer_key_filename: str, question_filename: str) -> Dict[int, int]:
    """Parse answer key file and return answers for the matching level.

    Returns dict mapping question_number -> answer_number (A=1, B=2, C=3, D=4, E=5).
    """
    answer_map = {'A': 1, 'B': 2, 'C': 3, 'D': 4, 'E': 5}
    answers = {}

    try:
        if answer_key_filename.lower().endswith('.docx'):
            doc = Document(io.BytesIO(answer_key_content))
            for table in doc.tables:
                if len(table.rows) < 2 or len(table.columns) < 2:
                    continue

                # Parse header row to find level columns
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]

                # Find matching level based on question filename
                # e.g., "Benjamin.docx" -> match "Benjamin" in header
                q_name = question_filename.rsplit('.', 1)[0].lower().strip()
                level_col = -1
                for col_idx, header in enumerate(header_cells):
                    if header.strip().lower() == q_name:
                        level_col = col_idx
                        break

                if level_col < 0:
                    # Try partial match (e.g., "PreEcolier" in "Pre-Ecolier")
                    for col_idx, header in enumerate(header_cells):
                        h = header.strip().lower().replace('-', '').replace(' ', '')
                        qn = q_name.replace('-', '').replace(' ', '')
                        if h == qn:
                            level_col = col_idx
                            break

                if level_col < 0:
                    continue

                # Answer is in the column after the question number column
                # Format: [Q#, Answer, Q#, Answer, ...]
                # Ensure level_col is the question number column (even index in pair)
                if level_col % 2 == 1:
                    level_col -= 1  # Go back to question number column

                answer_col = level_col + 1

                # Parse answer rows
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if answer_col >= len(cells):
                        continue
                    q_num_text = cells[level_col].strip()
                    answer_text = cells[answer_col].strip().upper()
                    if q_num_text and q_num_text.isdigit() and answer_text in answer_map:
                        answers[int(q_num_text)] = answer_map[answer_text]

        elif answer_key_filename.lower().endswith(('.xlsx', '.xls')):
            # Excel support: columns are levels, rows are Q# + answer
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(answer_key_content), read_only=True)
            ws = wb.active

            rows_data = list(ws.iter_rows(values_only=True))
            if rows_data:
                # Row 0 is header: [Q#_or_empty, Level1, Level2, ...]
                headers = [str(h).strip().lower() if h else '' for h in rows_data[0]]

                # Match level based on question filename
                q_name = question_filename.rsplit('.', 1)[0].lower().strip()
                # Remove common suffixes like " Edited 1-2026"
                q_name_clean = re.sub(r'\s*edited.*$', '', q_name, flags=re.IGNORECASE).strip()

                level_col = -1
                for col_idx, header in enumerate(headers):
                    h = header.replace('-', '').replace(' ', '')
                    qn = q_name_clean.replace('-', '').replace(' ', '')
                    if h == qn or h == q_name.replace('-', '').replace(' ', ''):
                        level_col = col_idx
                        break

                if level_col >= 0:
                    for row in rows_data[1:]:
                        if len(row) <= level_col:
                            continue
                        q_num_val = row[0]
                        answer_val = row[level_col]
                        if q_num_val is not None and answer_val is not None:
                            q_num_str = str(q_num_val).strip()
                            answer_text = str(answer_val).strip().upper()
                            if q_num_str.isdigit() and answer_text in answer_map:
                                answers[int(q_num_str)] = answer_map[answer_text]
            wb.close()

        elif answer_key_filename.lower().endswith('.pdf'):
            # PDF support: extract text and find answer patterns
            try:
                import fitz  # PyMuPDF
                pdf_doc = fitz.open(stream=answer_key_content, filetype="pdf")
                full_text = ""
                for page in pdf_doc:
                    full_text += page.get_text()
                pdf_doc.close()

                # Find answer patterns like "1. A", "2. B", "1 A", "1\tA"
                for match in re.finditer(r'(\d+)\s*[.\t)]\s*([A-E])\b', full_text):
                    q_num = int(match.group(1))
                    answer_letter = match.group(2).upper()
                    if answer_letter in answer_map:
                        answers[q_num] = answer_map[answer_letter]
            except ImportError:
                pass  # PyMuPDF not installed, skip PDF parsing

    except Exception:
        pass  # Return empty answers on any error

    return answers


@router.post("/convert-word-to-excel")
def convert_word_to_excel(
    file: UploadFile = Form(...),
    use_latex: str = Form("0"),
    subject: str = Form("general"),
    answer_key: Optional[UploadFile] = File(None),
) -> StreamingResponse:
    """Convert a Word file containing questions to Excel format."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    funcs = _get_parsing_functions()

    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")

    latex_enabled = use_latex == "1"

    try:
        content = file.file.read()
        doc = Document(io.BytesIO(content))

        # Try cell-based parser first
        questions = funcs['parse_cell_based_questions'](doc)

        is_math_format = False
        is_english_level_format = False
        is_envie_format = False

        if not questions:
            lines, table_options = funcs['extract_docx_lines'](doc, include_textboxes=True, use_latex=latex_enabled)

            is_math_format = any(re.match(r'^Question\s+\d+', line, re.IGNORECASE) for line in lines[:20])
            is_english_level_format = (
                any('Section A' in line or 'Section B' in line for line in lines[:15]) and
                file.filename and 'LEVEL' in file.filename.upper()
            )
            is_envie_format = file.filename and 'EN-VIE' in file.filename.upper()

            # Detect bilingual science (IKSC) format
            if not is_envie_format:
                # Check bilingual options (e.g., "A. Refraction / Khúc xạ")
                bilingual_option_count = sum(
                    1 for line in lines[:80]
                    if re.match(r'^[A-E][.)]\s*.+\s*/\s*.+', line)
                )
                # Also check for EN+VN question pairs (numbered EN → VN translation)
                bilingual_pair_count = 0
                for li in range(len(lines) - 1):
                    l1, l2 = lines[li], lines[li + 1]
                    if (re.match(r'^\d+\.\s*', l1)
                            and not any(ord(c) > 127 for c in l1[:200])
                            and any(ord(c) > 127 for c in l2[:200])
                            and not re.match(r'^[A-E][.)]\s*', l2)):
                        bilingual_pair_count += 1
                        if bilingual_pair_count >= 3:
                            break
                if bilingual_option_count >= 3 or bilingual_pair_count >= 3:
                    is_envie_format = True

            # Check for Word numbered list format
            is_word_numbered_format = False
            if not is_envie_format:
                body = doc._element.body
                num_level_0 = 0
                num_level_1 = 0
                for child in body:
                    tag = child.tag.split('}')[-1]
                    if tag == 'p':
                        pPr = child.find(docx_qn('w:pPr'))
                        if pPr is not None:
                            numPr = pPr.find(docx_qn('w:numPr'))
                            if numPr is not None:
                                ilvl_elem = numPr.find(docx_qn('w:ilvl'))
                                if ilvl_elem is not None:
                                    ilvl = int(ilvl_elem.get(docx_qn('w:val')) or '0')
                                    if ilvl == 0:
                                        num_level_0 += 1
                                    elif ilvl == 1:
                                        num_level_1 += 1
                if num_level_0 > 10 and num_level_0 == num_level_1:
                    is_word_numbered_format = True

            # Detect Vietnamese "Câu X:" exam format
            is_vn_cau_format = False
            if not is_envie_format and not is_word_numbered_format:
                cau_count = sum(1 for line in lines[:80] if re.match(r'^Câu\s+\d+', line, re.IGNORECASE))
                if cau_count >= 3:
                    is_vn_cau_format = True

            # Detect English reading/competition exam format
            is_english_reading_format = False
            if not is_math_format and not is_english_level_format and not is_envie_format and not is_word_numbered_format and not is_vn_cau_format:
                text_sample = ' '.join(lines[:50]).lower()
                has_instruction = any(
                    kw in text_sample
                    for kw in ['for each question', 'read the text', 'choose the correct answer', 'choose the best answer']
                )
                has_merged_options = any(
                    re.search(r'[A-D]\)\s*\w', line) for line in lines[:80]
                )
                if has_instruction or has_merged_options:
                    is_english_reading_format = True

            if is_math_format:
                questions = funcs['parse_math_exam_questions'](lines)
            elif is_english_level_format:
                questions = funcs['parse_english_exam_questions'](doc)
            elif is_envie_format or is_word_numbered_format or is_english_reading_format or is_vn_cau_format:
                questions = funcs['parse_envie_questions'](doc)
            else:
                questions = funcs['parse_bilingual_questions'](lines, table_options)

        # Post-process: split merged options like "A. xxx   B. xxx" into separate options
        _opt_split_re = re.compile(
            r'\s{2,}(?=[A-E][\.\)]\s)'        # 2+ spaces before "B. "
            r'|(?<=\S)\t+(?=[A-E][\.\)])'      # tab before "B."
            r'|(?<=\S)\s+(?=[B-E][\.\)]\s)'    # single space before "B. " (not A)
            r'|(?<=\.)(?=[B-E][\.\)]\s*\S)'    # "text.B. text"
            r'|(?<=[a-z\u00E0-\u1EF9])(?=[B-E][\.\)]\s*\S)'  # "textB. text" (lowercase→B.)
            r'|(?<=\S)(?=[B-E]\.\xa0)'         # non-breaking space: "textB.\xa0text"
        )
        for q in questions:
            opts = q.get('options', [])
            if opts and len(opts) < 4:
                split_opts = []
                for opt in opts:
                    parts = _opt_split_re.split(opt)
                    for part in parts:
                        part = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', part.strip())
                        if part:
                            split_opts.append(part)
                if len(split_opts) > len(opts):
                    q['options'] = split_opts

        # For science subject, apply bilingual dedup (IKSC format: EN+VN pairs)
        if subject == 'science' and questions:
            from app.parsers import _dedup_bilingual_science
            questions = _dedup_bilingual_science(questions)

        # Save to question bank
        from app.api.routers.parsing.helpers import _save_questions_to_bank
        detected_subject = subject if subject in ('english', 'math', 'science', 'general') else 'general'
        _save_questions_to_bank(
            questions,
            subject=detected_subject,
            source=f"word-import:{file.filename}",
            question_type='mcq' if questions and questions[0].get('options') else 'blank',
        )

        # Sort questions by number if most questions have a number assigned
        numbered_count = sum(1 for q in questions if q.get('number') is not None)
        if numbered_count > len(questions) * 0.5:
            questions.sort(key=lambda q: q.get('number') or float('inf'))

        # Parse answer key if provided
        answer_key_data = {}
        if answer_key is not None and answer_key.filename:
            ak_content = answer_key.file.read()
            if ak_content:
                answer_key_data = _parse_answer_key(ak_content, answer_key.filename, file.filename)

        # Create Excel workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Questions"

        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        highlight_fill = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
        center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Standard format: Question Type, Question, Option 1-5, Correct Answer,
        # Default Marks, Default Time To Solve, Difficulty Level, Hint, Solution
        headers = [
            "Question Type", "Question",
            "Option 1", "Option 2", "Option 3", "Option 4", "Option 5",
            "Correct Answer", "Default Marks", "Default Time To Solve",
            "Difficulty Level", "Hint", "Solution",
        ]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border

        # Determine level from filename for IKLC/IKSC scoring
        raw_name = file.filename.rsplit('.', 1)[0].strip().lower() if file.filename else ''
        total_questions = len(questions)

        # Extract level name from filename (e.g., "Benjamin Edited 1-2026" → "benjamin")
        level_name = raw_name
        for lvl in ['preecolier', 'pre-ecolier', 'ecolier', 'benjamin', 'cadet', 'junior', 'student']:
            if lvl in raw_name.replace('-', '').replace(' ', ''):
                level_name = lvl
                break

        # IKLC English scoring rules per level
        iklc_marks = {
            'preecolier': {'correct': 2.00, 'wrong': 0},
            'pre-ecolier': {'correct': 2.00, 'wrong': 0},
            'ecolier': {'correct': 2.00, 'wrong': 0},
            'benjamin': {'correct': 1.00, 'wrong': -0.50},
            'cadet': {'correct': 1.25, 'wrong': -0.50},
            'junior': {'correct': 1.50, 'wrong': -0.50},
            'student': {'correct': 1.75, 'wrong': -0.50},
        }

        # IKSC Science scoring: difficulty-based (easy=3, medium=4, hard=5), wrong=-1
        # Questions split into 3 equal groups by position

        def get_default_marks(q_idx_0based):
            """Return (default_mark, difficulty_level) for a question."""
            if subject == 'english' and level_name in iklc_marks:
                marks_info = iklc_marks[level_name]
                return marks_info['correct'], "EASY"
            elif subject == 'science':
                # IKSC: Split into 3 equal groups (easy=3, medium=4, hard=5)
                group_size = total_questions // 3 if total_questions >= 3 else total_questions
                if q_idx_0based < group_size:
                    return 3, "EASY"
                elif q_idx_0based < group_size * 2:
                    return 4, "MEDIUM"
                else:
                    return 5, "HARD"
            return None, "EASY"

        for row_idx, q in enumerate(questions, 2):
            question_text = q.get('question', '')
            options = q.get('options', [])
            answer = q.get('answer', '')
            q_idx_0based = row_idx - 2  # 0-based question index

            # Determine question type: MSA = trắc nghiệm, SAQ = tự luận
            q_type = "MSA" if options else "SAQ"

            # A: Question Type
            cell = ws.cell(row=row_idx, column=1, value=q_type)
            cell.alignment = center_align
            cell.border = thin_border

            # B: Question
            cell = ws.cell(row=row_idx, column=2, value=question_text)
            cell.alignment = left_align
            cell.border = thin_border

            # Determine correct answer from answer key or parsed answer
            # answer_key_data maps question_number -> answer_number (1-5)
            q_num = row_idx - 1  # 1-based question number
            correct_answer_num = answer_key_data.get(q_num)

            # Fallback to parsed answer if no answer key
            if correct_answer_num is None and answer and isinstance(answer, str):
                answer_upper = answer.upper().strip()
                if answer_upper in ('A', 'B', 'C', 'D', 'E'):
                    correct_answer_num = ord(answer_upper) - ord('A') + 1

            # C-G: Option 1-5
            for opt_idx in range(5):
                cell = ws.cell(row=row_idx, column=3 + opt_idx,
                               value=options[opt_idx] if opt_idx < len(options) else None)
                cell.alignment = left_align
                cell.border = thin_border

                # Highlight correct answer option
                if opt_idx < len(options) and correct_answer_num and (opt_idx + 1) == correct_answer_num:
                    cell.fill = highlight_fill

            # H: Correct Answer (A=1, B=2, C=3, D=4, E=5)
            cell = ws.cell(row=row_idx, column=8, value=correct_answer_num)
            cell.alignment = center_align
            cell.border = thin_border

            # I: Default Marks & K: Difficulty Level (based on subject/level)
            default_mark, difficulty = get_default_marks(q_idx_0based)
            cell = ws.cell(row=row_idx, column=9, value=default_mark)
            cell.alignment = center_align
            cell.border = thin_border

            # J: Default Time To Solve
            cell = ws.cell(row=row_idx, column=10, value=30)
            cell.alignment = center_align
            cell.border = thin_border

            # K: Difficulty Level
            cell = ws.cell(row=row_idx, column=11, value=difficulty)
            cell.alignment = center_align
            cell.border = thin_border

            # L: Hint
            cell = ws.cell(row=row_idx, column=12, value=None)
            cell.border = thin_border

            # M: Solution
            cell = ws.cell(row=row_idx, column=13, value=None)
            cell.border = thin_border

        # Adjust column widths
        ws.column_dimensions["A"].width = 15   # Question Type
        ws.column_dimensions["B"].width = 60   # Question
        ws.column_dimensions["C"].width = 25   # Option 1
        ws.column_dimensions["D"].width = 25   # Option 2
        ws.column_dimensions["E"].width = 25   # Option 3
        ws.column_dimensions["F"].width = 25   # Option 4
        ws.column_dimensions["G"].width = 25   # Option 5
        ws.column_dimensions["H"].width = 15   # Correct Answer
        ws.column_dimensions["I"].width = 14   # Default Marks
        ws.column_dimensions["J"].width = 20   # Default Time To Solve
        ws.column_dimensions["K"].width = 16   # Difficulty Level
        ws.column_dimensions["L"].width = 20   # Hint
        ws.column_dimensions["M"].width = 20   # Solution

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        base_name = Path(file.filename).stem
        safe_name = f"{base_name}_questions.xlsx"
        from urllib.parse import quote
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe_name)}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý file: {str(e)}")
