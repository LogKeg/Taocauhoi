"""OMR Grader: functions to grade answer sheets and parse answer keys.

Contains:
- _grade_single_sheet: grade a single OMR sheet with improved bubble detection
- _grade_mixed_format_sheet: grade mixed MCQ + fill-in answer sheets (e.g. SEAMO)
- _extract_answers_from_text: extract answers from plain text (PDF/Word fallback)
- _parse_answer_key_for_template: parse answer key file for a specific template
"""

import io
import re
from typing import List

from app.core import ANSWER_TEMPLATES
from app.services.omr.omr_template_detector_and_student_info_ocr import _extract_student_info_ocr
from app.services.omr.omr_bubble_detection_grid_fill_analysis_and_grouping import (
    _detect_bubbles_grid_based,
    _group_bubbles_to_questions_improved,
    _detect_bubbles,
    _group_bubbles_to_questions,
    _analyze_bubble_fill_improved,
    _analyze_bubble_fill,
    _detect_seamo_bubbles_fixed_grid,
)
from app.services.image import _preprocess_omr_image


def _get_easyocr_reader():
    """Create and return an EasyOCR reader instance (English only, CPU mode)."""
    import ssl
    ssl._create_default_https_context = ssl._create_unverified_context
    import easyocr
    return easyocr.Reader(['en'], gpu=False, verbose=False)


def _grade_single_sheet(image_bytes: bytes, answer_key: List[str], template_type: str = "IKSC_BENJAMIN", extract_info: bool = True):
    """Chấm một phiếu trả lời với thuật toán OMR cải tiến"""
    import cv2
    import numpy as np

    template = ANSWER_TEMPLATES.get(template_type, ANSWER_TEMPLATES["IKSC_BENJAMIN"])
    num_questions = template["questions"]
    scoring = template["scoring"]
    option_labels = ["A", "B", "C", "D", "E"]

    # Trích xuất thông tin học sinh bằng OCR
    student_info = {}
    if extract_info:
        try:
            student_info = _extract_student_info_ocr(image_bytes)
        except Exception:
            pass

    # Tiền xử lý ảnh với deskew và perspective correction
    result = _preprocess_omr_image(image_bytes)
    if result[0] is None:
        return {"error": "Không thể đọc ảnh"}

    original, gray, binary = result

    # Thử phương pháp mới trước: phát hiện dựa trên grid
    rows, all_rects = _detect_bubbles_grid_based(gray, binary, template_type)

    questions = []
    use_new_method = False

    if rows and len(rows) >= 3:
        # Sử dụng phương pháp mới nếu phát hiện đủ hàng
        questions = _group_bubbles_to_questions_improved(rows, template_type)
        use_new_method = True

    # Fallback: Sử dụng phương pháp cũ nếu phương pháp mới không hiệu quả
    if len(questions) < num_questions * 0.3:
        bubbles = _detect_bubbles(binary, template_type)

        if len(bubbles) >= num_questions * 5 * 0.3:
            questions = _group_bubbles_to_questions(bubbles, template_type)
            use_new_method = False

    if len(questions) < num_questions * 0.3:
        return {
            "error": f"Không phát hiện đủ câu hỏi. Tìm thấy: {len(questions)}, cần: {num_questions}. "
                     f"Vui lòng đảm bảo ảnh rõ nét và phiếu được căn chỉnh đúng."
        }

    # Phân tích từng câu hỏi
    student_answers = []
    details = []
    correct_count = 0
    wrong_count = 0
    blank_count = 0

    # Tạo dict để tra cứu câu hỏi theo index (thay vì dùng vị trí trong mảng)
    questions_by_index = {q["index"]: q for q in questions}

    for q_idx in range(num_questions):
        q_num = q_idx + 1  # Số thứ tự câu hỏi (1-based)
        if q_num not in questions_by_index:
            # Không tìm thấy câu hỏi này
            student_answers.append(None)
            blank_count += 1
            correct_answer = answer_key[q_idx] if q_idx < len(answer_key) else None
            details.append({
                "q": q_num,
                "student": None,
                "correct": correct_answer,
                "status": "not_found",
                "fill_ratios": []
            })
            continue

        question = questions_by_index[q_num]
        fill_ratios = []
        mean_vals = []  # Lưu mean value của từng option
        is_filled_list = []  # Lưu trạng thái is_filled của từng option
        max_fill = 0
        selected_option = None

        for opt_idx, bubble in enumerate(question["bubbles"]):
            if opt_idx >= len(option_labels):
                break

            if use_new_method:
                # Phương pháp mới: bubble là dict
                result = _analyze_bubble_fill_improved(gray, bubble)
                if len(result) == 3:
                    is_filled, fill_ratio, mean_val = result
                else:
                    is_filled, fill_ratio = result
                    mean_val = 128.0
                mean_vals.append(mean_val)
                is_filled_list.append(is_filled)
            else:
                # Phương pháp cũ: bubble là tuple với contour
                is_filled, fill_ratio = _analyze_bubble_fill(binary, bubble[4])
                mean_vals.append(128.0)
                is_filled_list.append(is_filled)

            fill_ratios.append(fill_ratio)

            if fill_ratio > max_fill:
                max_fill = fill_ratio
                if is_filled:
                    selected_option = option_labels[opt_idx]

        # Ưu tiên option có is_filled=True và mean thấp nhất (được tô đậm nhất)
        filled_options = [(i, mean_vals[i]) for i in range(len(is_filled_list)) if is_filled_list[i]]
        if filled_options and selected_option is None:
            # Có option được đánh dấu filled nhưng chưa được chọn
            # Chọn option có mean thấp nhất (tối nhất = được tô)
            darkest_filled = min(filled_options, key=lambda x: x[1])
            selected_option = option_labels[darkest_filled[0]]
        elif len(filled_options) == 1:
            # Chỉ có 1 option filled -> chọn option đó
            selected_option = option_labels[filled_options[0][0]]

        # Phát hiện vùng tối bất thường (bóng/rìa ảnh)
        # Nếu nhiều options có mean rất thấp (<50), đây có thể là vùng tối
        dark_region_count = sum(1 for m in mean_vals if m < 50)
        is_dark_region = dark_region_count >= 3

        if is_dark_region:
            # Vùng tối: chọn option có mean CAO nhất (sáng nhất = không bị bóng che)
            # vì các vùng tối là do bóng, không phải do được tô
            max_mean = max(mean_vals)
            min_mean = min(mean_vals)

            # Chỉ chọn nếu có 1 option sáng hơn hẳn (chênh lệch > 50)
            if max_mean - min_mean > 50:
                bright_option_idx = mean_vals.index(max_mean)
                # Kiểm tra option sáng này có được tô không
                if fill_ratios[bright_option_idx] > 0.1:
                    selected_option = option_labels[bright_option_idx]
                else:
                    # Option sáng nhưng không được tô -> có thể là blank hoặc tô option khác
                    # Trong vùng tối, tìm option có score cao nhất trong các option không quá tối
                    valid_options = [(i, fill_ratios[i]) for i in range(len(mean_vals))
                                    if mean_vals[i] > 100 or fill_ratios[i] > 0.3]
                    if valid_options:
                        best_idx = max(valid_options, key=lambda x: x[1])[0]
                        selected_option = option_labels[best_idx]
        else:
            # Vùng bình thường: sử dụng logic cũ với cải tiến

            # Ngưỡng động: nếu max_fill > 0.25 và vượt trội hơn các option khác
            if selected_option is None and max_fill > 0.25:
                # Kiểm tra xem có một option nào vượt trội không
                sorted_ratios = sorted(fill_ratios, reverse=True)
                if len(sorted_ratios) >= 2 and sorted_ratios[0] > sorted_ratios[1] * 1.3:
                    # Option đầu lớn hơn 30% so với option thứ 2
                    selected_option = option_labels[fill_ratios.index(max_fill)]

            # Kiểm tra nếu có nhiều đáp án được chọn
            # Sử dụng ngưỡng động dựa trên max_fill
            if max_fill > 0.5:
                # Nếu có option được tô đậm, các option khác cần đạt ít nhất 60% của max
                filled_threshold = max_fill * 0.6
            else:
                filled_threshold = 0.35

            filled_count = sum(1 for r in fill_ratios if r > filled_threshold)
            if filled_count > 1:
                # Kiểm tra xem có 1 option rõ ràng vượt trội không
                sorted_ratios = sorted(fill_ratios, reverse=True)

                # Tính chênh lệch mean giữa option cao nhất và thấp nhất
                max_score_idx = fill_ratios.index(sorted_ratios[0])
                max_score_mean = mean_vals[max_score_idx]

                # Nếu option có score cao nhất cũng có mean thấp nhất -> đây là bubble được tô
                if max_score_mean == min(mean_vals) or sorted_ratios[0] > sorted_ratios[1] * 1.3:
                    # Có 1 option vượt trội rõ ràng
                    selected_option = option_labels[max_score_idx]
                else:
                    # Kiểm tra thêm: nếu max > 0.5 và second < 0.4, vẫn chọn max
                    if sorted_ratios[0] > 0.5 and sorted_ratios[1] < 0.4:
                        selected_option = option_labels[fill_ratios.index(sorted_ratios[0])]
                    else:
                        # Phân tích thêm bằng mean value
                        # Option được tô sẽ có mean thấp hơn các option không tô
                        mean_diff = max(mean_vals) - min(mean_vals)
                        if mean_diff > 30:
                            # Có sự khác biệt rõ ràng về độ sáng
                            darkest_idx = mean_vals.index(min(mean_vals))
                            selected_option = option_labels[darkest_idx]
                        else:
                            selected_option = "MULTI"  # Đánh dấu chọn nhiều đáp án

        student_answers.append(selected_option)

        # So sánh với đáp án
        correct_answer = answer_key[q_idx] if q_idx < len(answer_key) else None

        if selected_option is None:
            status = "blank"
            blank_count += 1
        elif selected_option == "MULTI":
            status = "invalid"
            wrong_count += 1
        elif correct_answer and selected_option.upper() == correct_answer.upper():
            status = "correct"
            correct_count += 1
        else:
            status = "wrong"
            wrong_count += 1

        details.append({
            "q": q_idx + 1,
            "student": selected_option,
            "correct": correct_answer,
            "status": status,
            "fill_ratios": [round(r, 3) for r in fill_ratios]
        })

    # Tính điểm
    if scoring.get("type") == "tiered":
        # Tính điểm theo phần (tiered scoring) - dùng cho IKSC
        score = scoring.get("base", 0)
        tiers = scoring.get("tiers", [])

        for detail in details:
            q_num = detail["q"]
            status = detail["status"]

            # Tìm tier phù hợp cho câu hỏi này
            tier_points = {"correct": 1, "wrong": 0}  # Default
            for tier in tiers:
                if tier["start"] <= q_num <= tier["end"]:
                    tier_points = tier
                    break

            if status == "correct":
                score += tier_points.get("correct", 1)
            elif status in ["wrong", "invalid"]:
                score += tier_points.get("wrong", 0)
            # blank: không cộng/trừ điểm

    elif scoring.get("type") == "best_of":
        # Tính điểm chỉ lấy N câu tốt nhất - dùng cho IKLC Benjamin-Student
        # Công thức: base + (correct * points) + (wrong * penalty)
        # Chỉ tính count_best câu có điểm cao nhất
        count_best = scoring.get("count_best", 40)
        correct_pts = scoring.get("correct", 1)
        wrong_pts = scoring.get("wrong", -0.5)

        # Tính điểm từng câu
        question_scores = []
        for detail in details:
            status = detail["status"]
            if status == "correct":
                question_scores.append(correct_pts)
            elif status in ["wrong", "invalid"]:
                question_scores.append(wrong_pts)
            else:  # blank
                question_scores.append(0)

        # Sắp xếp giảm dần và lấy N câu tốt nhất
        question_scores.sort(reverse=True)
        best_scores = question_scores[:count_best]

        score = scoring.get("base", 0) + sum(best_scores)

    else:
        # Tính điểm đơn giản (flat scoring) - dùng cho IKLC Pre-Ecolier, Ecolier và các kỳ thi khác
        score = (
            scoring.get("base", 0) +
            correct_count * scoring.get("correct", 1) +
            wrong_count * scoring.get("wrong", 0) +
            blank_count * scoring.get("blank", 0)
        )

    return {
        "answers": student_answers,
        "score": round(score, 2),
        "correct": correct_count,
        "wrong": wrong_count,
        "blank": blank_count,
        "total": num_questions,
        "details": details,
        "student_info": student_info,
        "detection_method": "grid_based" if use_new_method else "contour_based",
        "questions_detected": len(questions)
    }


def _grade_mixed_format_sheet(
    image_bytes: bytes,
    answer_key: List[str],
    template_type: str,
    extract_info: bool = True
):
    """Chấm phiếu trả lời có format hỗn hợp (trắc nghiệm + điền đáp án)

    Dùng cho SEAMO Math: 20 câu trắc nghiệm + 5 câu điền đáp án
    """
    import cv2
    import numpy as np

    template = ANSWER_TEMPLATES.get(template_type, ANSWER_TEMPLATES["IKSC_BENJAMIN"])
    num_questions = template["questions"]
    scoring = template["scoring"]
    mixed_format = template.get("mixed_format", None)

    if not mixed_format:
        # Không có mixed format, dùng hàm chấm thông thường
        return _grade_single_sheet(image_bytes, answer_key, template_type, extract_info)

    mcq_count = mixed_format.get("mcq", 20)
    fill_in_count = mixed_format.get("fill_in", 5)

    # Trích xuất thông tin học sinh
    student_info = {}
    if extract_info:
        try:
            student_info = _extract_student_info_ocr(image_bytes)
        except Exception:
            pass

    # Tiền xử lý ảnh
    result = _preprocess_omr_image(image_bytes)
    if result[0] is None:
        return {"error": "Không thể đọc ảnh"}

    original, gray, binary = result

    # ========== PHẦN 1: Chấm 20 câu trắc nghiệm bằng OMR ==========
    # Kiểm tra nếu là SEAMO, sử dụng fixed grid detection
    is_seamo = "SEAMO" in template_type.upper()

    mcq_questions = []
    if is_seamo:
        # SEAMO có layout cố định, dùng fixed grid
        mcq_questions = _detect_seamo_bubbles_fixed_grid(gray)
    else:
        # Các template khác dùng dynamic detection
        mcq_template_type = template_type
        rows, all_rects = _detect_bubbles_grid_based(gray, binary, mcq_template_type)

        if rows and len(rows) >= 2:
            mcq_questions = _group_bubbles_to_questions_improved(rows, mcq_template_type)
            # Chỉ lấy các câu từ 1 đến mcq_count
            mcq_questions = [q for q in mcq_questions if q["index"] <= mcq_count]

    # Chấm phần trắc nghiệm
    option_labels = ["A", "B", "C", "D", "E"]
    student_answers = []
    details = []
    correct_count = 0
    wrong_count = 0
    blank_count = 0

    questions_by_index = {q["index"]: q for q in mcq_questions}

    for q_idx in range(mcq_count):
        q_num = q_idx + 1
        correct_answer = answer_key[q_idx] if q_idx < len(answer_key) else None

        if q_num not in questions_by_index:
            student_answers.append(None)
            blank_count += 1
            details.append({
                "q": q_num,
                "student": None,
                "correct": correct_answer,
                "status": "not_found",
                "type": "mcq",
                "fill_ratios": []
            })
            continue

        question = questions_by_index[q_num]
        fill_ratios = []
        max_fill = 0
        selected_option = None

        for opt_idx, bubble in enumerate(question["bubbles"]):
            if opt_idx >= len(option_labels):
                break

            result_fill = _analyze_bubble_fill_improved(gray, bubble)
            if len(result_fill) == 3:
                is_filled, fill_ratio, mean_val = result_fill
            else:
                is_filled, fill_ratio = result_fill

            fill_ratios.append(fill_ratio)

            if fill_ratio > max_fill:
                max_fill = fill_ratio
                if is_filled:
                    selected_option = option_labels[opt_idx]

        if selected_option is None and max_fill > 0.25:
            sorted_ratios = sorted(fill_ratios, reverse=True)
            if len(sorted_ratios) >= 2 and sorted_ratios[0] > sorted_ratios[1] * 1.3:
                selected_option = option_labels[fill_ratios.index(max_fill)]

        if selected_option is None:
            student_answers.append(None)
            blank_count += 1
            status = "blank"
        elif correct_answer and selected_option.upper() == correct_answer.upper():
            student_answers.append(selected_option)
            correct_count += 1
            status = "correct"
        else:
            student_answers.append(selected_option)
            wrong_count += 1
            status = "wrong"

        details.append({
            "q": q_num,
            "student": selected_option,
            "correct": correct_answer,
            "status": status,
            "type": "mcq",
            "fill_ratios": [round(r, 3) for r in fill_ratios]
        })

    # ========== PHẦN 2: Chấm 5 câu điền đáp án bằng OCR ==========
    try:
        reader = _get_easyocr_reader()
    except Exception as e:
        # Nếu không load được OCR, đánh dấu các câu fill-in là not_found
        for q_idx in range(mcq_count, num_questions):
            q_num = q_idx + 1
            correct_answer = answer_key[q_idx] if q_idx < len(answer_key) else None
            student_answers.append(None)
            blank_count += 1
            details.append({
                "q": q_num,
                "student": None,
                "correct": correct_answer,
                "status": "not_found",
                "type": "fill_in",
                "confidence": 0.0,
                "error": f"OCR error: {str(e)}"
            })

        # Tính điểm và trả về
        score = (
            scoring.get("base", 0) +
            correct_count * scoring.get("correct", 1) +
            wrong_count * scoring.get("wrong", 0) +
            blank_count * scoring.get("blank", 0)
        )

        return {
            "answers": student_answers,
            "score": round(score, 2),
            "correct": correct_count,
            "wrong": wrong_count,
            "blank": blank_count,
            "total": num_questions,
            "details": details,
            "student_info": student_info,
            "format": "mixed",
            "mcq_count": mcq_count,
            "fill_in_count": fill_in_count
        }

    # Tìm vùng chứa đáp án điền (thường ở phía dưới phiếu)
    # Phát hiện các ô điền đáp án
    height, width = gray.shape[:2]

    # Giả sử phần điền đáp án nằm ở 1/3 dưới của ảnh
    fill_in_region = gray[int(height * 0.6):, :]

    # Sử dụng OCR để đọc toàn bộ vùng
    try:
        ocr_results = reader.readtext(fill_in_region, detail=1, paragraph=False)
    except Exception:
        ocr_results = []

    # Tìm các đáp án số/chữ
    recognized_fill_ins = []
    for (bbox, text, confidence) in ocr_results:
        text = text.strip()
        # Lọc các text có vẻ là đáp án (số hoặc chữ ngắn)
        if text and len(text) <= 10:
            cx = (bbox[0][0] + bbox[2][0]) / 2
            cy = (bbox[0][1] + bbox[2][1]) / 2
            recognized_fill_ins.append({
                'text': text,
                'confidence': confidence,
                'cx': cx,
                'cy': cy + int(height * 0.6)  # Offset lại vị trí
            })

    # Sắp xếp theo vị trí (trái sang phải, trên xuống dưới)
    recognized_fill_ins.sort(key=lambda r: (r['cy'], r['cx']))

    # Gán đáp án cho các câu fill-in
    for i, q_idx in enumerate(range(mcq_count, num_questions)):
        q_num = q_idx + 1
        correct_answer = answer_key[q_idx] if q_idx < len(answer_key) else None

        if i < len(recognized_fill_ins):
            recognized = recognized_fill_ins[i]['text']
            confidence = recognized_fill_ins[i]['confidence']

            # So sánh đáp án (có thể là số hoặc chữ)
            if correct_answer:
                # Chuẩn hóa để so sánh
                student_norm = recognized.upper().strip()
                correct_norm = str(correct_answer).upper().strip()

                if student_norm == correct_norm:
                    status = "correct"
                    correct_count += 1
                else:
                    status = "wrong"
                    wrong_count += 1
            else:
                status = "unknown"

            student_answers.append(recognized)
            details.append({
                "q": q_num,
                "student": recognized,
                "correct": correct_answer,
                "status": status,
                "type": "fill_in",
                "confidence": round(confidence, 3)
            })
        else:
            student_answers.append(None)
            blank_count += 1
            details.append({
                "q": q_num,
                "student": None,
                "correct": correct_answer,
                "status": "not_found",
                "type": "fill_in",
                "confidence": 0.0
            })

    # Tính điểm
    score = (
        scoring.get("base", 0) +
        correct_count * scoring.get("correct", 1) +
        wrong_count * scoring.get("wrong", 0) +
        blank_count * scoring.get("blank", 0)
    )

    return {
        "answers": student_answers,
        "score": round(score, 2),
        "correct": correct_count,
        "wrong": wrong_count,
        "blank": blank_count,
        "total": num_questions,
        "details": details,
        "student_info": student_info,
        "format": "mixed",
        "mcq_count": mcq_count,
        "fill_in_count": fill_in_count
    }


def _extract_answers_from_text(text: str, num_questions: int) -> dict:
    """Trích xuất đáp án từ text (PDF/Word)"""
    # Hỗ trợ các format: "1. A", "1) A", "1: A", "1 A", "Câu 1: A"
    answer_patterns = [
        r'(?:Câu\s*)?(\d+)\s*[.:)]\s*([A-Ea-e])',  # Câu 1: A, 1. A, 1) A
        r'(\d+)\s+([A-Ea-e])\b',  # 1 A
    ]

    found_answers = {}
    for pattern in answer_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            q_num = int(match[0])
            answer = match[1].upper()
            if 1 <= q_num <= num_questions:
                found_answers[q_num] = answer

    return found_answers


def _parse_answer_key_for_template(answer_file_content: bytes, file_ext: str, template_type: str) -> List[str]:
    """Parse đáp án từ file cho một template cụ thể"""
    from collections import defaultdict

    template = ANSWER_TEMPLATES.get(template_type)
    if not template:
        return []

    num_questions = template["questions"]
    answers = []

    if file_ext in ["xlsx", "xls"]:
        # Đọc từ file Excel
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(answer_file_content))
        ws = wb.active

        # Check if header row has level names to select the right column
        ans_col = 1  # Default: column B (index 1)
        if ws.max_column > 2:
            level_keywords = {
                "pre_ecolier": ["preecolier", "pre-ecolier", "pre ecolier", "pre_ecolier"],
                "ecolier": ["ecolier"],
                "benjamin": ["benjamin"],
                "cadet": ["cadet"],
                "junior": ["junior"],
                "student": ["student"],
            }
            headers = [str(ws.cell(1, c).value or '').strip().lower() for c in range(1, ws.max_column + 1)]
            is_ecolier_only = "ecolier" in template_type.lower() and "pre" not in template_type.lower()

            for key, keywords in level_keywords.items():
                if key in template_type.lower():
                    for col_idx, h in enumerate(headers):
                        if is_ecolier_only:
                            if h == "ecolier" or (h.endswith("ecolier") and not h.startswith("pre")):
                                ans_col = col_idx
                                break
                        else:
                            if any(kw in h for kw in keywords):
                                ans_col = col_idx
                                break
                    break

        for row in ws.iter_rows(min_row=2, max_col=ws.max_column):
            if ans_col < len(row) and row[ans_col].value:
                val = str(row[ans_col].value).strip().upper()
                if val and val[0] in "ABCDE":
                    answers.append(val[0])

    elif file_ext == "pdf":
        # Đọc từ file PDF
        import fitz  # PyMuPDF

        pdf_doc = fitz.open(stream=answer_file_content, filetype="pdf")
        pdf_text = ""
        for page in pdf_doc:
            pdf_text += page.get_text() + "\n"

        found_answers = {}

        # Kiểm tra nếu là file IKLC (Linguistic Kangaroo) với format đặc biệt
        is_iklc_format = "LINGUISTIC KANGAROO" in pdf_text.upper() or all(
            level in pdf_text for level in ["Joey", "Wallaby"]
        )

        if is_iklc_format and "IKLC" in template_type.upper():
            # Parse IKLC PDF với format nhiều cột theo vị trí x
            # Cột: Start (25 câu), Story (30 câu), Joey (50 câu), Wallaby (50 câu), Grey K. (50 câu), Red K. (50 câu)
            iklc_levels = [
                ("Start", 25),      # Pre-Ecolier (Lớp 1-2)
                ("Story", 30),      # Ecolier (Lớp 3-4)
                ("Joey", 50),       # Benjamin (Lớp 5-6)
                ("Wallaby", 50),    # Cadet (Lớp 7-8)
                ("Grey", 50),       # Junior (Lớp 9-10)
                ("Red", 50),        # Student (Lớp 11-12)
            ]

            level_map = {
                "IKLC_PRE_ECOLIER": 0,
                "IKLC_ECOLIER": 1,
                "IKLC_BENJAMIN": 2,
                "IKLC_CADET": 3,
                "IKLC_JUNIOR": 4,
                "IKLC_STUDENT": 5,
            }

            target_level_idx = level_map.get(template_type.upper(), -1)

            if target_level_idx >= 0:
                target_level_name, target_num_q = iklc_levels[target_level_idx]

                # Đọc tất cả text blocks với vị trí từ tất cả các trang
                all_blocks = []
                for page in pdf_doc:
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = span["text"].strip()
                                    x0 = span["bbox"][0]
                                    y0 = span["bbox"][1]
                                    if text:
                                        all_blocks.append({"text": text, "x": x0, "y": y0})

                # Tách số câu và đáp án
                numbers = []  # Số câu hỏi (1-50)
                answers_list = []  # Đáp án A-E

                for b in all_blocks:
                    text = b["text"]
                    if text.isdigit() and 1 <= int(text) <= 50:
                        numbers.append({"num": int(text), "x": b["x"], "y": b["y"]})
                    elif len(text) == 1 and text in "ABCDE":
                        answers_list.append({"ans": text, "x": b["x"], "y": b["y"]})
                    elif len(text) >= 1 and text[0] in "ABCDE" and "," in text:
                        # Trường hợp "B, C" -> lấy ký tự đầu
                        answers_list.append({"ans": text[0], "x": b["x"], "y": b["y"]})

                # Tìm vị trí x của số 1 cho mỗi cột (mỗi level bắt đầu từ câu 1)
                ones = [n for n in numbers if n["num"] == 1]
                ones.sort(key=lambda o: o["x"])

                # Có 6 cột (6 số 1), gán level theo thứ tự x
                # ones[0] = Start, ones[1] = Story, ones[2] = Joey, ...
                if len(ones) >= 6:
                    # Xác định x boundaries giữa các cột
                    x_boundaries = []
                    for i in range(len(ones) - 1):
                        mid_x = (ones[i]["x"] + ones[i + 1]["x"]) / 2
                        x_boundaries.append(mid_x)
                    x_boundaries.append(9999)  # Boundary cuối cùng

                    # Hàm xác định cột của một số dựa trên vị trí x
                    def get_column_idx(x):
                        for i, boundary in enumerate(x_boundaries):
                            if x < boundary:
                                return i
                        return len(x_boundaries) - 1

                    # Nhóm số câu theo cột
                    column_numbers = defaultdict(list)
                    for n in numbers:
                        col_idx = get_column_idx(n["x"])
                        column_numbers[col_idx].append(n)

                    # Lấy số câu của cột target
                    target_numbers = column_numbers.get(target_level_idx, [])

                    # Với mỗi số câu, tìm đáp án gần nhất bên phải
                    for num_block in target_numbers:
                        q_num = num_block["num"]
                        q_x = num_block["x"]
                        q_y = num_block["y"]

                        if q_num > target_num_q:
                            continue

                        # Tìm đáp án gần nhất: cùng y (tolerance 3px trước, sau đó 8px) và x lớn hơn số câu
                        best_answer = None
                        best_dist = 9999
                        best_y_diff = 9999

                        for ans_block in answers_list:
                            ans_x = ans_block["x"]
                            ans_y = ans_block["y"]

                            # Đáp án ở bên phải số câu và cùng hàng
                            y_diff = abs(ans_y - q_y)
                            if ans_x > q_x and y_diff < 8:
                                dist = ans_x - q_x
                                if dist < 50:  # Không quá xa
                                    # Ưu tiên đáp án cùng y hơn (y_diff nhỏ hơn)
                                    # Nếu y_diff gần bằng nhau (< 3px), chọn x gần nhất
                                    if y_diff < best_y_diff - 3 or (abs(y_diff - best_y_diff) <= 3 and dist < best_dist):
                                        best_dist = dist
                                        best_y_diff = y_diff
                                        best_answer = ans_block["ans"]

                        if best_answer and q_num not in found_answers:
                            found_answers[q_num] = best_answer

        pdf_doc.close()

        # Fallback: parse đơn giản
        if not found_answers:
            found_answers = _extract_answers_from_text(pdf_text, num_questions)

        for i in range(1, num_questions + 1):
            answers.append(found_answers.get(i, ""))

    elif file_ext in ["docx", "doc"]:
        # Đọc từ file Word
        from docx import Document
        doc = Document(io.BytesIO(answer_file_content))
        found_answers = {}

        level_keywords = {
            "pre_ecolier": ["preecolier", "pre-ecolier", "pre ecolier", "pre_ecolier"],
            "ecolier": ["ecolier"],
            "benjamin": ["benjamin"],
            "cadet": ["cadet"],
            "junior": ["junior"],
            "student": ["student"],
        }

        for table in doc.tables:
            if len(table.rows) > 1 and len(table.columns) >= 2:
                header = [cell.text.strip().lower() for cell in table.rows[0].cells]

                level_col = -1
                search_keywords = []
                is_ecolier_only = False

                for key, keywords in level_keywords.items():
                    if key in template_type.lower():
                        search_keywords = keywords
                        if key == "ecolier" and "pre" not in template_type.lower():
                            is_ecolier_only = True
                        break

                for col_idx, col_header in enumerate(header):
                    if is_ecolier_only:
                        if col_header == "ecolier" or (col_header.endswith("ecolier") and not col_header.startswith("pre")):
                            level_col = col_idx
                            break
                    else:
                        for keyword in search_keywords:
                            if keyword in col_header:
                                level_col = col_idx
                                break
                    if level_col >= 0:
                        break

                if level_col >= 0:
                    # Detect paired-column format: each level has 2 cols (number, answer)
                    # Check if header has duplicate names (merged cells)
                    is_paired = (level_col + 1 < len(header)
                                 and header[level_col] == header[level_col + 1])
                    if is_paired:
                        # Paired format: level_col = numbers, level_col+1 = answers
                        num_col = level_col
                        ans_col = level_col + 1
                    else:
                        # Standard format: col 0 = numbers, level_col = answers
                        num_col = 0
                        ans_col = level_col

                    for row in table.rows[1:]:
                        try:
                            q_num = int(row.cells[num_col].text.strip())
                            answer = row.cells[ans_col].text.strip().upper()
                            if answer and answer in "ABCDE":
                                found_answers[q_num] = answer
                        except (ValueError, IndexError):
                            continue

                if not found_answers and len(table.columns) == 2:
                    for row in table.rows[1:]:
                        try:
                            q_num = int(row.cells[0].text.strip())
                            answer = row.cells[1].text.strip().upper()
                            if answer and answer in "ABCDE":
                                found_answers[q_num] = answer
                        except (ValueError, IndexError):
                            continue

        if not found_answers:
            doc_text = ""
            for para in doc.paragraphs:
                doc_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        doc_text += cell.text + " "
                    doc_text += "\n"
            found_answers = _extract_answers_from_text(doc_text, num_questions)

        for i in range(1, num_questions + 1):
            answers.append(found_answers.get(i, ""))

    return answers
