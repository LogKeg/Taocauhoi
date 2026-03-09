"""Answer key parsing and text extraction for OMR grading.

Parse answer keys from Excel, PDF, and Word files for various exam templates.
Extract answers from plain text as fallback.
"""

import io
import re
from collections import defaultdict
from typing import List

from app.core import ANSWER_TEMPLATES


def _extract_answers_from_text(text: str, num_questions: int) -> dict:
    """Trích xuất đáp án từ text (PDF/Word)"""
    # Hỗ trợ các format: "1. A", "1) A", "1: A", "1 A", "Câu 1: A"
    answer_patterns = [
        r'(?:Câu\s*)?(\d+)\s*[.:)]\s*([A-Ea-e])',  # Câu 1: A, 1. A, 1) A
        r'(\d+)\s+([A-Ea-e])\b',  # 1 A
    ]

    found_answers = {}
    for pattern in answer_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            q_num = int(match[0])
            answer = match[1].upper()
            if 1 <= q_num <= num_questions:
                found_answers[q_num] = answer

    return found_answers


def _parse_answer_key_for_template(answer_file_content: bytes, file_ext: str, template_type: str) -> List[str]:
    """Parse đáp án từ file cho một template cụ thể"""

    template = ANSWER_TEMPLATES.get(template_type)
    if not template:
        return []

    num_questions = template["questions"]
    answers = []

    if file_ext in ["xlsx", "xls"]:
        answers = _parse_excel_answer_key(answer_file_content, template_type, num_questions)

    elif file_ext == "pdf":
        answers = _parse_pdf_answer_key(answer_file_content, template_type, num_questions)

    elif file_ext in ["docx", "doc"]:
        answers = _parse_word_answer_key(answer_file_content, template_type, num_questions)

    return answers


def _parse_excel_answer_key(answer_file_content: bytes, template_type: str, num_questions: int) -> List[str]:
    """Parse answer key from Excel file."""
    import openpyxl

    answers = []
    wb = openpyxl.load_workbook(io.BytesIO(answer_file_content))
    ws = wb.active

    # Check if header row has level names to select the right column
    ans_col = 1  # Default: column B (index 1)
    if ws.max_column > 2:
        level_keywords = {
            "pre_ecolier": ["preecolier", "pre-ecolier", "pre ecolier", "pre_ecolier"],
            "ecolier": ["ecolier"],
            "benjamin": ["benjamin"],
            "cadet": ["cadet"],
            "junior": ["junior"],
            "student": ["student"],
        }
        headers = [str(ws.cell(1, c).value or '').strip().lower() for c in range(1, ws.max_column + 1)]
        is_ecolier_only = "ecolier" in template_type.lower() and "pre" not in template_type.lower()

        for key, keywords in level_keywords.items():
            if key in template_type.lower():
                for col_idx, h in enumerate(headers):
                    if is_ecolier_only:
                        if h == "ecolier" or (h.endswith("ecolier") and not h.startswith("pre")):
                            ans_col = col_idx
                            break
                    else:
                        if any(kw in h for kw in keywords):
                            ans_col = col_idx
                            break
                break

    for row in ws.iter_rows(min_row=2, max_col=ws.max_column):
        if ans_col < len(row) and row[ans_col].value:
            val = str(row[ans_col].value).strip().upper()
            if val and val[0] in "ABCDE":
                answers.append(val[0])

    return answers


def _parse_pdf_answer_key(answer_file_content: bytes, template_type: str, num_questions: int) -> List[str]:
    """Parse answer key from PDF file."""
    import fitz  # PyMuPDF

    answers = []
    pdf_doc = fitz.open(stream=answer_file_content, filetype="pdf")
    pdf_text = ""
    for page in pdf_doc:
        pdf_text += page.get_text() + "\n"

    found_answers = {}

    # Kiểm tra nếu là file IKLC (Linguistic Kangaroo) với format đặc biệt
    is_iklc_format = "LINGUISTIC KANGAROO" in pdf_text.upper() or all(
        level in pdf_text for level in ["Joey", "Wallaby"]
    )

    if is_iklc_format and "IKLC" in template_type.upper():
        found_answers = _parse_iklc_pdf_format(pdf_doc, template_type, num_questions)

    pdf_doc.close()

    # Fallback: parse đơn giản
    if not found_answers:
        found_answers = _extract_answers_from_text(pdf_text, num_questions)

    for i in range(1, num_questions + 1):
        answers.append(found_answers.get(i, ""))

    return answers


def _parse_iklc_pdf_format(pdf_doc, template_type: str, num_questions: int) -> dict:
    """Parse IKLC PDF với format nhiều cột theo vị trí x.

    Cột: Start (25 câu), Story (30 câu), Joey (50 câu), Wallaby (50 câu), Grey K. (50 câu), Red K. (50 câu)
    """
    iklc_levels = [
        ("Start", 25),      # Pre-Ecolier (Lớp 1-2)
        ("Story", 30),      # Ecolier (Lớp 3-4)
        ("Joey", 50),       # Benjamin (Lớp 5-6)
        ("Wallaby", 50),    # Cadet (Lớp 7-8)
        ("Grey", 50),       # Junior (Lớp 9-10)
        ("Red", 50),        # Student (Lớp 11-12)
    ]

    level_map = {
        "IKLC_PRE_ECOLIER": 0,
        "IKLC_ECOLIER": 1,
        "IKLC_BENJAMIN": 2,
        "IKLC_CADET": 3,
        "IKLC_JUNIOR": 4,
        "IKLC_STUDENT": 5,
    }

    found_answers = {}
    target_level_idx = level_map.get(template_type.upper(), -1)

    if target_level_idx >= 0:
        target_level_name, target_num_q = iklc_levels[target_level_idx]

        # Đọc tất cả text blocks với vị trí từ tất cả các trang
        all_blocks = []
        for page in pdf_doc:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip()
                            x0 = span["bbox"][0]
                            y0 = span["bbox"][1]
                            if text:
                                all_blocks.append({"text": text, "x": x0, "y": y0})

        # Tách số câu và đáp án
        numbers = []  # Số câu hỏi (1-50)
        answers_list = []  # Đáp án A-E

        for b in all_blocks:
            text = b["text"]
            if text.isdigit() and 1 <= int(text) <= 50:
                numbers.append({"num": int(text), "x": b["x"], "y": b["y"]})
            elif len(text) == 1 and text in "ABCDE":
                answers_list.append({"ans": text, "x": b["x"], "y": b["y"]})
            elif len(text) >= 1 and text[0] in "ABCDE" and "," in text:
                # Trường hợp "B, C" -> lấy ký tự đầu
                answers_list.append({"ans": text[0], "x": b["x"], "y": b["y"]})

        # Tìm vị trí x của số 1 cho mỗi cột (mỗi level bắt đầu từ câu 1)
        ones = [n for n in numbers if n["num"] == 1]
        ones.sort(key=lambda o: o["x"])

        # Có 6 cột (6 số 1), gán level theo thứ tự x
        # ones[0] = Start, ones[1] = Story, ones[2] = Joey, ...
        if len(ones) >= 6:
            # Xác định x boundaries giữa các cột
            x_boundaries = []
            for i in range(len(ones) - 1):
                mid_x = (ones[i]["x"] + ones[i + 1]["x"]) / 2
                x_boundaries.append(mid_x)
            x_boundaries.append(9999)  # Boundary cuối cùng

            # Hàm xác định cột của một số dựa trên vị trí x
            def get_column_idx(x):
                for i, boundary in enumerate(x_boundaries):
                    if x < boundary:
                        return i
                return len(x_boundaries) - 1

            # Nhóm số câu theo cột
            column_numbers = defaultdict(list)
            for n in numbers:
                col_idx = get_column_idx(n["x"])
                column_numbers[col_idx].append(n)

            # Lấy số câu của cột target
            target_numbers = column_numbers.get(target_level_idx, [])

            # Với mỗi số câu, tìm đáp án gần nhất bên phải
            for num_block in target_numbers:
                q_num = num_block["num"]
                q_x = num_block["x"]
                q_y = num_block["y"]

                if q_num > target_num_q:
                    continue

                # Tìm đáp án gần nhất: cùng y (tolerance 3px trước, sau đó 8px) và x lớn hơn số câu
                best_answer = None
                best_dist = 9999
                best_y_diff = 9999

                for ans_block in answers_list:
                    ans_x = ans_block["x"]
                    ans_y = ans_block["y"]

                    # Đáp án ở bên phải số câu và cùng hàng
                    y_diff = abs(ans_y - q_y)
                    if ans_x > q_x and y_diff < 8:
                        dist = ans_x - q_x
                        if dist < 50:  # Không quá xa
                            # Ưu tiên đáp án cùng y hơn (y_diff nhỏ hơn)
                            # Nếu y_diff gần bằng nhau (< 3px), chọn x gần nhất
                            if y_diff < best_y_diff - 3 or (abs(y_diff - best_y_diff) <= 3 and dist < best_dist):
                                best_dist = dist
                                best_y_diff = y_diff
                                best_answer = ans_block["ans"]

                if best_answer and q_num not in found_answers:
                    found_answers[q_num] = best_answer

    return found_answers


def _parse_word_answer_key(answer_file_content: bytes, template_type: str, num_questions: int) -> List[str]:
    """Parse answer key from Word file."""
    from docx import Document

    answers = []
    doc = Document(io.BytesIO(answer_file_content))
    found_answers = {}

    level_keywords = {
        "pre_ecolier": ["preecolier", "pre-ecolier", "pre ecolier", "pre_ecolier"],
        "ecolier": ["ecolier"],
        "benjamin": ["benjamin"],
        "cadet": ["cadet"],
        "junior": ["junior"],
        "student": ["student"],
    }

    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) >= 2:
            header = [cell.text.strip().lower() for cell in table.rows[0].cells]

            level_col = -1
            search_keywords = []
            is_ecolier_only = False

            for key, keywords in level_keywords.items():
                if key in template_type.lower():
                    search_keywords = keywords
                    if key == "ecolier" and "pre" not in template_type.lower():
                        is_ecolier_only = True
                    break

            for col_idx, col_header in enumerate(header):
                if is_ecolier_only:
                    if col_header == "ecolier" or (col_header.endswith("ecolier") and not col_header.startswith("pre")):
                        level_col = col_idx
                        break
                else:
                    for keyword in search_keywords:
                        if keyword in col_header:
                            level_col = col_idx
                            break
                if level_col >= 0:
                    break

            if level_col >= 0:
                # Detect paired-column format: each level has 2 cols (number, answer)
                # Check if header has duplicate names (merged cells)
                is_paired = (level_col + 1 < len(header)
                             and header[level_col] == header[level_col + 1])
                if is_paired:
                    # Paired format: level_col = numbers, level_col+1 = answers
                    num_col = level_col
                    ans_col = level_col + 1
                else:
                    # Standard format: col 0 = numbers, level_col = answers
                    num_col = 0
                    ans_col = level_col

                for row in table.rows[1:]:
                    try:
                        q_num = int(row.cells[num_col].text.strip())
                        answer = row.cells[ans_col].text.strip().upper()
                        if answer and answer in "ABCDE":
                            found_answers[q_num] = answer
                    except (ValueError, IndexError):
                        continue

            if not found_answers and len(table.columns) == 2:
                for row in table.rows[1:]:
                    try:
                        q_num = int(row.cells[0].text.strip())
                        answer = row.cells[1].text.strip().upper()
                        if answer and answer in "ABCDE":
                            found_answers[q_num] = answer
                    except (ValueError, IndexError):
                        continue

    if not found_answers:
        doc_text = ""
        for para in doc.paragraphs:
            doc_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    doc_text += cell.text + " "
                doc_text += "\n"
        found_answers = _extract_answers_from_text(doc_text, num_questions)

    for i in range(1, num_questions + 1):
        answers.append(found_answers.get(i, ""))

    return answers
