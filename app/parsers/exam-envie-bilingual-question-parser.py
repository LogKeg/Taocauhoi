"""
Parser for EN-VIE bilingual English exam questions and bilingual science deduplication.

Handles formats used in Kangaroo, IKSC, and similar bilingual exam documents:
- Fill-blank questions with tab-separated options
- Questions ending with ? followed by paragraph options
- Matching questions with A) B) C) D) E) options
- Questions in paragraphs with options in separate tables (1x4 or 2x2 grid)
"""

import re
from typing import List

from docx import Document


def _parse_envie_questions(doc: Document) -> List[dict]:
    """
    Parse EN-VIE bilingual English exam questions.
    These files have different formats:
    1. Fill-blank questions followed by tab-separated options: "text\tB) text\tC) text"
    2. Questions ending with ? followed by paragraph options (no A/B/C markers)
    3. Matching questions with A) B) C) D) E) options
    4. Questions in paragraphs with options in separate tables (1x4 or 2x2 grid)
    """
    from docx.oxml.ns import qn as docx_qn

    questions = []

    # Extract document elements in order (paragraphs and tables interleaved)
    # This is important for formats where questions are in paragraphs and options in tables
    def get_document_elements():
        """Get paragraphs and tables in document order."""
        elements = []
        body = doc._element.body

        # Track numbering counters per numId and ilvl
        # Format: {(numId, ilvl): current_count}
        numbering_counters = {}

        for child in body:
            tag = child.tag.split('}')[-1]

            if tag == 'p':  # Paragraph
                # Build text with line breaks preserved (w:br → \n)
                text_parts = []
                for run_elem in child.findall('.//' + docx_qn('w:r')):
                    for sub in run_elem:
                        sub_tag = sub.tag.split('}')[-1]
                        if sub_tag == 't':
                            text_parts.append(sub.text or '')
                        elif sub_tag == 'br':
                            text_parts.append('\n')
                text = ''.join(text_parts).strip()
                if not text:
                    # Fallback to w:t only
                    t_elements = child.findall('.//' + docx_qn('w:t'))
                    text = ''.join([t.text or '' for t in t_elements]).strip()
                if text:
                    # Check for highlighted text in paragraph (for inline options)
                    # Find which option (by position) has yellow highlight
                    highlighted_opt_idx = 0
                    if re.search(r'[A-E]\)', text):  # Has inline options
                        # Parse runs to find highlighted option
                        runs = child.findall('.//' + docx_qn('w:r'))
                        current_opt_letter = None
                        for run in runs:
                            # Get run text
                            run_texts = run.findall('.//' + docx_qn('w:t'))
                            run_text = ''.join([t.text or '' for t in run_texts])

                            # Check for option marker
                            opt_match = re.search(r'([A-E])\)', run_text)
                            if opt_match:
                                current_opt_letter = opt_match.group(1)

                            # Check for highlight in this run
                            shd_elements = run.findall('.//' + docx_qn('w:shd'))
                            hl_elements = run.findall('.//' + docx_qn('w:highlight'))
                            has_yellow = False
                            for shd in shd_elements:
                                fill = shd.get(docx_qn('w:fill'))
                                if fill and fill.upper() == 'FFFF00':
                                    has_yellow = True
                            for hl in hl_elements:
                                val = hl.get(docx_qn('w:val'))
                                if val and val.lower() == 'yellow':
                                    has_yellow = True

                            # If highlighted and we know which option, record it
                            if has_yellow and current_opt_letter:
                                highlighted_opt_idx = ord(current_opt_letter) - ord('A') + 1

                    # Check for Word numbering (List Paragraph style)
                    # This handles files where question numbers are in Word's numbering system
                    num_level = None  # 0 = question level, 1 = option level
                    num_value = None  # The actual number (1, 2, 3, ...)

                    pPr = child.find(docx_qn('w:pPr'))
                    if pPr is not None:
                        numPr = pPr.find(docx_qn('w:numPr'))
                        if numPr is not None:
                            numId_elem = numPr.find(docx_qn('w:numId'))
                            ilvl_elem = numPr.find(docx_qn('w:ilvl'))

                            if numId_elem is not None and ilvl_elem is not None:
                                numId = numId_elem.get(docx_qn('w:val'))
                                ilvl = int(ilvl_elem.get(docx_qn('w:val')) or '0')
                                num_level = ilvl

                                # Track and increment counter for this numbering
                                key = (numId, ilvl)
                                if key not in numbering_counters:
                                    numbering_counters[key] = 0
                                numbering_counters[key] += 1
                                num_value = numbering_counters[key]

                                # Reset sub-level counters when parent level increments
                                # e.g., when question (ilvl=0) increments, reset option counter (ilvl=1)
                                if ilvl == 0:
                                    sub_key = (numId, 1)
                                    numbering_counters[sub_key] = 0

                    # Split paragraphs containing \n with option patterns
                    if '\n' in text and re.search(r'\n[A-E]\.?\s*\S', text):
                        for sub_line in text.split('\n'):
                            sub_line = sub_line.strip()
                            if sub_line:
                                elements.append({
                                    'type': 'paragraph',
                                    'text': sub_line,
                                    'highlighted_option': highlighted_opt_idx,
                                    'num_level': num_level,
                                    'num_value': num_value
                                })
                    else:
                        # Remove any remaining \n from text
                        text = text.replace('\n', ' ').strip()
                        elements.append({
                            'type': 'paragraph',
                            'text': text,
                            'highlighted_option': highlighted_opt_idx,
                            'num_level': num_level,  # None = no numbering, 0 = question, 1 = option
                            'num_value': num_value   # The number (1, 2, 3, ...)
                        })

            elif tag == 'tbl':  # Table
                # Extract options from table cells
                rows = child.findall('.//' + docx_qn('w:tr'))
                options = []
                highlighted_idx = 0  # Track which option is highlighted (for single-question tables)
                highlighted_map = {}  # Track highlights per Cloze question number
                opt_counter = 0
                current_cloze_num = None
                opt_idx_in_cloze = 0

                def split_cell_options(cell_text: str) -> List[str]:
                    """Split cell text that may contain multiple tab-separated options.
                    E.g., 'Selfish.\tB) Emancipatory.' -> ['Selfish.', 'Emancipatory.']
                    """
                    # Check for tab-separated options with markers
                    if '\t' in cell_text and re.search(r'\t[A-E]\s*[.)]', cell_text):
                        parts = []
                        # Split by tab and process each part
                        raw_parts = cell_text.split('\t')
                        for part in raw_parts:
                            part = part.strip()
                            if not part:
                                continue
                            # Remove option marker (A), B., etc.)
                            cleaned = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', part)
                            if cleaned:
                                parts.append(cleaned)
                        return parts if parts else [cell_text]
                    return [cell_text]

                # Collect rows, then check if table contains embedded questions
                # (e.g., options for Q38 in row 0, then Q39 + options in rows 1-2)
                table_rows_data = []  # List of lists: each row -> list of (cell_text, is_highlighted)
                for tr in rows:
                    row_cells = []
                    for tc in tr.findall('.//' + docx_qn('w:tc')):
                        t_elements = tc.findall('.//' + docx_qn('w:t'))
                        cell_text = ''.join([t.text or '' for t in t_elements]).strip()
                        cell_is_highlighted = False
                        if cell_text:
                            shd_elements = tc.findall('.//' + docx_qn('w:shd'))
                            for shd in shd_elements:
                                fill = shd.get(docx_qn('w:fill'))
                                if fill and fill.upper() == 'FFFF00':
                                    cell_is_highlighted = True
                                    break
                        row_cells.append((cell_text, cell_is_highlighted))
                    table_rows_data.append(row_cells)

                # Detect embedded questions: a cell contains "Question text" + "Option A text"
                # concatenated (e.g., "Who discovered gravity?Albert Einstein.")
                # Heuristic: first cell of a row has text ending with '?' or '.' followed by more text
                # and other cells in that row + next rows look like B), C), D), E) options
                embedded_questions = []  # List of (question_text, options, row_start_idx)
                embedded_row_indices = set()

                for ri, row_cells in enumerate(table_rows_data):
                    if ri in embedded_row_indices:
                        continue
                    if not row_cells or not row_cells[0][0]:
                        continue
                    first_cell = row_cells[0][0]
                    # Check if first cell contains question+optionA pattern
                    # e.g., "Who discovered gravity?Albert Einstein."
                    q_opt_match = re.match(r'^(.+\?)\s*([A-Z].+)$', first_cell, re.DOTALL)
                    if not q_opt_match:
                        continue
                    q_text = q_opt_match.group(1).strip()
                    opt_a_text = q_opt_match.group(2).strip()
                    # Verify other cells have B), C), etc.
                    other_cells_text = [c[0] for c in row_cells[1:] if c[0]]
                    has_bc_markers = any(re.match(r'^[B-E]\)', ct) for ct in other_cells_text)
                    if not has_bc_markers and len(other_cells_text) < 1:
                        continue
                    # This row has an embedded question - collect its options
                    emb_opts = [opt_a_text]
                    emb_highlighted = 0
                    if row_cells[0][1]:  # first cell highlighted
                        emb_highlighted = 1
                    opt_num = 1
                    for ci in range(1, len(row_cells)):
                        ct, ch = row_cells[ci]
                        if ct:
                            cell_options = split_cell_options(ct)
                            for co in cell_options:
                                clean_co = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', co).strip()
                                if clean_co:
                                    opt_num += 1
                                    emb_opts.append(clean_co)
                                    if ch:
                                        emb_highlighted = opt_num
                    embedded_row_indices.add(ri)
                    # Check next rows for continuation (D), E))
                    for ri2 in range(ri + 1, len(table_rows_data)):
                        next_row_cells = table_rows_data[ri2]
                        next_texts = [c[0] for c in next_row_cells if c[0]]
                        if next_texts and all(re.match(r'^[A-E]\)', t) or not t for t in next_texts):
                            for ct, ch in next_row_cells:
                                if ct:
                                    cell_options = split_cell_options(ct)
                                    for co in cell_options:
                                        clean_co = re.sub(r'^[A-Ea-e]\s*[.)]\s*', '', co).strip()
                                        if clean_co:
                                            opt_num += 1
                                            emb_opts.append(clean_co)
                                            if ch:
                                                emb_highlighted = opt_num
                            embedded_row_indices.add(ri2)
                        else:
                            break
                    if len(emb_opts) >= 2:
                        embedded_questions.append((q_text, emb_opts, emb_highlighted, ri))

                # Process non-embedded rows as regular options
                # Track embedded questions found in cells (e.g., "A) Table tennis.34. Pick out the odd one.")
                row_groups = []  # List of (options_list, highlighted_idx, highlighted_map, embedded_q_text)
                current_group_options = []
                current_group_highlighted = 0
                current_group_hl_map = {}
                current_group_opt_counter = 0

                for ri, row_cells in enumerate(table_rows_data):
                    if ri in embedded_row_indices:
                        continue
                    embedded_q_in_row = None
                    for cell_text, cell_is_highlighted in row_cells:
                        if cell_text:
                            # Check for embedded question number in cell
                            # Pattern: "A) Table tennis.34. Pick out the odd one."
                            # Split into option part and embedded question part
                            emb_q_match = re.search(r'(?<=\.)\s*(\d+\.\s+.+)$', cell_text)
                            if emb_q_match and re.match(r'^[A-E]\)', cell_text):
                                # Found embedded question - split the cell
                                option_part = cell_text[:emb_q_match.start()].strip()
                                embedded_q_in_row = emb_q_match.group(1).strip()
                                cell_text = option_part

                            cell_options = split_cell_options(cell_text)
                            for cell_opt in cell_options:
                                current_group_opt_counter += 1
                                current_group_options.append(cell_opt)

                                cloze_match = re.match(r'^(\d+)\.\s*[A-E]\)', cell_opt)
                                if cloze_match:
                                    current_cloze_num = int(cloze_match.group(1))
                                    opt_idx_in_cloze = 1
                                elif re.match(r'^[B-E]\)', cell_opt) and current_cloze_num:
                                    opt_idx_in_cloze += 1

                                if cell_is_highlighted:
                                    if current_cloze_num:
                                        current_group_hl_map[current_cloze_num] = opt_idx_in_cloze
                                    else:
                                        current_group_highlighted = current_group_opt_counter

                    # If this row had an embedded question, save current group and start new one
                    if embedded_q_in_row:
                        row_groups.append((
                            current_group_options[:],
                            current_group_highlighted,
                            current_group_hl_map.copy(),
                            embedded_q_in_row
                        ))
                        current_group_options = []
                        current_group_highlighted = 0
                        current_group_hl_map = {}
                        current_group_opt_counter = 0

                # Save last group
                if current_group_options:
                    row_groups.append((current_group_options, current_group_highlighted, current_group_hl_map, None))

                # Count total cells across all rows (including empty ones)
                total_cell_count = sum(len(row) for row in table_rows_data)

                # Emit table elements with embedded question paragraphs between them
                if not row_groups:
                    # No row groups — emit table with whatever options we have
                    # Even if options is empty, emit if there are cells (image-based options)
                    if options or total_cell_count >= 3:
                        elements.append({
                            'type': 'table',
                            'options': options,
                            'highlighted': highlighted_idx,
                            'cloze_highlights': highlighted_map,
                            'cell_count': total_cell_count
                        })
                else:
                    for grp_opts, grp_hl, grp_hl_map, grp_emb_q in row_groups:
                        if grp_opts:
                            elements.append({
                                'type': 'table',
                                'options': grp_opts,
                                'highlighted': grp_hl,
                                'cloze_highlights': grp_hl_map,
                                'cell_count': total_cell_count
                            })
                        if grp_emb_q:
                            # Remove leading question number for clean question text
                            q_text_clean = re.sub(r'^\d+\.\s*', '', grp_emb_q)
                            elements.append({
                                'type': 'paragraph',
                                'text': q_text_clean,
                                'highlighted_option': 0,
                                'num_level': None,
                                'num_value': None
                            })

                # Append embedded questions as separate paragraph+table pairs
                for q_text, emb_opts, emb_hl, _ in embedded_questions:
                    elements.append({
                        'type': 'paragraph',
                        'text': q_text,
                        'highlighted_option': 0,
                        'num_level': None,
                        'num_value': None
                    })
                    elements.append({
                        'type': 'table',
                        'options': emb_opts,
                        'highlighted': emb_hl,
                        'cloze_highlights': {}
                    })

        return elements

    doc_elements = get_document_elements()

    # Merge bilingual EN+VN paragraph pairs for science exams (IKSC format)
    # Pattern: EN paragraph (numbered, no VN chars) followed by VN paragraph (has VN chars)
    # followed by option lines (A./B./C.) → merge EN+VN into one bilingual paragraph
    def _has_vietnamese_chars(text):
        return any(ord(c) > 127 for c in text[:200])

    # Check if this looks like a bilingual science document
    # Detect by bilingual options (e.g., "A. Refraction / Khúc xạ") or EN+VN question pairs
    bilingual_option_count = sum(
        1 for elem in doc_elements[:60]
        if elem['type'] == 'paragraph' and re.match(r'^[A-E][.)]\s*.+\s*/\s*.+', elem['text'])
    )
    bilingual_pair_count = 0
    for idx_e in range(len(doc_elements) - 1):
        e1 = doc_elements[idx_e]
        e2 = doc_elements[idx_e + 1]
        if (e1['type'] == 'paragraph' and e2['type'] == 'paragraph'
                and re.match(r'^\d+\.\s*', e1['text'])
                and not _has_vietnamese_chars(e1['text'])
                and _has_vietnamese_chars(e2['text'])
                and not re.match(r'^[A-E][.)]\s*', e2['text'])):
            bilingual_pair_count += 1
            if bilingual_pair_count >= 3:
                break
    is_bilingual_doc = bilingual_option_count >= 3 or bilingual_pair_count >= 3
    if is_bilingual_doc:
        merged_elements = []
        i = 0
        while i < len(doc_elements):
            elem = doc_elements[i]
            if (elem['type'] == 'paragraph'
                    and not _has_vietnamese_chars(elem['text'])
                    and re.match(r'^\d+\.\s*', elem['text'])
                    and i + 1 < len(doc_elements)
                    and doc_elements[i + 1]['type'] == 'paragraph'
                    and _has_vietnamese_chars(doc_elements[i + 1]['text'])):
                # Merge EN + VN into one bilingual paragraph
                en_text = re.sub(r'^\d+\.\s*', '', elem['text']).strip()
                vn_text = doc_elements[i + 1]['text']
                merged_elem = dict(doc_elements[i + 1])
                merged_elem['text'] = en_text + '\n' + vn_text
                merged_elements.append(merged_elem)
                i += 2
            else:
                merged_elements.append(elem)
                i += 1
        doc_elements = merged_elements

    # Build a map of option paragraph text -> highlighted option index
    # This helps when processing inline options in paragraphs
    para_highlight_map = {}
    # Build list of (doc_elem_idx, normalized_text) for fuzzy matching
    doc_elem_texts = []
    for idx, elem in enumerate(doc_elements):
        if elem['type'] == 'paragraph':
            text = elem['text']
            if elem.get('highlighted_option', 0) > 0:
                para_highlight_map[text] = elem['highlighted_option']
            # Store normalized text for matching
            normalized_text = text.replace('\n', '').replace('\r', '')
            doc_elem_texts.append((idx, normalized_text))

    # Also keep traditional paragraph extraction for backward compatibility
    # Split paragraphs containing \n with option patterns (e.g., "A.Moss\nB. Wheat\nC. Corn")
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if '\n' in text and re.search(r'\n[A-E]\.?\s*\S', text):
            for sub_line in text.split('\n'):
                sub_line = sub_line.strip()
                if sub_line:
                    paragraphs.append(sub_line)
        else:
            paragraphs.append(text)

    # Merge bilingual EN+VN paragraph pairs in paragraphs list too
    if is_bilingual_doc:
        merged_paras = []
        i = 0
        while i < len(paragraphs):
            text = paragraphs[i]
            if (not _has_vietnamese_chars(text)
                    and re.match(r'^\d+\.\s*', text)
                    and i + 1 < len(paragraphs)
                    and _has_vietnamese_chars(paragraphs[i + 1])
                    and not re.match(r'^[A-E][.)]\s*', paragraphs[i + 1])):
                en_text = re.sub(r'^\d+\.\s*', '', text).strip()
                vn_text = paragraphs[i + 1]
                merged_paras.append(en_text + '\n' + vn_text)
                i += 2
            else:
                merged_paras.append(text)
                i += 1
        paragraphs = merged_paras

    # Create a mapping from paragraphs index to doc_elements index
    # This ensures consistent ordering when sorting by _doc_pos
    # Use fuzzy matching: find doc_element that contains the paragraph text
    para_idx_to_doc_idx = {}
    for pi, p_text in enumerate(paragraphs):
        if not p_text:
            para_idx_to_doc_idx[pi] = 10000 + pi
            continue
        # Normalize paragraph text for matching
        normalized_p = p_text.replace('\n', '').replace('\r', '')
        # First try exact match
        found = False
        for doc_idx, doc_text in doc_elem_texts:
            if normalized_p == doc_text:
                para_idx_to_doc_idx[pi] = doc_idx
                found = True
                break
        if not found:
            # Try fuzzy match: check if doc_text contains or ends with normalized_p
            for doc_idx, doc_text in doc_elem_texts:
                if doc_text.endswith(normalized_p) or normalized_p in doc_text:
                    para_idx_to_doc_idx[pi] = doc_idx
                    found = True
                    break
        if not found:
            # Fallback: use paragraph index (paragraphs are roughly in order)
            para_idx_to_doc_idx[pi] = pi

    def extract_options_envie(line: str) -> List[str]:
        """Extract options from EN-VIE format line."""
        opts = []
        # Format: "optA\tB) optB\tC) optC\tD) optD" or "optA  B) optB C) optC"
        # Also handle format without tabs: "Beijing, China.	B) Athens, Greece.   C) Rome"
        if re.search(r'B\s*\)', line, re.IGNORECASE):
            # Find B) marker position
            b_match = re.search(r'B\s*\)', line, re.IGNORECASE)
            if b_match:
                # Option A is text before B)
                opt_a = line[:b_match.start()].strip().rstrip('\t ')
                if opt_a:
                    # Clean option A - remove A) or A. marker if present
                    opt_a_clean = re.sub(r'^[Aa]\s*[.)]\s*', '', opt_a)
                    # Also handle newline - take only first part
                    if '\n' in opt_a_clean:
                        opt_a_clean = opt_a_clean.split('\n')[0].strip()
                    opts.append(opt_a_clean if opt_a_clean else opt_a)
                else:
                    # Image option - add placeholder
                    opts.append('[Image A]')
                # Find all B-E markers
                markers = list(re.finditer(r'([B-E])\s*\)', line, re.IGNORECASE))
                for idx, m in enumerate(markers):
                    start = m.end()
                    if idx + 1 < len(markers):
                        end = markers[idx + 1].start()
                    else:
                        end = len(line)
                    opt_text = line[start:end].strip().rstrip('\t ')
                    if opt_text:
                        # Handle newline - take only first part
                        if '\n' in opt_text:
                            opt_text = opt_text.split('\n')[0].strip()
                        opts.append(opt_text)
                    else:
                        # Image option - add placeholder
                        opts.append(f'[Image {m.group(1).upper()}]')
        # Format: "A) optA B) optB C) optC" (with A marker)
        elif re.match(r'^A\s*\)', line, re.IGNORECASE):
            markers = list(re.finditer(r'([A-E])\s*\)', line, re.IGNORECASE))
            for idx, m in enumerate(markers):
                start = m.end()
                if idx + 1 < len(markers):
                    end = markers[idx + 1].start()
                else:
                    end = len(line)
                opt_text = line[start:end].strip().rstrip('\t ')
                if opt_text:
                    # Handle newline - take only first part
                    if '\n' in opt_text:
                        opt_text = opt_text.split('\n')[0].strip()
                    opts.append(opt_text)
        # Format: "C) optC D) optD" or "D) optD E) optE" (continuation line)
        elif re.match(r'^[C-E]\s*\)', line, re.IGNORECASE):
            markers = list(re.finditer(r'([C-E])\s*\)', line, re.IGNORECASE))
            for idx, m in enumerate(markers):
                start = m.end()
                if idx + 1 < len(markers):
                    end = markers[idx + 1].start()
                else:
                    end = len(line)
                opt_text = line[start:end].strip().rstrip('\t ')
                if opt_text:
                    # Handle newline - take only first part
                    if '\n' in opt_text:
                        opt_text = opt_text.split('\n')[0].strip()
                    opts.append(opt_text)
        # Format: "A. optA\tB. optB\tC. optC..." (dot-separated, tab/space delimited)
        elif re.match(r'^A\.\s*\S', line):
            markers = list(re.finditer(r'([A-E])\.\s*', line))
            for idx, m in enumerate(markers):
                start = m.end()
                if idx + 1 < len(markers):
                    end = markers[idx + 1].start()
                else:
                    end = len(line)
                opt_text = line[start:end].strip().rstrip('\t ')
                if opt_text:
                    if '\n' in opt_text:
                        opt_text = opt_text.split('\n')[0].strip()
                    opts.append(opt_text)
        return opts

    def get_answer_from_option_line(opt_line: str) -> str:
        """Get answer from paragraph highlight map for an option line."""
        if opt_line in para_highlight_map:
            return str(para_highlight_map[opt_line])
        return ''

    def clean_option_text(opt: str) -> str:
        """Clean option text by removing markers and extracting just the content.

        Handles formats like:
        - 'A) Option text' -> 'Option text'
        - 'A. Option text' -> 'Option text'
        - 'a) Option text' -> 'Option text'
        - 'A) Option text\n34. Next question' -> 'Option text' (split at newline)
        """
        if not opt:
            return opt

        # First, split at newline and take only first part (in case question is embedded)
        if '\n' in opt:
            opt = opt.split('\n')[0].strip()

        # Remove option markers: A), A., a), a., etc.
        opt_match = re.match(r'^([A-Ea-e])[.)]\s*(.*)$', opt)
        if opt_match:
            return opt_match.group(2).strip()

        return opt.strip()

    def clean_options_list(options: List[str]) -> List[str]:
        """Clean a list of options, removing markers and embedded questions."""
        cleaned = []
        for opt in options:
            clean_opt = clean_option_text(opt)
            if clean_opt:
                cleaned.append(clean_opt)
        return cleaned

    def is_option_line_envie(line: str) -> bool:
        """Check if line contains options."""
        # Options with B) marker (tab or space separated)
        if re.search(r'B\s*\)', line, re.IGNORECASE):
            return True
        # Options starting with A)
        if re.match(r'^A\s*\)', line, re.IGNORECASE):
            return True
        # Image-only options: "B)\tC)\tD)"
        if re.match(r'^B\s*\)\s*\t', line, re.IGNORECASE):
            return True
        # Continuation line: "C) opt D) opt" or "D) opt E) opt"
        if re.match(r'^[C-E]\s*\)', line, re.IGNORECASE):
            return True
        # Options with A. B. format (e.g., "A. Stork" or "A.It occurs")
        if re.match(r'^A\.\s*\S', line):
            return True
        return False

    def has_fill_blank(text: str) -> bool:
        return '…' in text or '___' in text or '...' in text

    def is_sentence_with_blank(text: str) -> bool:
        """Check if text is a sentence containing a fill-blank (transform sentence)."""
        # Pattern: sentence with … marker, ending with period
        if has_fill_blank(text) and text.endswith('.'):
            return True
        return False

    def is_instruction_line(text: str) -> bool:
        """Check if line is an instruction (not a question)."""
        # If it has fill-blank marker, it's a question, not instruction
        if has_fill_blank(text):
            return False
        lower = text.lower()
        # Filter out artifact/watermark names (level names from competition exams)
        artifact_names = {'benjamin', 'ecolier', 'preecolier', 'student', 'junior', 'cadet', 'wallaby',
                          'pre-ecolier', 'pre ecolier', 'kadett', 'koala'}
        stripped = text.strip().lower()
        if stripped in artifact_names:
            return True
        # Also match repeated artifact like "StudentStudent" or "JuniorJunior"
        for name in artifact_names:
            if stripped == name * 2 or stripped == name + name:
                return True
        patterns = [
            'for each question',
            'read the following',
            'read and look',
            'read, look',
            'choose the correct',
            'choose the best',
            'answer the questions',
            'questions 1-',
            'questions (1-',
            'for questions',
            'for each sentence',
            'for each group',
            'choose the right answer to define',
            'read the text',
        ]
        return any(p in lower for p in patterns)

    # Track question numbers we've seen to detect standalone number patterns
    seen_q_numbers = set()

    # Track which paragraphs were already processed
    processed_paragraphs = set()

    # Track dialogue sections (e.g., Q11-15 with 3 response options each)
    in_dialogue_section = False
    dialogue_max_opts = 3

    # ============ WORD NUMBERING PASS: Handle Word List Paragraph format ============
    # For files like "Kangaroo Start R2 demo.docx" where:
    # - Questions have num_level=0 (numbered list items)
    # - Options have num_level=1 (sub-list items) with tab-separated B) C) D)
    # - IMPORTANT: Each question has exactly ONE option paragraph (1:1 ratio)
    # Check if document uses Word numbering format
    num_level_0_count = sum(1 for elem in doc_elements if elem['type'] == 'paragraph' and elem.get('num_level') == 0)
    num_level_1_count = sum(1 for elem in doc_elements if elem['type'] == 'paragraph' and elem.get('num_level') == 1)

    has_word_numbering = num_level_0_count > 0

    # Check if this is a clean 1:1 format (each question has exactly one option paragraph)
    # Story file has 30 questions but 54 options (options split across multiple paragraphs)
    # Start R2 file has 25 questions and 25 options (1:1)
    is_1to1_format = (num_level_0_count > 0 and num_level_0_count == num_level_1_count)

    # Check if options have duplicate pattern (old format files)
    # Old format: "book.book.sign.sign.map.map." - words repeated due to Word formatting
    # New R2 format: "park\tB) beach\tC) school" - clean options
    has_duplicate_options = False
    if has_word_numbering and is_1to1_format:
        for elem in doc_elements:
            if elem['type'] == 'paragraph' and elem.get('num_level') == 1:
                opt_text = elem['text']
                # Check for duplicate pattern like "word.word." or "wordword"
                # Split by B) and check the first option (A)
                if 'B)' in opt_text:
                    before_b = opt_text.split('B)')[0].strip()
                    # Remove StartStart prefix for checking
                    before_b = re.sub(r'^(Start)+', '', before_b)
                    # Check for duplicate words: "book.book." or "a cameraa camera"
                    if re.search(r'(\w{3,}\.)\1', before_b) or re.search(r'(\w{4,})\1', before_b):
                        has_duplicate_options = True
                        break

    # Only use Word numbering pass for:
    # 1. Clean 1:1 format (each question has exactly one option paragraph)
    # 2. No duplicate patterns in options
    if has_word_numbering and is_1to1_format and not has_duplicate_options:
        word_num_questions = []
        elem_idx = 0

        while elem_idx < len(doc_elements):
            elem = doc_elements[elem_idx]

            # Skip non-paragraph elements
            if elem['type'] != 'paragraph':
                elem_idx += 1
                continue

            # Check if this is a question (num_level=0) or instruction
            if elem.get('num_level') == 0:
                text = elem['text']
                q_num = elem.get('num_value')

                # Skip instruction lines
                if is_instruction_line(text):
                    elem_idx += 1
                    continue

                # Check if next element is options (num_level=1)
                if elem_idx + 1 < len(doc_elements):
                    next_elem = doc_elements[elem_idx + 1]

                    if next_elem['type'] == 'paragraph' and next_elem.get('num_level') == 1:
                        opt_text = next_elem['text']
                        # Clean watermark/header artifacts (e.g., "StartStart" prefix)
                        # These appear when document has repeated header text runs
                        opt_text = re.sub(r'^(Start)+', '', opt_text)
                        # Parse options from format: "optA\tB) optB\tC) optC\tD) optD"
                        opts = extract_options_envie(opt_text)

                        if len(opts) >= 2:
                            # Get answer from highlighted option if available
                            answer = ''
                            hl = next_elem.get('highlighted_option', 0)
                            if hl > 0:
                                answer = str(hl)

                            q_dict = {
                                'question': text,
                                'options': opts,
                                '_doc_pos': elem_idx,
                                '_q_num': q_num
                            }
                            if answer:
                                q_dict['answer'] = answer

                            word_num_questions.append(q_dict)
                            processed_paragraphs.add(text)
                            processed_paragraphs.add(opt_text)
                            elem_idx += 2
                            continue

            elem_idx += 1

        # If we found questions using Word numbering, use them and skip other parsing
        if word_num_questions:
            # Sort by question number
            word_num_questions.sort(key=lambda q: q.get('_q_num', 0))
            questions.extend(word_num_questions)
            # Return early - no need for other parsing passes
            return questions
    # ============ END WORD NUMBERING PASS ============

    # ============ PRE-PROCESS: Parse paragraph + table pairs ============
    # For formats where question is in paragraph and options are in following table
    # This handles Red Kangaroo style: "Question text ending with ? or :" followed by table with options
    para_table_questions = []

    elem_idx = 0
    while elem_idx < len(doc_elements):
        elem = doc_elements[elem_idx]

        if elem['type'] == 'paragraph':
            text = elem['text']
            # Check if this looks like a question
            # Patterns: ends with ? or :, has fill-blank (___), ends with quote
            has_fill_blank_marker = '___' in text or '______' in text
            # Also check if followed by table (synthetic paragraph from table splitting)
            next_is_table = (
                elem_idx + 1 < len(doc_elements) and
                doc_elements[elem_idx + 1]['type'] == 'table'
            )
            ends_with_ellipsis = text.endswith('...') or text.endswith('…')
            # Lower min length for texts ending with ... or followed by table
            min_len = 10 if (next_is_table or ends_with_ellipsis) else 15
            # Check for ? inside text followed by parenthetical, e.g., "...bao nhiêu? (g = 10 m/s²)"
            has_question_mark_inside = bool(re.search(r'\?\s*\([^)]+\)\s*$', text))
            is_question_text = (
                len(text) > min_len and
                not is_instruction_line(text) and
                (
                    text.endswith(':') or
                    text.endswith('?') or
                    text.endswith('"') or  # Dialogue/quote questions
                    has_fill_blank_marker or  # Fill-in-blank questions
                    ends_with_ellipsis or  # Questions ending with "..."
                    has_question_mark_inside or  # "...? (params)" format
                    (text.endswith('.') and next_is_table)  # Table-split questions like "Pick out the odd one."
                )
            )

            if is_question_text:
                # Look ahead for a table with options OR paragraph options (A., B., C., D.)
                if elem_idx + 1 < len(doc_elements):
                    next_elem = doc_elements[elem_idx + 1]

                    # Case 1: Options in table
                    table_cell_count = next_elem.get('cell_count', 0) if next_elem['type'] == 'table' else 0
                    if next_elem['type'] == 'table' and (len(next_elem['options']) >= 2 or table_cell_count >= 3):
                        options = next_elem['options']

                        # Check for Cloze passage tables (options start with number like "15.A)", "16.A)")
                        first_opt = options[0] if options else ''
                        is_cloze_table = re.match(r'^\d+\.\s*[A-E]\)', first_opt)
                        if is_cloze_table:
                            # Parse Cloze table - extract questions with passage
                            # Find passage before this table (look back for instruction line)
                            passage_lines = []
                            for back_idx in range(elem_idx - 1, max(0, elem_idx - 15), -1):
                                back_elem = doc_elements[back_idx]
                                if back_elem['type'] == 'paragraph':
                                    back_text = back_elem['text']
                                    # Stop at instruction line or previous question
                                    if 'Read the following' in back_text or 'choose the best' in back_text.lower():
                                        passage_lines.insert(0, back_text)
                                        break
                                    # Skip header/watermark lines
                                    if back_text.startswith('Red Kangaroo'):
                                        continue
                                    passage_lines.insert(0, back_text)

                            cloze_passage = '\n'.join(passage_lines) if passage_lines else ''

                            # Parse Cloze options - group by question number
                            cloze_opts_by_num = {}
                            cloze_highlights = {}  # Track highlighted option per question
                            current_q_num = None
                            current_opts = []
                            opt_idx_in_q = 0

                            for opt in options:
                                # Check if starts with question number (15.A), 16.A), etc.)
                                q_match = re.match(r'^(\d+)\.\s*[A-E]\)\s*(.*)$', opt)
                                if q_match:
                                    # Save previous question
                                    if current_q_num and current_opts:
                                        cloze_opts_by_num[current_q_num] = current_opts
                                    current_q_num = int(q_match.group(1))
                                    current_opts = [q_match.group(2)]
                                    opt_idx_in_q = 1
                                else:
                                    # Continuation option (B), C), D))
                                    opt_match = re.match(r'^[B-E]\)\s*(.*)$', opt)
                                    if opt_match:
                                        current_opts.append(opt_match.group(1))
                                        opt_idx_in_q += 1

                            # Save last question
                            if current_q_num and current_opts:
                                cloze_opts_by_num[current_q_num] = current_opts

                            # Get highlighted answers from cloze_highlights map
                            cloze_highlights = next_elem.get('cloze_highlights', {})

                            # Create Cloze questions
                            for q_num in sorted(cloze_opts_by_num.keys()):
                                opts = cloze_opts_by_num[q_num]
                                q_dict = {
                                    'question': f'Cloze question {q_num}\n\n{cloze_passage}' if cloze_passage else f'Cloze question {q_num}',
                                    'options': opts,
                                    '_doc_pos': elem_idx  # Track document position
                                }
                                # Add answer from highlighted option
                                if q_num in cloze_highlights:
                                    q_dict['answer'] = str(cloze_highlights[q_num])
                                para_table_questions.append(q_dict)

                            # Mark passage paragraphs as processed
                            for pl in passage_lines:
                                processed_paragraphs.add(pl)

                            elem_idx += 2  # Skip paragraph and table
                            continue

                        # Found question + table options pattern
                        # Clean option prefixes (A., B., C., D., E. or A), B), etc.)
                        cleaned_options = clean_options_list(options)

                        # If table has cells but few text options, fill with placeholders
                        if len(cleaned_options) < 3 and table_cell_count >= 3:
                            # Determine expected option count from cell layout
                            expected_opts = min(table_cell_count, 5)
                            while len(cleaned_options) < expected_opts:
                                letter = chr(ord('A') + len(cleaned_options))
                                cleaned_options.append(f'[{letter}]')

                        q_dict = {
                            'question': text,
                            'options': cleaned_options,
                            '_doc_pos': elem_idx  # Track document position
                        }
                        # Add correct answer from highlighted option
                        if next_elem.get('highlighted', 0) > 0:
                            q_dict['answer'] = str(next_elem['highlighted'])

                        para_table_questions.append(q_dict)
                        processed_paragraphs.add(text)
                        elem_idx += 2  # Skip both paragraph and table
                        continue

                    # Case 2: Options in paragraphs (A., B., C., D. as separate paragraphs)
                    elif next_elem['type'] == 'paragraph' and re.match(r'^A\.\s*\S', next_elem['text']):
                        # Collect paragraph options
                        para_options = []
                        opt_idx = elem_idx + 1
                        while opt_idx < len(doc_elements):
                            opt_elem = doc_elements[opt_idx]
                            if opt_elem['type'] != 'paragraph':
                                break
                            opt_match = re.match(r'^([A-E])\.\s*(.+)$', opt_elem['text'])
                            if opt_match:
                                para_options.append(opt_match.group(2))
                                processed_paragraphs.add(opt_elem['text'])
                                opt_idx += 1
                            else:
                                break

                        if len(para_options) >= 2:
                            q_dict = {
                                'question': text,
                                'options': para_options,
                                '_doc_pos': elem_idx  # Track document position
                            }
                            para_table_questions.append(q_dict)
                            processed_paragraphs.add(text)
                            elem_idx = opt_idx
                            continue

        elem_idx += 1

    # ============ SECOND PASS: Handle standalone Cloze tables ============
    # Cloze tables not preceded by a question paragraph (e.g., Red Kangaroo format)
    # These tables have options starting with "11.A)", "12.A)", etc.
    processed_tables = set()  # Track which tables were processed
    for pt_q in para_table_questions:
        # Mark tables used in para_table_questions
        pass

    elem_idx = 0
    while elem_idx < len(doc_elements):
        elem = doc_elements[elem_idx]

        if elem['type'] == 'table':
            options = elem.get('options', [])
            if options:
                first_opt = options[0]
                # Check for Cloze table format (starts with number like "11.A)", "12.A)")
                is_cloze_table = re.match(r'^(\d+)\.\s*[A-E]\)', first_opt)

                if is_cloze_table:
                    first_q_num = int(is_cloze_table.group(1))

                    # Check if this Cloze table was already processed
                    already_processed = any(
                        f'Cloze question {first_q_num}' in q.get('question', '')
                        for q in para_table_questions
                    )

                    if not already_processed:
                        # Find passage before this table (look back for instruction + content)
                        # Strategy: First find the instruction line, then collect passage after it
                        instruction_idx = -1
                        instruction_text = ''
                        for back_idx in range(elem_idx - 1, max(0, elem_idx - 25), -1):
                            back_elem = doc_elements[back_idx]
                            if back_elem['type'] == 'paragraph':
                                back_text = back_elem['text']
                                # Check for instruction line with question range (e.g., "(11-20)")
                                # Must check BEFORE skipping headers, as instruction may start with header text
                                if f'({first_q_num}-' in back_text or f'({first_q_num} -' in back_text:
                                    # Extract the instruction part (after header if present)
                                    if 'Read the' in back_text:
                                        read_pos = back_text.find('Read the')
                                        instruction_text = back_text[read_pos:]
                                    else:
                                        instruction_text = back_text
                                    instruction_idx = back_idx
                                    break
                                # Also check for generic instruction patterns
                                if ('Read the text' in back_text or 'Read the following' in back_text) and 'choose the best' in back_text.lower():
                                    instruction_idx = back_idx
                                    instruction_text = back_text
                                    break
                                # Skip header/watermark lines (only if no instruction pattern found)
                                if back_text.startswith('Red Kangaroo') or back_text.startswith('Grey Kangaroo'):
                                    continue

                        # Now collect passage lines AFTER instruction (between instruction and Cloze table)
                        passage_lines = []
                        if instruction_idx >= 0:
                            passage_lines.append(instruction_text)
                            for fwd_idx in range(instruction_idx + 1, elem_idx):
                                fwd_elem = doc_elements[fwd_idx]
                                if fwd_elem['type'] == 'paragraph':
                                    fwd_text = fwd_elem['text']
                                    # Skip header/watermark lines
                                    if fwd_text.startswith('Red Kangaroo') or fwd_text.startswith('Grey Kangaroo'):
                                        continue
                                    passage_lines.append(fwd_text)
                        else:
                            # Fallback: collect a few paragraphs before table if no instruction found
                            for back_idx in range(elem_idx - 1, max(0, elem_idx - 5), -1):
                                back_elem = doc_elements[back_idx]
                                if back_elem['type'] == 'paragraph':
                                    back_text = back_elem['text']
                                    if back_text.startswith('Red Kangaroo') or back_text.startswith('Grey Kangaroo'):
                                        continue
                                    passage_lines.insert(0, back_text)

                        cloze_passage = '\n'.join(passage_lines) if passage_lines else ''

                        # Parse Cloze options - group by question number
                        cloze_opts_by_num = {}
                        current_q_num = None
                        current_opts = []

                        for opt in options:
                            # Check if starts with question number (11.A), 12.A), etc.)
                            q_match = re.match(r'^(\d+)\.\s*[A-E]\)\s*(.*)$', opt)
                            if q_match:
                                # Save previous question
                                if current_q_num and current_opts:
                                    cloze_opts_by_num[current_q_num] = current_opts
                                current_q_num = int(q_match.group(1))
                                current_opts = [q_match.group(2)]
                            else:
                                # Continuation option (B), C), D))
                                opt_match = re.match(r'^[B-E]\)\s*(.*)$', opt)
                                if opt_match:
                                    current_opts.append(opt_match.group(1))

                        # Save last question
                        if current_q_num and current_opts:
                            cloze_opts_by_num[current_q_num] = current_opts

                        # Get highlighted answers from cloze_highlights map
                        cloze_highlights = elem.get('cloze_highlights', {})

                        # Create Cloze questions
                        for q_num in sorted(cloze_opts_by_num.keys()):
                            opts = cloze_opts_by_num[q_num]
                            q_dict = {
                                'question': f'Cloze question {q_num}\n\n{cloze_passage}' if cloze_passage else f'Cloze question {q_num}',
                                'options': opts,
                                '_doc_pos': elem_idx  # Track document position
                            }
                            # Add answer from highlighted option
                            if q_num in cloze_highlights:
                                q_dict['answer'] = str(cloze_highlights[q_num])
                            para_table_questions.append(q_dict)

                        # Mark passage paragraphs as processed
                        for pl in passage_lines:
                            processed_paragraphs.add(pl)

        elem_idx += 1
    # ============ END SECOND PASS ============

    # Add para_table_questions with document position for later sorting
    # Each question needs _doc_pos to maintain document order
    for q in para_table_questions:
        if '_doc_pos' not in q:
            # Find position based on question text in paragraphs
            q_text = q.get('question', '')[:50]
            for pi, para in enumerate(paragraphs):
                if q_text and q_text in para:
                    q['_doc_pos'] = pi
                    break
            else:
                q['_doc_pos'] = 9999  # Default to end if not found

    questions.extend(para_table_questions)
    # ============ END PRE-PROCESS ============

    i = 0
    while i < len(paragraphs):
        line = paragraphs[i].strip()
        if not line:
            i += 1
            continue

        # Skip paragraphs already processed in para+table phase
        if line in processed_paragraphs:
            i += 1
            continue

        # Skip pure option lines (they belong to previous question)
        # Cloze option lines are handled separately in the textbox section
        if is_option_line_envie(line) and not has_fill_blank(line):
            # Check if this is a numbered cloze question "22. A) opt B) opt..."
            # Skip these as they're processed in the textbox section later
            is_numbered_cloze = re.match(r'^(\d+)\.\s*A\s*\)', line)
            if is_numbered_cloze:
                i += 1
                continue
            # Check if this is a standalone cloze option (A) at start with tab-separated B) C) D))
            # Also skip these - handled in textbox section
            is_standalone_cloze = re.match(r'^A\s*\)', line) and '\t' in line and re.search(r'B\s*\).*C\s*\).*D\s*\)', line)
            if is_standalone_cloze:
                i += 1
                continue
            # Regular option line - skip
            i += 1
            continue

        # Case 0a: Wallaby Q1-5 special format (must be BEFORE instruction skip)
        # Structure: "For each question (1-5)..." instruction
        # Then: passage paragraphs for Q1 (no number marker, can be multiple paragraphs)
        # Then: "2.", "3.", "4.", "5." markers (all consecutive)
        # Then: 15 option paragraphs (3 for each Q1-5)
        # Question content is in textboxes (Q2-5) or paragraphs before "2." (Q1)
        if 'question (1-5)' in line.lower() or 'questions (1-5)' in line.lower():
            # Extract textbox content for Q2-5 passages
            from lxml import etree
            try:
                body = doc.element.body
                xml_str = etree.tostring(body, encoding='unicode')
                textbox_matches = re.findall(r'<w:txbxContent[^>]*>(.*?)</w:txbxContent>', xml_str, re.DOTALL)

                # Get unique textbox contents (they're often duplicated)
                textbox_contents = []
                seen_tb = set()
                for tb in textbox_matches:
                    texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', tb)
                    content = ''.join(texts).strip()
                    # Skip short content, headers, and cloze options
                    if content and len(content) > 30 and content not in seen_tb:
                        if not re.match(r'^\d+\.?\s*A\s*\)', content):  # Not cloze options
                            if 'Wallaby' not in content and 'Answer a maximum' not in content:
                                seen_tb.add(content)
                                textbox_contents.append(content)
            except:
                textbox_contents = []

            # Find all consecutive standalone numbers and passages before them
            j = i + 1
            passages_before_nums = []
            standalone_nums = []

            # First pass: collect everything before we see standalone numbers
            while j < len(paragraphs):
                next_para = paragraphs[j].strip()
                if not next_para:
                    j += 1
                    continue
                # Check for standalone number
                num_match = re.match(r'^(\d+)\.$', next_para)
                if num_match:
                    standalone_nums.append(int(num_match.group(1)))
                    j += 1
                    # Continue collecting more numbers
                    while j < len(paragraphs):
                        np = paragraphs[j].strip()
                        if not np:
                            j += 1
                            continue
                        nm = re.match(r'^(\d+)\.$', np)
                        if nm:
                            standalone_nums.append(int(nm.group(1)))
                            j += 1
                        else:
                            break
                    break
                else:
                    passages_before_nums.append(next_para)
                    j += 1

            # If we found numbers 2-5 pattern
            if standalone_nums and 2 in standalone_nums:
                # Count questions (1 + number of standalone markers)
                num_questions = 1 + len(standalone_nums)
                # Collect 3 options per question
                opts_list = []
                while j < len(paragraphs) and len(opts_list) < num_questions * 3:
                    opt_para = paragraphs[j].strip()
                    if not opt_para:
                        j += 1
                        continue
                    if is_instruction_line(opt_para) or is_option_line_envie(opt_para):
                        break
                    opts_list.append(opt_para)
                    j += 1

                # Build question contents
                # Q1: from paragraphs before "2." marker
                q1_content = ' '.join(passages_before_nums) if passages_before_nums else 'Question 1'

                # Q2-5: from textboxes (first 4 unique textbox contents)
                q_contents = [q1_content]
                for tb_idx in range(min(4, len(textbox_contents))):
                    # Truncate long passages for display
                    tb_text = textbox_contents[tb_idx]
                    if len(tb_text) > 200:
                        tb_text = tb_text[:200] + '...'
                    q_contents.append(tb_text)

                # Pad with placeholders if not enough textboxes
                while len(q_contents) < num_questions:
                    q_contents.append(f'Question {len(q_contents) + 1}')

                # Create questions from collected options
                if len(opts_list) >= num_questions * 3:
                    for q_idx in range(num_questions):
                        q_opts = opts_list[q_idx*3:(q_idx+1)*3]
                        if len(q_opts) == 3:
                            questions.append({
                                'question': q_contents[q_idx],
                                'options': q_opts,
                                '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                            })
                            seen_q_numbers.add(str(q_idx + 1))
                    i = j
                    continue
            i += 1
            continue

        # Skip general instruction lines (after handling Q1-5 special case)
        # BUT: If instruction line is followed by option line, treat it as a question
        # EXCEPT: If instruction mentions question range like "(21-30)", always skip
        if is_instruction_line(line):
            # Detect dialogue section: "For each question (11-15), read and choose"
            # Dialogue questions have exactly 3 response options
            lower_line = line.lower()
            if re.search(r'\(\d+-\d+\)', line):
                # Reset dialogue section flag by default for any new section
                in_dialogue_section = False
                if 'read and choose' in lower_line or 'choose the best answer' in lower_line:
                    range_match = re.search(r'\((\d+)-(\d+)\)', line)
                    if range_match:
                        r_start, r_end = int(range_match.group(1)), int(range_match.group(2))
                        if r_end - r_start == 4:  # 5 questions like (11-15)
                            in_dialogue_section = True
                            dialogue_max_opts = 3
            # Always skip if it's a section instruction with question range
            if re.search(r'\(\d+-\d+\)', line):
                i += 1
                continue
            # Check if next non-empty line is an option line
            j = i + 1
            while j < len(paragraphs):
                np = paragraphs[j].strip()
                if np:
                    break
                j += 1
            if j < len(paragraphs) and is_option_line_envie(paragraphs[j].strip()):
                # Check if next line is a numbered cloze option (e.g., "22.\tA) neither")
                # If so, skip this instruction - it's for Cloze section
                next_line_text = paragraphs[j].strip()
                if re.match(r'^\d+\.\s*\t?A\s*\)', next_line_text):
                    i += 1
                    continue
                # This instruction acts as a question - process it below
                pass
            else:
                i += 1
                continue

        # Case 0b: Standalone question number "1." or "2." or just "1", "2" etc.
        # Ecolier format: number → question text → merged options (A+B on one line, C+D on next)
        # EN-VIE format: number → 3 plain option paragraphs
        q_num_match = re.match(r'^(\d+)\.?$', line)
        if q_num_match:
            q_num = q_num_match.group(1)
            if q_num not in seen_q_numbers:
                seen_q_numbers.add(q_num)
                # Look ahead: next non-empty line could be question text (Ecolier) or option (EN-VIE)
                j = i + 1

                # Find next non-empty paragraph
                while j < len(paragraphs) and not paragraphs[j].strip():
                    j += 1

                if j < len(paragraphs):
                    next_text = paragraphs[j].strip()

                    # Ecolier pattern: question text on next line, then merged options
                    # Check if the line AFTER the question text has merged options
                    if next_text and not is_option_line_envie(next_text) and not re.match(r'^\d+\.?$', next_text):
                        # This looks like a question text, check what follows
                        k = j + 1
                        while k < len(paragraphs) and not paragraphs[k].strip():
                            k += 1
                        if k < len(paragraphs) and is_option_line_envie(paragraphs[k].strip()):
                            # Ecolier format: number → question → merged options
                            question_text = next_text
                            opts = extract_options_envie(paragraphs[k].strip())
                            m = k + 1
                            # Check for continuation lines (C) D) on next line)
                            while m < len(paragraphs) and len(opts) < 5:
                                cont_line = paragraphs[m].strip()
                                if not cont_line:
                                    m += 1
                                    continue
                                if re.match(r'^[C-E]\s*\)', cont_line, re.IGNORECASE):
                                    more_opts = extract_options_envie(cont_line)
                                    if more_opts:
                                        opts.extend(more_opts)
                                        m += 1
                                        continue
                                break
                            if opts:
                                questions.append({
                                    'question': question_text,
                                    'options': opts[:5],
                                    'number': int(q_num),
                                    '_doc_pos': para_idx_to_doc_idx.get(i, i)
                                })
                                i = m
                                continue

                # Fallback: EN-VIE format - collect next 3 non-empty paragraphs as options
                opts = []
                j = i + 1
                while j < len(paragraphs) and len(opts) < 3:
                    opt_line = paragraphs[j].strip()
                    if not opt_line:
                        j += 1
                        continue
                    # Stop if we hit another question number or option line
                    if re.match(r'^\d+\.?$', opt_line) or is_option_line_envie(opt_line):
                        break
                    if is_instruction_line(opt_line):
                        j += 1
                        continue
                    opts.append(opt_line)
                    j += 1

                if len(opts) == 3:
                    questions.append({
                        'question': f'Question {q_num} (reading comprehension)',
                        'options': opts,
                        'number': int(q_num),
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)
                    })
                    i = j
                    continue
            i += 1
            continue

        # Find next non-empty line
        next_line = None
        next_idx = i + 1
        while next_idx < len(paragraphs):
            if paragraphs[next_idx].strip():
                next_line = paragraphs[next_idx].strip()
                break
            next_idx += 1

        # Case 1: Fill-blank question followed by option line(s)
        if has_fill_blank(line):
            # Case 1a: Options in tab-separated format
            if next_line and is_option_line_envie(next_line):
                opts = extract_options_envie(next_line)
                j = next_idx + 1
                # Check for multi-line options (C) D) on next line)
                while j < len(paragraphs) and len(opts) < 4:
                    cont_line = paragraphs[j].strip()
                    if not cont_line:
                        j += 1
                        continue
                    # Check if line starts with C) or D) (continuation)
                    if re.match(r'^[C-E]\s*\)', cont_line, re.IGNORECASE):
                        more_opts = extract_options_envie(cont_line)
                        if more_opts:
                            opts.extend(more_opts)
                            j += 1
                            continue
                    break
                if opts:
                    # Word Groups format: Look back to collect additional fill-blank paragraphs
                    question_parts = [line]
                    for back_idx in range(i - 1, max(0, i - 10), -1):
                        back_text = paragraphs[back_idx].strip()
                        if not back_text:
                            continue
                        # Stop at instruction line
                        if is_instruction_line(back_text):
                            break
                        # Stop at options line (previous question's options)
                        if is_option_line_envie(back_text) and not has_fill_blank(back_text):
                            break
                        # Collect consecutive fill-blank paragraphs
                        if has_fill_blank(back_text):
                            question_parts.insert(0, back_text)
                        else:
                            break

                    combined_question = ' '.join(question_parts)
                    questions.append({
                        'question': combined_question,
                        'options': opts,
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                    })
                    i = j
                    continue

            # Case 1b: Options as separate paragraphs (no A/B/C markers)
            # Common in Story format, Red Kangaroo has 5 options (A-E)
            if next_line and not is_option_line_envie(next_line):
                opts = []
                j = i + 1
                while j < len(paragraphs) and len(opts) < 5:  # Max 5 options (A-E)
                    opt_line = paragraphs[j].strip()
                    if not opt_line:
                        j += 1
                        continue
                    # Stop if we hit another question or option line
                    if opt_line.endswith('?') or is_option_line_envie(opt_line) or has_fill_blank(opt_line):
                        break
                    # Skip instruction lines
                    if is_instruction_line(opt_line):
                        j += 1
                        continue
                    opts.append(opt_line)
                    j += 1

                if len(opts) >= 3:  # Need at least 3 options
                    questions.append({
                        'question': line,
                        'options': opts[:5],  # Max 5 options (A-E)
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                    })
                    i = j
                    continue

        # Case 2: Question ending with ? followed by option line (text or image)
        if line.endswith('?') and len(line) > 10:
            # First check if next line is an option line (image options)
            if next_line and is_option_line_envie(next_line):
                opts = extract_options_envie(next_line)
                j = next_idx + 1
                # Check for continuation lines (C) D) on next line)
                while j < len(paragraphs) and len(opts) < 5:
                    cont_line = paragraphs[j].strip()
                    if not cont_line:
                        j += 1
                        continue
                    # Check if line starts with C), D), or E) (continuation)
                    if re.match(r'^[C-E]\s*\)', cont_line, re.IGNORECASE):
                        more_opts = extract_options_envie(cont_line)
                        if more_opts:
                            opts.extend(more_opts)
                            j += 1
                            continue
                    break
                if opts:
                    q_dict = {
                        'question': line,
                        'options': opts[:5],  # Max 5 options (A-E)
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                    }
                    # Check for highlighted answer in option line
                    ans = get_answer_from_option_line(next_line)
                    if ans:
                        q_dict['answer'] = ans
                    questions.append(q_dict)
                    i = j
                    continue

            # Otherwise collect next 3-5 paragraphs as options (no markers)
            # Common in Q1-5 sections with reading comprehension
            # Benjamin has 5 options (A-E), other levels may have 3-4
            # In dialogue sections, limit to 3 options to avoid eating next question
            max_opts = dialogue_max_opts if in_dialogue_section else 5
            opts = []
            j = i + 1
            while j < len(paragraphs) and len(opts) < max_opts:
                opt_line = paragraphs[j].strip()
                if not opt_line:
                    j += 1
                    continue
                # Stop if we hit another question or option line
                if opt_line.endswith('?') or is_option_line_envie(opt_line) or has_fill_blank(opt_line):
                    break
                # Stop if line ends with comma (likely start of next question, not an option)
                if opt_line.endswith(','):
                    break
                # Skip instruction lines
                if is_instruction_line(opt_line):
                    j += 1
                    continue
                # Stop if line is too long (likely a passage, not an option)
                if len(opt_line) > 200:
                    break
                # This looks like an option
                opts.append(opt_line)
                j += 1

            if len(opts) >= 3:  # Reading comprehension has 3-5 options
                questions.append({
                    'question': line,
                    'options': opts[:max_opts],
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i = j
                continue

        # Case 2b: Dialogue question NOT ending with ? (Grey Kangaroo Q12, Q14, Q15)
        # Format: "I think we should come up with a plan B." followed by 3 response options
        # Distinguishing features:
        # - Short sentence ending with period (dialogue prompt)
        # - Following 3 lines are responses (not option markers like A) B))
        # - Must be in dialogue section (after "For each question (11-15)")
        if line.endswith('.') and len(line) > 15 and len(line) < 80 and not has_fill_blank(line):
            # First check for special format: next line has "optA\tB) optB" and following line has "C) optC"
            # This is Grey Kangaroo Q14 format
            j = i + 1
            while j < len(paragraphs):
                np = paragraphs[j].strip()
                if not np:
                    j += 1
                    continue
                break

            if j < len(paragraphs):
                first_opt_line = paragraphs[j].strip()
                # Check for "optA\tB) optB" pattern - use extract_options_envie for proper parsing
                if re.search(r'\tB\s*\)', first_opt_line):
                    first_opts = extract_options_envie(first_opt_line)
                    if len(first_opts) >= 2:
                        # Look for continuation options (C), D), E)) on next lines
                        j += 1
                        while j < len(paragraphs) and len(first_opts) < 5:
                            np = paragraphs[j].strip()
                            if not np:
                                j += 1
                                continue
                            # Check if line starts with C), D), or E)
                            if re.match(r'^[C-E]\s*\)', np, re.IGNORECASE):
                                cont_opts = extract_options_envie(np)
                                if cont_opts:
                                    first_opts.extend(cont_opts)
                                    j += 1
                                    continue
                            break
                        if len(first_opts) >= 3:
                            all_opts = first_opts
                        questions.append({
                            'question': line,
                            'options': all_opts[:5],  # Max 5 options (A-E)
                            '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                        })
                        i = j
                        continue

            # Otherwise check for dialogue responses (no markers)
            dialogue_opts = []
            j = i + 1
            while j < len(paragraphs) and len(dialogue_opts) < 4:
                opt_line = paragraphs[j].strip()
                if not opt_line:
                    j += 1
                    continue
                # Stop if we hit a question, option line with markers, or instruction
                if opt_line.endswith('?') or has_fill_blank(opt_line) or is_instruction_line(opt_line):
                    break
                # Stop if line has B) C) markers (this is an option line for different question)
                if re.search(r'B\s*\)', opt_line):
                    break
                # Dialogue responses are short sentences (< 60 chars) ending with period
                if len(opt_line) < 70 and opt_line.endswith('.'):
                    dialogue_opts.append(opt_line)
                    j += 1
                else:
                    break

            # Dialogue questions have exactly 3-4 response options
            if len(dialogue_opts) == 3 or len(dialogue_opts) == 4:
                questions.append({
                    'question': line,
                    'options': dialogue_opts[:3],
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i = j
                continue

        # Case 3: Matching question - description followed by option line with 3+ choices
        # e.g., "Maria loves comfort food..." followed by "The Safari. B) Origo. C)..."
        # Also handles instruction line as question: "Fill in the missing letter" + "A\tB) I\tC) U"
        # Skip numbered cloze format "22. A) ..."
        if len(line) > 30 and not line.endswith('?') and not has_fill_blank(line) and not re.match(r'^(\d+)\.\s*A\s*\)', line):
            if next_line and is_option_line_envie(next_line):
                # Check if options are on separate paragraphs (A., B., C., D., E.)
                if re.match(r'^[A]\.\s*\S', next_line) and not re.search(r'B\s*[.)]\s*\S', next_line):
                    # Collect A., B., C., D., E. from separate paragraphs
                    para_opts = []
                    j = next_idx
                    while j < len(paragraphs) and len(para_opts) < 5:
                        opt_line = paragraphs[j].strip()
                        if not opt_line:
                            j += 1
                            continue
                        opt_m = re.match(r'^([A-E])\.\s*(.+)$', opt_line)
                        if opt_m:
                            para_opts.append(opt_m.group(2).strip())
                            j += 1
                        else:
                            break
                    if len(para_opts) >= 3:
                        questions.append({
                            'question': line,
                            'options': para_opts[:5],
                            '_doc_pos': para_idx_to_doc_idx.get(i, i)
                        })
                        i = j
                        continue
                # Extract options from option line (inline format)
                opts = extract_options_envie(next_line)
                j = next_idx + 1
                # Check for continuation lines (D), E) on next line)
                while j < len(paragraphs) and len(opts) < 5:
                    cont_line = paragraphs[j].strip()
                    if not cont_line:
                        j += 1
                        continue
                    # Check if line starts with D) or E) (continuation)
                    if re.match(r'^[D-E]\s*\)', cont_line, re.IGNORECASE):
                        more_opts = extract_options_envie(cont_line)
                        if more_opts:
                            opts.extend(more_opts)
                            j += 1
                            continue
                    break
                if len(opts) >= 3:  # At least 3 options (A, B, C)
                    q_dict = {
                        'question': line,
                        'options': opts[:5],
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                    }
                    ans = get_answer_from_option_line(next_line)
                    if ans:
                        q_dict['answer'] = ans
                    questions.append(q_dict)
                    i = j
                    continue

        # Case 4: Question ending with , followed by paragraph options (incomplete sentence)
        # e.g., "According to the notice," followed by 3 options
        # Or: Question split across paragraphs ending with , then . followed by option line
        if line.endswith(',') and len(line) > 15:
            # First check if next paragraph is a continuation (ends with .) followed by option line
            j = i + 1
            while j < len(paragraphs):
                np = paragraphs[j].strip()
                if np:
                    break
                j += 1

            if j < len(paragraphs):
                continuation = paragraphs[j].strip()
                # Check if this is a continuation ending with . (not a question)
                if continuation.endswith('.') and not continuation.endswith('?') and len(continuation) < 80:
                    # Check if next line after continuation is an option line
                    k = j + 1
                    while k < len(paragraphs):
                        np = paragraphs[k].strip()
                        if np:
                            break
                        k += 1
                    if k < len(paragraphs) and is_option_line_envie(paragraphs[k].strip()):
                        # Merge question parts and get options
                        combined_question = line + ' ' + continuation
                        opt_line = paragraphs[k].strip()
                        opts = extract_options_envie(opt_line)
                        if opts:
                            questions.append({
                                'question': combined_question,
                                'options': opts[:5],
                                '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                            })
                            i = k + 1
                            continue

            # Otherwise, collect paragraph options
            opts = []
            j = i + 1
            while j < len(paragraphs) and len(opts) < 3:
                opt_line = paragraphs[j].strip()
                if not opt_line:
                    j += 1
                    continue
                if opt_line.endswith('?') or is_option_line_envie(opt_line) or has_fill_blank(opt_line):
                    break
                if is_instruction_line(opt_line):
                    j += 1
                    continue
                opts.append(opt_line)
                j += 1

            if len(opts) == 3:
                questions.append({
                    'question': line,
                    'options': opts,
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i = j
                continue

        # Case 5a: Reading comprehension with C) marker on option 2 line
        # Format: "Abbreviations in the likes of BRB and LOL" + opt1 + "opt2\tC) opt3"
        # Skip numbered cloze format "22. A) ..."
        # Skip if line ends with '.' AND has commas (looks like option list, not question)
        looks_like_option = line.endswith('.') and ',' in line and len(line) < 60
        if len(line) > 20 and len(line) < 150 and not line.endswith('?') and not has_fill_blank(line) and not is_instruction_line(line) and not re.match(r'^(\d+)\.\s*A\s*\)', line) and not looks_like_option:
            j = i + 1
            opt1 = None
            while j < len(paragraphs):
                np = paragraphs[j].strip()
                if not np:
                    j += 1
                    continue
                opt1 = np
                break

            if opt1 and j + 1 < len(paragraphs):
                opt2_line = paragraphs[j + 1].strip() if j + 1 < len(paragraphs) else ""
                # Skip if opt2_line has B) marker - this is a full option line, not Case 5a format
                if not re.search(r'B\s*\)', opt2_line):
                    c_marker = re.search(r'\tC\s*\)|(\s{2,})C\s*\)', opt2_line)
                    if c_marker:
                        opt2 = opt2_line[:c_marker.start()].strip()
                        opt3_match = re.search(r'C\s*\)\s*(.+)$', opt2_line, re.IGNORECASE)
                        opt3 = opt3_match.group(1).strip() if opt3_match else ""

                        if opt2 and opt3:
                            questions.append({
                                'question': line,
                                'options': [opt1, opt2, opt3],
                                '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                            })
                            i = j + 2
                            continue

        # Case 5b-new: Grey Kangaroo Q1-5 reading comprehension (MUST come before 5b)
        # Format: stem line (no ? ending), then "opt1\tB) opt2" then "C) opt3"
        # Example: "Traditionally, child narrators are regarded as"
        #          "voices that employ a gloomy tone.\tB) innocent and genuine."
        #          "C) devoid of the depth needed to explore serious themes."
        if len(line) > 15 and len(line) < 80 and not line.endswith('?') and not line.endswith('.') and not has_fill_blank(line) and not is_instruction_line(line) and not re.match(r'^(\d+)\.\s*A\s*\)', line):
            # Check if next line has format "opt1\tB) opt2" (option A + B in one line)
            j = i + 1
            while j < len(paragraphs):
                np = paragraphs[j].strip()
                if not np:
                    j += 1
                    continue
                break

            if j < len(paragraphs):
                first_opt_line = paragraphs[j].strip()
                # Check for "opt1\tB) opt2" or "opt1  B) opt2" pattern
                b_match = re.search(r'(.+?)(?:\t|\s{2,})B\s*\)\s*(.+)$', first_opt_line)
                if b_match:
                    opt_a = b_match.group(1).strip()
                    opt_b = b_match.group(2).strip()

                    # Look for C) option on next line
                    j += 1
                    while j < len(paragraphs):
                        np = paragraphs[j].strip()
                        if not np:
                            j += 1
                            continue
                        break

                    opt_c = None
                    if j < len(paragraphs):
                        c_line = paragraphs[j].strip()
                        c_match = re.match(r'^C\s*\)\s*(.+)$', c_line, re.IGNORECASE)
                        if c_match:
                            opt_c = c_match.group(1).strip()
                            j += 1

                    if opt_a and opt_b and opt_c:
                        questions.append({
                            'question': line,
                            'options': [opt_a, opt_b, opt_c],
                            '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                        })
                        i = j
                        continue

        # Case 5b: Reading comprehension with 3 plain paragraph options (Q12-15 format)
        # Stem is short (< 70 chars), followed by 3 options without markers
        # Each option starts with lowercase (continuation of stem sentence)
        # Skip numbered cloze format "22. A) ..."
        if len(line) > 8 and len(line) < 70 and not line.endswith('?') and not line.endswith('.') and not has_fill_blank(line) and not is_instruction_line(line) and not re.match(r'^(\d+)\.\s*A\s*\)', line):
            # Collect next 3 non-empty paragraphs
            opts = []
            j = i + 1
            while j < len(paragraphs) and len(opts) < 3:
                np = paragraphs[j].strip()
                if not np:
                    j += 1
                    continue
                # Stop if we hit instruction or question-like line
                if is_instruction_line(np) or np.endswith('?') or is_option_line_envie(np):
                    break
                # Options should start with lowercase (continuation) or be short sentences
                if len(np) < 100:
                    opts.append(np)
                    j += 1
                else:
                    break

            if len(opts) == 3:
                questions.append({
                    'question': line,
                    'options': opts,
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i = j
                continue

        # Case 5c: Matching question - "Match each prefix..." or "Match the questions..." followed by table rows then options
        if 'match' in line.lower() and ('column' in line.lower() or 'prefix' in line.lower() or 'questions' in line.lower() or 'left' in line.lower()):
            # Skip table rows until we find options A) 1.../2.../
            j = i + 1
            table_rows = []
            while j < len(paragraphs):
                np = paragraphs[j].strip()
                if not np:
                    j += 1
                    continue
                # Check for option line "A) 1e/2b/..."
                if re.match(r'^A\s*\)\s*\d', np):
                    break
                # Collect table rows
                if '\t' in np or re.match(r'^\w+\s+[a-e]\.', np):
                    table_rows.append(np)
                j += 1

            # Extract options from option lines
            opts = []
            while j < len(paragraphs) and len(opts) < 5:
                np = paragraphs[j].strip()
                if not np:
                    j += 1
                    continue
                # Extract A-E options
                if re.match(r'^[A-E]\s*\)', np):
                    markers = list(re.finditer(r'([A-E])\s*\)', np, re.IGNORECASE))
                    for idx, m in enumerate(markers):
                        start = m.end()
                        if idx + 1 < len(markers):
                            end = markers[idx + 1].start()
                        else:
                            end = len(np)
                        opt_text = np[start:end].strip()
                        if opt_text:
                            opts.append(opt_text)
                    j += 1
                    continue
                break

            if len(opts) >= 3:
                questions.append({
                    'question': line,
                    'options': opts[:5],
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i = j
                continue

        # Case 5d: Question with options in image (no text options)
        # Format: "Pick the sentence with the correct punctuation." followed by another question
        # These questions have their options shown as images
        if line.endswith('.') and len(line) > 20 and len(line) < 80:
            # Check if this looks like a question (not a regular statement)
            lower_line = line.lower()
            is_question_like = any(w in lower_line for w in ['pick', 'choose', 'select', 'which', 'what'])
            # Check if next line is NOT an option line (options are in images)
            if is_question_like and next_line and not is_option_line_envie(next_line):
                # Check that next line is another question (not table data)
                next_is_question = len(next_line) > 20 and not re.match(r'^\w+\s+[a-e]\.', next_line)
                if next_is_question:
                    questions.append({
                        'question': line,
                        'options': ['[Option A in image]', '[Option B in image]', '[Option C in image]', '[Option D in image]', '[Option E in image]'],
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                    })
                    i += 1
                    continue

        # Case 6: Question + next line options with 5 choices (Wallaby Q31-50)
        # Format: "The Forbidden City is located in …" on one line
        # Then: "Beijing, China.\tB) Athens, Greece.   C) Rome..." on next line
        # Also handles multi-line options (D) E) on subsequent line)
        # Skip numbered cloze format "22. A) ..." which is handled by Case 7
        if len(line) > 15 and not is_instruction_line(line) and not re.match(r'^(\d+)\.\s*A\s*\)', line):
            if next_line and is_option_line_envie(next_line):
                opts = extract_options_envie(next_line)
                j = next_idx + 1
                # Check for continuation line (D) E) options)
                while j < len(paragraphs) and len(opts) < 5:
                    cont_line = paragraphs[j].strip()
                    if not cont_line:
                        j += 1
                        continue
                    # Check for D) E) continuation
                    if re.match(r'^[D-E]\s*\)', cont_line, re.IGNORECASE):
                        more_opts = extract_options_envie(cont_line)
                        if more_opts:
                            opts.extend(more_opts)
                            j += 1
                            continue
                    break
                if len(opts) >= 3:
                    # Check if this is a Word Groups format (multiple fill-blank sentences)
                    # Look back to find additional fill-blank paragraphs that belong to this question
                    # Strategy: Collect consecutive fill-blank paragraphs, stop at options line
                    question_parts = [line]
                    if has_fill_blank(line):
                        found_fill_blank_sequence = False
                        for back_idx in range(i - 1, max(0, i - 10), -1):
                            back_text = paragraphs[back_idx].strip()
                            if not back_text:
                                continue
                            # Stop if we hit instruction line
                            if is_instruction_line(back_text):
                                break
                            # Options line marks the boundary between questions
                            # Stop when we hit an options line (previous question's options)
                            if is_option_line_envie(back_text) and not has_fill_blank(back_text):
                                break
                            if has_fill_blank(back_text):
                                question_parts.insert(0, back_text)
                                found_fill_blank_sequence = True
                            else:
                                # Non-fill-blank, non-option line - stop
                                break

                    combined_question = ' '.join(question_parts)
                    questions.append({
                        'question': combined_question,
                        'options': opts[:5],
                        '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                    })
                    i = j
                    continue

        # Case 7: Cloze question in paragraph format
        # Format: "22.\tA) neither\tB) both\tC) either\tD) whether"
        cloze_match = re.match(r'^(\d+)\.\s*A\s*\)', line)
        if cloze_match:
            q_num = cloze_match.group(1)
            opts = extract_options_envie(line[cloze_match.start():])
            # If A) not extracted properly, re-extract starting from A)
            a_pos = re.search(r'A\s*\)', line)
            if a_pos:
                opts_text = line[a_pos.start():]
                markers = list(re.finditer(r'([A-E])\s*\)', opts_text, re.IGNORECASE))
                opts = []
                for idx, m in enumerate(markers):
                    start = m.end()
                    if idx + 1 < len(markers):
                        end = markers[idx + 1].start()
                    else:
                        end = len(opts_text)
                    opt_text = opts_text[start:end].strip().rstrip('\t ')
                    if opt_text:
                        opts.append(opt_text)
            if len(opts) >= 3:
                questions.append({
                    'question': f'Cloze question {q_num}',
                    'options': opts[:4],
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i += 1
                continue

        # Case 8: Cloze options without number (Q25-30 format)
        # Format: "A) yet\tB) whereas\tC) also\tD) while" - only options, no question number
        # These appear after numbered cloze questions in the same section
        # Each line is OPTIONS for ONE cloze question (not question + options from next line)
        if re.match(r'^A\s*\)', line) and '\t' in line:
            opts_text = line
            markers = list(re.finditer(r'([A-E])\s*\)', opts_text, re.IGNORECASE))
            opts = []
            for idx, m in enumerate(markers):
                start = m.end()
                if idx + 1 < len(markers):
                    end = markers[idx + 1].start()
                else:
                    end = len(opts_text)
                opt_text = opts_text[start:end].strip().rstrip('\t ')
                if opt_text:
                    opts.append(opt_text)
            if len(opts) >= 3:
                # This line IS the options for a cloze question (no question text, just options)
                # Question text is in the passage as a numbered blank like (26)
                questions.append({
                    'question': 'Cloze question (unnumbered)',
                    'options': opts[:4],
                    '_doc_pos': para_idx_to_doc_idx.get(i, i)  # Track document position
                })
                i += 1
                continue

        i += 1

    # Parse cloze questions from tables
    # Format: Table rows with "16.\tA) option" in first cell, B)/C)/D) in other cells
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not cells[0]:
                continue

            # Check if first cell has numbered question format "16.\tA) option"
            first_cell = cells[0]
            match = re.match(r'^(\d+)\.\s*A\s*\)\s*(.+)$', first_cell.replace('\t', ' '))
            if match:
                q_num = match.group(1)
                opt_a = match.group(2).strip()
                opts = [opt_a]

                # Extract B, C, D options from other cells
                for cell_text in cells[1:]:
                    opt_match = re.match(r'^([B-E])\s*\)\s*(.+)$', cell_text.strip())
                    if opt_match:
                        opts.append(opt_match.group(2).strip())

                if len(opts) >= 3:
                    questions.append({
                        'question': f'Cloze question {q_num}',
                        'options': opts[:4],
                        '_doc_pos': 9000  # Tables come later in document
                    })
                continue

            # Check for table with mixed question/options (Wallaby format)
            # Cell might have "Question?\nOption A." format
            for cell_text in cells:
                # Look for question ending with ? followed by option on new line
                if '?' in cell_text and '\n' in cell_text:
                    lines = cell_text.split('\n')
                    for li, line in enumerate(lines):
                        if line.strip().endswith('?'):
                            q_text = line.strip()
                            # Next line is option A
                            opts = []
                            if li + 1 < len(lines):
                                opt_a = lines[li + 1].strip().rstrip('.')
                                if opt_a:
                                    opts.append(opt_a)
                            # Get B, C options from other cells in same row
                            for other_cell in cells[1:]:
                                opt_match = re.match(r'^([B-E])\s*\)\s*(.+)$', other_cell.strip())
                                if opt_match:
                                    opts.append(opt_match.group(2).strip().rstrip('.'))

                            if q_text and len(opts) >= 3:
                                questions.append({
                                    'question': q_text,
                                    'options': opts[:5],
                                    '_doc_pos': 9000  # Tables come later in document
                                })
                            break

            # Note: Joey format (A) option cells with embedded question numbers)
            # is now handled by get_document_elements() table splitting

    # Parse cloze questions from textboxes and paragraphs
    # Textboxes may contain: "21. A) sparked B) spotted C) split D) squandered"
    from lxml import etree
    try:
        body = doc.element.body
        xml_str = etree.tostring(body, encoding='unicode')
        textbox_matches = re.findall(r'<w:txbxContent[^>]*>(.*?)</w:txbxContent>', xml_str, re.DOTALL)

        seen_cloze_nums = set()
        cloze_opts_from_textbox = {}  # Store options by question number

        for match in textbox_matches:
            texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', match)
            content = ' '.join(texts).strip()

            # Look for cloze option format: "21. A) option B) option C) option D) option"
            cloze_match = re.match(r'^(\d+)\.\s*A\s*\)\s*', content)
            if cloze_match:
                q_num = int(cloze_match.group(1))
                if q_num in seen_cloze_nums:
                    continue
                seen_cloze_nums.add(q_num)

                # Extract options
                markers = list(re.finditer(r'([A-D])\s*\)\s*', content))
                opts = []
                for idx, m in enumerate(markers):
                    start = m.end()
                    if idx + 1 < len(markers):
                        end = markers[idx + 1].start()
                    else:
                        end = len(content)
                    opt_text = content[start:end].strip()
                    if opt_text:
                        opts.append(opt_text)

                if len(opts) >= 3:
                    cloze_opts_from_textbox[q_num] = opts[:4]

        # Also check paragraphs for numbered cloze options (e.g., "22. A) neither B) both...")
        for para in doc.paragraphs:
            text = para.text.strip()
            opt_match = re.match(r'^(\d+)\.\s*A\s*\)', text)
            if opt_match:
                q_num = int(opt_match.group(1))
                if q_num not in cloze_opts_from_textbox:
                    # Extract options
                    markers = list(re.finditer(r'([A-D])\s*\)\s*', text))
                    opts = []
                    for idx, m in enumerate(markers):
                        start = m.end()
                        if idx + 1 < len(markers):
                            end = markers[idx + 1].start()
                        else:
                            end = len(text)
                        opt_text = text[start:end].strip()
                        if opt_text:
                            opts.append(opt_text)
                    if len(opts) >= 3:
                        cloze_opts_from_textbox[q_num] = opts[:4]

        # Collect unnumbered cloze option lines from paragraphs
        # These are lines like "A) yet B) whereas C) also D) while" in cloze section
        unnumbered_cloze_opts = []
        in_cloze_section = False
        for para in doc.paragraphs:
            text = para.text.strip()
            # Detect start of cloze section
            if 'space (21-30)' in text.lower() or '(21-30)' in text:
                in_cloze_section = True
                continue
            # Detect end of cloze section
            if in_cloze_section and ('questions 31-' in text.lower() or 'For questions 31' in text):
                in_cloze_section = False
                continue
            # Collect unnumbered option lines in cloze section
            if in_cloze_section and re.match(r'^A\s*\)', text):
                # This is an unnumbered cloze option line
                markers = list(re.finditer(r'([A-D])\s*\)\s*', text))
                opts = []
                for idx, m in enumerate(markers):
                    start = m.end()
                    if idx + 1 < len(markers):
                        end = markers[idx + 1].start()
                    else:
                        end = len(text)
                    opt_text = text[start:end].strip()
                    if opt_text:
                        opts.append(opt_text)
                if len(opts) >= 3:
                    unnumbered_cloze_opts.append(opts[:4])

        # Find cloze numbers in passages (from textboxes)
        cloze_in_passage = set()
        for match in textbox_matches:
            texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', match)
            content = ''.join(texts).strip()
            # Passage blanks: "(26)" or "(27)"
            blanks = re.findall(r'\((\d+)\)', content)
            for b in blanks:
                cloze_in_passage.add(int(b))

        # Determine missing cloze numbers (numbers in passage but no options)
        cloze_with_options = set(cloze_opts_from_textbox.keys())
        missing_cloze = sorted(cloze_in_passage - cloze_with_options)

        # Assign unnumbered options to missing cloze numbers
        if missing_cloze and unnumbered_cloze_opts:
            for i, cloze_num in enumerate(missing_cloze):
                if i < len(unnumbered_cloze_opts):
                    cloze_opts_from_textbox[cloze_num] = unnumbered_cloze_opts[i]

        # Also handle cloze 30 which may be the last unnumbered option
        if 30 not in cloze_opts_from_textbox and unnumbered_cloze_opts:
            # Use the last available unnumbered option for cloze 30
            remaining_unnumbered = len(unnumbered_cloze_opts) - len(missing_cloze)
            if remaining_unnumbered > 0:
                cloze_opts_from_textbox[30] = unnumbered_cloze_opts[-1]

        # Add all cloze questions
        for q_num in sorted(cloze_opts_from_textbox.keys()):
            opts = cloze_opts_from_textbox[q_num]
            questions.append({
                'question': f'Cloze question {q_num}',
                'options': opts,
                '_doc_pos': 8000 + q_num  # Textbox cloze, position by question number
            })

        # Add placeholder for any remaining missing cloze numbers (image-based)
        all_cloze_nums = set(cloze_opts_from_textbox.keys())
        still_missing = cloze_in_passage - all_cloze_nums
        if all_cloze_nums:
            max_cloze = max(all_cloze_nums)
            for cloze_num in sorted(still_missing):
                if 20 <= cloze_num <= max_cloze:
                    questions.append({
                        'question': f'Cloze question {cloze_num} (image-based)',
                        'options': ['[Option A in image]', '[Option B in image]', '[Option C in image]', '[Option D in image]'],
                        '_doc_pos': 8000 + cloze_num  # Position by question number
                    })

    except Exception:
        pass  # lxml may not be available

    # Remove incorrectly parsed "A) ..." lines as questions
    questions = [q for q in questions if not q.get('question', '').startswith('A)')]

    # Remove duplicate Cloze questions (keep ones with passage/answer)
    # Group by Cloze number and keep the best version
    cloze_by_num = {}
    non_cloze = []
    for q in questions:
        text = q.get('question', '')
        m = re.search(r'Cloze question (\d+)', text)
        if m:
            q_num = int(m.group(1))
            existing = cloze_by_num.get(q_num)
            # Prefer version with passage (longer text) or with answer
            if existing is None:
                cloze_by_num[q_num] = q
            else:
                # Keep the one with more content or answer
                existing_len = len(existing.get('question', ''))
                new_len = len(text)
                existing_has_ans = bool(existing.get('answer'))
                new_has_ans = bool(q.get('answer'))
                if (new_has_ans and not existing_has_ans) or (new_len > existing_len and not existing_has_ans):
                    cloze_by_num[q_num] = q
        else:
            non_cloze.append(q)

    # Rebuild questions list
    questions = non_cloze + list(cloze_by_num.values())

    # Sort cloze questions by their question number and insert at correct position
    def get_cloze_num(q):
        text = q.get('question', '')
        m = re.search(r'Cloze question (\d+)', text)
        if m:
            return int(m.group(1))
        return 999

    # Separate cloze questions from others
    cloze_questions = [q for q in questions if 'Cloze question' in q.get('question', '')]
    other_questions = [q for q in questions if 'Cloze question' not in q.get('question', '')]

    # Sort other_questions by document position to maintain correct order
    other_questions.sort(key=lambda q: q.get('_doc_pos', 9999))

    # Sort cloze questions by their number
    cloze_questions.sort(key=get_cloze_num)

    # Insert cloze questions at correct position based on cloze number
    # E.g., "Cloze question 21" should be inserted at position 20 (0-indexed)
    if cloze_questions:
        # Get the first cloze question number to determine insertion position
        first_cloze_num = get_cloze_num(cloze_questions[0])
        # Insert at position = first_cloze_num - 1 (0-indexed)
        # E.g., Cloze Q11 -> insert at index 10, Cloze Q21 -> insert at index 20
        insert_idx = first_cloze_num - 1
        # Make sure we don't go beyond the available questions
        insert_idx = min(insert_idx, len(other_questions))
        questions = other_questions[:insert_idx] + cloze_questions + other_questions[insert_idx:]
    else:
        questions = other_questions

    # Deduplicate questions (same question text)
    seen_texts = set()
    unique_questions = []
    for q in questions:
        qt = q.get('question', '').strip()
        if qt and qt not in seen_texts:
            seen_texts.add(qt)
            unique_questions.append(q)
        elif not qt:
            unique_questions.append(q)
    questions = unique_questions

    # Clean up _doc_pos from final output
    for q in questions:
        q.pop('_doc_pos', None)

    return questions


# Note: _dedup_bilingual_science has been moved to envie-bilingual-dedup-helper.py
# Import via: from app.parsers import _dedup_bilingual_science
